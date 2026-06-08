import os
import sqlite3
import logging
import json
import re
from aiogram import Bot, Dispatcher, types, F, BaseMiddleware
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode, ChatType
from aiogram.filters import Command
from aiogram.utils.chat_action import ChatActionSender
from aiogram.types import TelegramObject
import httpx
import asyncio
from bs4 import BeautifulSoup
from datetime import datetime
import time
import uuid
import base64
from io import BytesIO
from contextvars import ContextVar
import tempfile
import math






def write_research_log(chat_id: int, content: str):
    """Дописать блок пересказов в физический файл data/research_log_{chat_id}.txt."""
    if chat_id is None:
        return
    os.makedirs("data", exist_ok=True)
    file_path = f"data/research_log_{chat_id}.txt"
    try:
        with open(file_path, "a", encoding="utf-8") as f:
            f.write(content + "\n\n")
        logger.info(f"[RESEARCH-LOG] Записан блок в файл: {file_path}")
    except Exception as e:
        logger.error(f"[RESEARCH-LOG] Ошибка записи в файл {file_path}: {repr(e)}")


def read_research_log(chat_id: int) -> str:
    """Считать накопленный конспект из файла data/research_log_{chat_id}.txt."""
    if chat_id is None:
        return ""
    file_path = f"data/research_log_{chat_id}.txt"
    if not os.path.exists(file_path):
        return ""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception as e:
        logger.error(f"[RESEARCH-LOG] Ошибка чтения файла {file_path}: {repr(e)}")
        return ""


def clear_research_log(chat_id: int):
    """Очистить файл исследования при старте."""
    if chat_id is None:
        return
    file_path = f"data/research_log_{chat_id}.txt"
    if os.path.exists(file_path):
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("")
            logger.info(f"[RESEARCH-LOG] Очищен файл исследования: {file_path}")
        except Exception as e:
            logger.error(f"[RESEARCH-LOG] Ошибка очистки файла {file_path}: {repr(e)}")


DB_PATH = "/app/data/search_sessions.db"

def init_db():
    try:
        os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS search_sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                chat_id INTEGER NOT NULL,
                title TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                research_log TEXT,
                chat_history TEXT
            );
        """)
        conn.commit()
        conn.close()
        logger.info("[DB] Инициализация SQLite базы данных сессий успешна.")
    except Exception as e:
        logger.error(f"[DB] Ошибка инициализации базы данных: {e}")

async def save_session_to_db(user_id: int, chat_id: int, title: str, research_log: str, chat_history: list) -> int:
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        history_json = json.dumps(chat_history, ensure_ascii=False)
        cursor.execute(
            "INSERT INTO search_sessions (user_id, chat_id, title, research_log, chat_history) VALUES (?, ?, ?, ?, ?)",
            (user_id, chat_id, title, research_log, history_json)
        )
        conn.commit()
        last_id = cursor.lastrowid
        conn.close()
        logger.info(f"[DB] Сессия '{title}' успешно сохранена для пользователя {user_id} (ID: {last_id}).")
        return last_id
    except Exception as e:
        logger.error(f"[DB] Ошибка сохранения сессии: {e}")
        return 0

def get_user_sessions(user_id: int, limit: int = 5, offset: int = 0) -> list:
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, title, created_at FROM search_sessions WHERE user_id = ? ORDER BY created_at DESC LIMIT ? OFFSET ?",
            (user_id, limit, offset)
        )
        sessions = cursor.fetchall()
        conn.close()
        return sessions
    except Exception as e:
        logger.error(f"[DB] Ошибка получения сессий: {e}")
        return []

def get_user_sessions_count(user_id: int) -> int:
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM search_sessions WHERE user_id = ?", (user_id,))
        count = cursor.fetchone()[0]
        conn.close()
        return count
    except Exception as e:
        logger.error(f"[DB] Ошибка получения количества сессий: {e}")
        return 0

def get_session_by_id(session_id: int) -> tuple:
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, user_id, chat_id, title, created_at, research_log, chat_history FROM search_sessions WHERE id = ?",
            (session_id,)
        )
        session = cursor.fetchone()
        conn.close()
        return session
    except Exception as e:
        logger.error(f"[DB] Ошибка получения сессии по ID: {e}")
        return None

def delete_session_by_id(session_id: int):
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM search_sessions WHERE id = ?", (session_id,))
        conn.commit()
        conn.close()
        logger.info(f"[DB] Сессия ID {session_id} удалена.")
    except Exception as e:
        logger.error(f"[DB] Ошибка удаления сессии: {e}")

async def generate_session_title(chat_history: list, model_name: str = None) -> str:
    if not chat_history:
        return "Исследование без вопросов"
        
    dialogue_parts = []
    for msg in chat_history:
        role = "Пользователь" if msg.get("role") == "user" else "ИИ-Помощник"
        content = msg.get("content", "")
        content_snippet = content[:300] + "..." if len(content) > 300 else content
        dialogue_parts.append(f"{role}: {content_snippet}")
        
    dialogue_text = "\n".join(dialogue_parts)
    
    system_prompt = (
        "Ты — генератор коротких и емких заголовков для архива поисковых сессий.\n"
        "На основе предоставленного диалога пользователя и поискового бота сформулируй ОДИН короткий заголовок (не более 3-5 слов), отражающий главную суть исследования.\n"
        "Ответь СТРОГО заголовком, без кавычек, вводных слов, пояснений и знаков препинания в конце."
    )
    
    try:
        title = await ask_ollama(
            f"Диалог:\n{dialogue_text}\n\nСформулируй короткий заголовок:",
            system_prompt=system_prompt,
            model_name=model_name,
            timeout=10.0
        )
        if "<think>" in title and "</think>" in title:
            title = title.split("</think>")[-1].strip()
        cleaned_title = title.strip().replace('"', '').replace("'", "")
        if cleaned_title and len(cleaned_title) < 80:
            return cleaned_title
    except Exception as e:
        logger.warning(f"[TITLE-GEN] Ошибка генерации заголовка: {e}")
        
    for msg in chat_history:
        if msg.get("role") == "user":
            first_query = msg.get("content", "")
            first_query = get_clean_fallback_query(first_query)
            if len(first_query) > 40:
                return first_query[:40].strip() + "..."
            return first_query.strip()
            
    return f"Исследование от {datetime.now().strftime('%d.%m.%Y %H:%M')}"


def make_summary_status_text(stage_name: str, summary_states: dict) -> str:
    lines = [f"<b>{stage_name}:</b>"]
    for q, state in summary_states.items():
        short_q = q[:40] + "..." if len(q) > 40 else q
        lines.append(f"  • <code>{short_q}</code> — {state}")
    return "\n".join(lines)


async def generate_query_summary(
    query_str: str,
    user_query: str,
    sites: list[dict],
    model_name: str,
    stage_name: str,
    is_opinion: bool = False,
    chat_id: int = None,
    status_updater = None,
    summary_states: dict = None,
    progress_percent: int = 35,
    iteration: int = 1,
    current_thoughts: str = None,
    search_queries_state: list = None
) -> str:
    async def update_state(state: str):
        if summary_states is not None and status_updater is not None:
            summary_states[query_str] = state
            status_text = make_summary_status_text(stage_name, summary_states)
            try:
                await status_updater(
                    make_status_html(
                        query=user_query,
                        status_text=status_text,
                        progress_percent=progress_percent,
                        search_queries=search_queries_state,
                        attempt=iteration,
                        thoughts=current_thoughts
                    )
                )
            except Exception as e:
                logger.error(f"[QUERY-SUMMARY] Не удалось обновить статус: {e}")

    await update_state("⏳ конспектируется...")

    if not sites:
        await update_state("⚠️ нет полезных материалов")
        return f"По запросу '{query_str}' полезные материалы не найдены."

    # Берем от 3 до 5 сайтов (пользовательское ограничение)
    selected_sites = sites[:5]
    if len(selected_sites) < 3:
        logger.warning(f"[QUERY-SUMMARY] Для запроса '{query_str}' найдено меньше 3 сайтов (всего: {len(selected_sites)}).")

    # Формируем объединенный текст сайтов для этого запроса
    parts = []
    for s in selected_sites:
        parts.append(f"--- Источник: {s['url']} (Дата: {s['publish_date'] or 'неизвестно'}) ---\n{s['content']}\n")
    combined_text = "\n".join(parts)

    current_date = "2026-06-03" # Фиксированная сегодняшняя дата согласно метаданным
    
    # Формируем системный промпт в зависимости от типа этапа
    if is_opinion:
        system_prompt = (
            f"Текущая дата: {current_date}.\n"
            "Ты — аналитик данных, исследующий мнения, отзывы и личный опыт пользователей.\n"
            "Твоя задача — составить подробный конспект исключительно мнений людей, отзывов и обсуждений по теме запроса.\n"
            "КРИТИЧЕСКИ ВАЖНЫЕ ПРАВИЛА:\n"
            "1. Сконцентрируйся только на том, что говорят реальные люди (форумы, блоги, комментарии, обзоры).\n"
            f"2. Делай жесткий акцент на самых свежих и последних данных (актуальных на {current_date}).\n"
            "3. Если ты сомневаешься в свежести данных, все равно перескажи их.\n"
            "4. Вся информация, которая не касается запроса пользователя или не представляет собой мнение людей, должна быть ПОЛНОСТЬЮ проигнорирована и отсечена.\n"
            "5. Пиши тезисно, подробно, объективно, без вводных слов и воды. Твой ответ должен быть СТРОГО на русском языке, даже если исходные материалы написаны на английском или других языках."
        )
    else:
        system_prompt = (
            f"Текущая дата: {current_date}.\n"
            "Ты — аналитик данных. Твоя задача — составить подробный и информативный конспект предоставленного текста по конкретному поисковому направлению.\n"
            "КРИТИЧЕСКИ ВАЖНЫЕ ПРАВИЛА:\n"
            "1. Извлеки всю полезную информацию, которая СТРОГО относится к теме запроса пользователя.\n"
            f"2. Делай обязательный акцент на самых последних, актуальных и свежих данных (на {current_date}).\n"
            "3. Всю информацию, которая в целом не касается запроса пользователя или текущего этапа исследования, полностью проигнорируй (не включай ее в конспект вообще).\n"
            "4. Пиши тезисно, подробно, сухо, без воды, мета-сообщений и вводных фраз. Твой ответ должен быть СТРОГО на русском языке, даже если исходные материалы написаны на английском или других языках."
        )

    prompt = (
        f"Этап исследования: {stage_name}\n"
        f"Поисковый запрос: {query_str}\n"
        f"Общий запрос пользователя: {user_query}\n\n"
        f"Материалы с найденных сайтов:\n{combined_text}\n\n"
        "Сделай подробный структурированный пересказ материалов строго по теме запроса пользователя, следуя системным правилам."
    )

    try:
        summary = await ask_ollama(prompt, system_prompt=system_prompt, model_name=model_name, timeout=40.0)
        if "<think>" in summary and "</think>" in summary:
            summary = summary.split("</think>")[-1].strip()
        
        # Добавляем инфо в физический лог
        log_block = (
            f"=== НАПРАВЛЕНИЕ ПОИСКА: {query_str} ===\n"
            f"Найдено сайтов: {len(selected_sites)}\n"
            f"Конспект материалов:\n{summary}\n"
            f"======================================"
        )
        write_research_log(chat_id, log_block)
        
        await update_state("✅ готово")
        return summary
    except Exception as e:
        logger.error(f"[QUERY-SUMMARY] Ошибка конспектирования для запроса '{query_str}': {repr(e)}")
        fallback = f"Ошибка конспектирования направления '{query_str}'."
        write_research_log(chat_id, f"=== НАПРАВЛЕНИЕ ПОИСКА: {query_str} ===\n[СБОЙ ИИ-КОНСПЕКТИРОВАНИЯ]\n======================================")
        await update_state("❌ ошибка")
        return fallback


# Для совместимости импортов документов
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import xlrd
except ImportError:
    xlrd = None

current_model_var = ContextVar('current_model_var', default=None)
current_status_updater_var = ContextVar('current_status_updater_var', default=None)
current_query_var = ContextVar('current_query_var', default="")
current_progress_percent_var = ContextVar('current_progress_percent_var', default=0)
current_search_queries_var = ContextVar('current_search_queries_var', default=None)
current_attempt_var = ContextVar('current_attempt_var', default=1)
current_thoughts_var = ContextVar('current_thoughts_var', default=None)
current_status_text_var = ContextVar('current_status_text_var', default="")



# Глобальные списки и кэши для мультимедиа и файлов
MULTIMODAL_MODELS = []
PREFERRED_BACKUP_MODELS = []
PREFERRED_VISION_MODELS = []

def clean_multimodal_messages(messages: list) -> list:
    cleaned = []
    for msg in messages:
        role = msg.get("role")
        content = msg.get("content")
        if isinstance(content, list):
            text_parts = []
            for part in content:
                if isinstance(part, dict):
                    if part.get("type") == "text":
                        text_parts.append(part.get("text", ""))
                else:
                    text_parts.append(str(part))
            cleaned.append({"role": role, "content": " ".join(text_parts).strip()})
        else:
            cleaned.append({"role": role, "content": content})
    return cleaned


def get_model_for_attempt(start_model: str, attempt: int, all_api_models: list[str]) -> str:
    if attempt == 1:
        return start_model or OLLAMA_MODEL
        
    fallback_chain = []
    for bm in PREFERRED_BACKUP_MODELS:
        if bm != start_model and bm in all_api_models:
            fallback_chain.append(bm)
            
    for am in all_api_models:
        if am != start_model and am not in fallback_chain:
            fallback_chain.append(am)
            
    if not fallback_chain:
        return start_model or OLLAMA_MODEL
        
    idx = attempt - 2
    if idx < len(fallback_chain):
        return fallback_chain[idx]
    return fallback_chain[-1]


media_groups_data = {}


# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Инициализация переменных окружения
BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "deepseek-v4-flash")
OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY", "")
SEARXNG_URL = os.getenv("SEARXNG_URL", "http://127.0.0.1:8080")

if not BOT_TOKEN:
    raise ValueError("TELEGRAM_BOT_TOKEN не задан в переменных окружения!")

bot = Bot(token=BOT_TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))
dp = Dispatcher()

class TelegramUpdateManager:
    def __init__(self, bot: Bot):
        self.bot = bot
        # Хранит отложенные обновления: { key: {"target": target, "text": str, "reply_markup": markup} }
        self.pending_updates = {}
        # Хранит таймстамп последней успешной отправки: { key: float }
        self.last_send_times = {}
        self.worker_task = None
        self.lock = asyncio.Lock()

    def start(self):
        """Запуск фоновой задачи обновления"""
        self.worker_task = asyncio.create_task(self._worker_loop())

    def _get_key(self, target) -> tuple | str:
        if isinstance(target, str):
            return target
        elif isinstance(target, types.Message):
            return (target.chat.id, target.message_id)
        else:
            raise ValueError("Неверный тип целевого сообщения")

    async def request_update(self, target, text: str, reply_markup=None, is_final=False):
        """Регистрация запроса на обновление текста"""
        key = self._get_key(target)
        
        if is_final:
            async with self.lock:
                self.pending_updates.pop(key, None)
                await self._send_to_telegram(target, text, reply_markup)
                self.last_send_times[key] = time.time()
                # Чистим после отправки финала через некоторое время
                self.last_send_times.pop(key, None)
            return

        async with self.lock:
            self.pending_updates[key] = {
                "target": target,
                "text": text,
                "reply_markup": reply_markup
            }

    async def _worker_loop(self):
        """Периодический опрос очереди отложенных обновлений (раз в 0.5 сек)"""
        while True:
            try:
                await asyncio.sleep(0.5)
                now = time.time()
                
                to_update = []
                async with self.lock:
                    for key, data in list(self.pending_updates.items()):
                        last_sent = self.last_send_times.get(key, 0.0)
                        if now - last_sent >= 2.0:
                            to_update.append((key, data))
                            
                for key, data in to_update:
                    success = await self._send_to_telegram(data["target"], data["text"], data["reply_markup"])
                    async with self.lock:
                        if success:
                            self.last_send_times[key] = now
                            if key in self.pending_updates and self.pending_updates[key]["text"] == data["text"]:
                                self.pending_updates.pop(key, None)
            except Exception as e:
                logger.error(f"[UPDATE-MANAGER] Критическая ошибка в worker_loop: {e}")

    async def _send_to_telegram(self, target, text: str, reply_markup) -> bool:
        """Реальный вызов Telegram API с обработкой ошибок"""
        try:
            if isinstance(target, str):
                await self.bot.edit_message_text(
                    text=text,
                    inline_message_id=target,
                    parse_mode=ParseMode.HTML,
                    reply_markup=reply_markup,
                    disable_web_page_preview=True
                )
            else:
                await target.edit_text(
                    text=text,
                    parse_mode=ParseMode.HTML,
                    reply_markup=reply_markup,
                    disable_web_page_preview=True
                )
            return True
        except Exception as e:
            if "message is not modified" in str(e):
                return True
            logger.warning(f"[UPDATE-MANAGER] Ошибка обновления сообщения: {e}")
            return False

update_manager = TelegramUpdateManager(bot)

# Глобальный словарь для хранения неограниченной истории диалогов
user_histories = {}
# Глобальный словарь для отслеживания сообщений очистки чата
last_clear_message_ids = {}

# Кэши для обхода лимита в 64 символа у Telegram API
inline_queries_cache = {}   # {cache_key: query}
private_queries_cache = {}  # {cache_key: {"query": query}}
awaiting_photo_text = {}    # {user_id: cache_key}
active_search_tasks = {}    # {task_key: asyncio.Task}


# Глобальный словарь моделей пользователей {user_id: model_name}
user_models = {}
# Глобальный словарь для отслеживания активных сессий пользователей {user_id: session_id}
active_session_ids = {}
USER_MODELS_FILE = "/app/data/user_models.json"
USER_HISTORIES_FILE = "/app/data/user_histories.json"

# Семафор для предотвращения параллельного ранжирования сайтов и перегрузки ИИ-шлюза
ranking_semaphore = asyncio.Semaphore(1)

def load_user_models():
    global user_models
    path = USER_MODELS_FILE
    if not os.path.exists(path):
        # Попробуем локальный путь, если мы вне докера
        path = "./user_models.json"
    
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                # Приводим ключи к int для совместимости в runtime
                user_models = {int(k): v for k, v in data.items() if k.isdigit()}
                logger.info(f"Успешно загружены настройки моделей для {len(user_models)} пользователей.")
        except Exception as e:
            logger.error(f"Ошибка при загрузке user_models из {path}: {e}")
            user_models = {}
    else:
        user_models = {}

def save_user_models():
    path = USER_MODELS_FILE
    # Если директория /app/data не существует, сохраняем в текущую
    dir_path = os.path.dirname(path)
    if dir_path and not os.path.exists(dir_path):
        path = "./user_models.json"
    
    try:
        # Для сохранения приводим ключи к строкам (JSON поддерживает только строковые ключи)
        data = {str(k): v for k, v in user_models.items()}
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        logger.info(f"Настройки моделей пользователей сохранены в {path}")
    except Exception as e:
        logger.error(f"Ошибка при сохранении user_models в {path}: {e}")

def load_user_histories():
    global user_histories
    path = USER_HISTORIES_FILE
    if not os.path.exists(path):
        path = "./user_histories.json"
    
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                user_histories = {int(k): v for k, v in data.items() if k.isdigit()}
                logger.info(f"Успешно загружена история диалогов для {len(user_histories)} пользователей.")
        except Exception as e:
            logger.error(f"Ошибка при загрузке user_histories из {path}: {e}")
            user_histories = {}
    else:
        user_histories = {}

def save_user_histories():
    path = USER_HISTORIES_FILE
    dir_path = os.path.dirname(path)
    if dir_path and not os.path.exists(dir_path):
        path = "./user_histories.json"
    
    try:
        data = {str(k): v for k, v in user_histories.items()}
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        logger.info(f"История диалогов пользователей сохранена в {path}")
    except Exception as e:
        logger.error(f"Ошибка при сохранении user_histories в {path}: {e}")

# Кэш мультимодальных моделей
multimodal_cache = {}
MULTIMODAL_CACHE_FILE = "/app/data/multimodal_cache.json"

def load_multimodal_cache():
    global multimodal_cache
    path = MULTIMODAL_CACHE_FILE
    if not os.path.exists(path):
        path = "./multimodal_cache.json"
    
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                multimodal_cache = json.load(f)
                logger.info(f"Успешно загружен кэш мультимодальных моделей: {len(multimodal_cache)} записей.")
        except Exception as e:
            logger.error(f"Ошибка при загрузке multimodal_cache из {path}: {e}")
            multimodal_cache = {}
    else:
        multimodal_cache = {}

def save_multimodal_cache():
    path = MULTIMODAL_CACHE_FILE
    dir_path = os.path.dirname(path)
    if dir_path and not os.path.exists(dir_path):
        path = "./multimodal_cache.json"
    
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(multimodal_cache, f, ensure_ascii=False, indent=4)
        logger.info(f"Кэш мультимодальных моделей сохранен в {path}")
    except Exception as e:
        logger.error(f"Ошибка при сохранении multimodal_cache в {path}: {e}")

# Объявим семафор для проверок мультимодальности (максимум 2 параллельных запроса)
multimodal_check_semaphore = asyncio.Semaphore(2)

async def check_model_multimodal(model_name: str) -> bool:
    # Изображение 64x64 красного цвета в формате JPEG
    sample_jpeg_b64 = (
        "/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAgGBgcGBQgHBwcJCQgKDBQNDAsLDBkSEw8UHRofH"
        "h0aHBwgJC4nICIsIxwcKDcpLDAxNDQ0Hyc5PTgyPC4zNDL/2wBDAQkJCQwLDBgNDRgyIRwhMj"
        "IyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjL/wAARCABA"
        "AEADASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgED"
        "AwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcY"
        "GRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJ"
        "ipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo"
        "6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgEC"
        "BAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl"
        "8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaH"
        "iImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn"
        "6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwDi6KKK+ZP3EKKKKACiiigAooooAKKKKACiiigA"
        "ooooA//Z"
    )
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/chat/completions"
    headers = {"Content-Type": "application/json"}
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
        
    payload = {
        "model": model_name,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Is this image red?"},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{sample_jpeg_b64}"}}
                ]
            }
        ],
        "temperature": 0.1
    }
    
    async with multimodal_check_semaphore:
        retry_delay = 2.0
        for attempt in range(1, 5):
            try:
                async with httpx.AsyncClient(timeout=15.0) as client:
                    logger.info(f"[MULTIMODAL-CHECK] Проверка модели {model_name} (Попытка {attempt}/4)...")
                    resp = await client.post(url, json=payload, headers=headers)
                    if resp.status_code == 200:
                        logger.info(f"[MULTIMODAL-CHECK] Модель {model_name} поддерживает изображения (HTTP 200).")
                        return True
                    elif resp.status_code == 429:
                        logger.warning(f"[MULTIMODAL-CHECK] Модель {model_name} вернула 429. Повтор через {retry_delay} сек...")
                        await asyncio.sleep(retry_delay)
                        retry_delay *= 2.0
                        continue
                    else:
                        logger.info(f"[MULTIMODAL-CHECK] Модель {model_name} НЕ поддерживает изображения (HTTP {resp.status_code}).")
                        return False
            except Exception as e:
                logger.warning(f"[MULTIMODAL-CHECK] Сетевая ошибка при проверке {model_name} (Попытка {attempt}/4): {e}")
                if attempt == 4:
                    raise e
                await asyncio.sleep(retry_delay)
                retry_delay *= 2.0
        return False

def is_model_multimodal(model_name: str) -> bool:
    return multimodal_cache.get(model_name, False)

# Инициализируем при запуске
load_user_models()
load_user_histories()
load_multimodal_cache()

class LoggingMiddleware(BaseMiddleware):
    async def __call__(self, handler, event: TelegramObject, data):
        print(f"RAW UPDATE PRINT: {event.model_dump_json(exclude_none=True)}", flush=True)
        return await handler(event, data)

dp.message.outer_middleware(LoggingMiddleware())
dp.inline_query.outer_middleware(LoggingMiddleware())
dp.chosen_inline_result.outer_middleware(LoggingMiddleware())

def get_current_datetime_str() -> str:
    return datetime.now().strftime("%d.%m.%Y %H:%M:%S")

def clean_model_answer(answer: str) -> str:
    text = answer.strip()
    prefixes = [
        "ответ на запрос:",
        "ответ на вопрос:",
        "ответ пользователя:",
        "ответ:",
        "запрос:",
        "вывод:"
    ]
    changed = True
    while changed:
        changed = False
        lower_text = text.lower()
        for prefix in prefixes:
            if lower_text.startswith(prefix):
                text = text[len(prefix):].strip()
                changed = True
                break
    if text.startswith('"') and text.endswith('"'):
        text = text[1:-1].strip()
    elif text.startswith("'") and text.endswith("'"):
        text = text[1:-1].strip()
    return text

def clean_markdown_from_thoughts(text: str) -> str:
    if not text:
        return ""
    # Удаляем заголовки (символы # в начале строк)
    t = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    # Удаляем ссылки вида [text](url) -> оставляем только text
    t = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', t)
    # Удаляем жирный шрифт, курсив, зачеркивание, моноширинный, спойлеры: **, *, __, _, `, ~, ||
    t = re.sub(r'\*\*|__|\*|_|`|~|\|\|', '', t)
    # Удаляем маркеры списков в начале строк (- , * , + , •)
    t = re.sub(r'^\s*[-*+•]\s+', '', t, flags=re.MULTILINE)
    # Схлопываем лишние переводы строк
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()

def auto_close_html_tags(html: str) -> str:
    tag_re = re.compile(r'</?([a-zA-Z0-9_-]+)(?:\s+[^>]*)?>')
    open_tags = []
    
    for match in tag_re.finditer(html):
        full_tag = match.group(0)
        tag_name = match.group(1).lower()
        
        if full_tag.startswith('</'):
            if open_tags and open_tags[-1] == tag_name:
                open_tags.pop()
        else:
            open_tags.append(tag_name)
            
    closed_html = html
    for tag_name in reversed(open_tags):
        closed_html += f"</{tag_name}>"
        
    return closed_html

def convert_markdown_to_html(text: str) -> str:
    if not text:
        return ""
    
    # 1. Заменяем списки в начале строк (* , - , +) на юникод-маркеры (•)
    text = re.sub(r'(?m)^\s*[*+-]\s+', '• ', text)

    # 2. Заменяем заголовки #, ##, ### в начале строк на жирный текст
    text = re.sub(r'(?m)^#+\s+(.*?)$', r'<b>\1</b>', text)

    # 2.5. Преобразование блоков цитат Markdown: > текст -> <blockquote>текст</blockquote>
    def replace_blockquote(match):
        block = match.group(0)
        lines = [re.sub(r'^\s*>\s*', '', line) for line in block.split('\n')]
        content = "\n".join(lines).strip()
        return f"<blockquote>{content}</blockquote>"
    text = re.sub(r'(?m)(^\s*>.*(?:\n\s*>.*)*)', replace_blockquote, text)

    # 3. Выделяем блоки кода с тройными кавычками
    code_blocks = []
    def replace_code_block(match):
        code = match.group(2)
        placeholder = f"CodeBlockPlaceholder{len(code_blocks)}"
        code_blocks.append(code)
        return placeholder

    text = re.sub(r'```([a-zA-Z0-9+#-_]*)\n?(.*?)\n?```', replace_code_block, text, flags=re.DOTALL)

    # 4. Выделяем инлайновый код `code`
    inline_codes = []
    def replace_inline_code(match):
        code = match.group(1)
        placeholder = f"InlineCodePlaceholder{len(inline_codes)}"
        inline_codes.append(code)
        return placeholder
    
    text = re.sub(r'`(.*?)`', replace_inline_code, text)

    # 5. Спойлеры: ||text|| -> <tg-spoiler>text</tg-spoiler>
    text = re.sub(r'\|\|(.*?)\|\|', r'<tg-spoiler>\1</tg-spoiler>', text)

    # 6. Жирный шрифт: **text** -> <b>text</b> и __text__ -> <b>text</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.*?)__', r'<b>\1</b>', text)

    # 7. Курсив: *text* -> <i>text</i> и _text_ -> <i>text</i>
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    text = re.sub(r'_(.*?)_', r'<i>\1</i>', text)

    # 8. Зачеркивание: ~~text~~ -> <s>text</s>
    text = re.sub(r'~~(.*?)~~', r'<s>\1</s>', text)

    # 8.5. Ссылки: [text](url) -> <a href="url">text</a>
    text = re.sub(r'\[([^\]\n]+)\]\(([^)\n]+)\)', r'<a href="\2">\1</a>', text)

    # 9. Возвращаем инлайновый код, экранируя спецсимволы
    for i, code in enumerate(inline_codes):
        escaped_code = code.replace('<', '&lt;').replace('>', '&gt;')
        text = text.replace(f"InlineCodePlaceholder{i}", f"<code>{escaped_code}</code>")

    # 10. Возвращаем блоки кода, экранируя спецсимволы
    for i, code in enumerate(code_blocks):
        escaped_code = code.replace('<', '&lt;').replace('>', '&gt;')
        text = text.replace(f"CodeBlockPlaceholder{i}", f"<pre>{escaped_code}</pre>")

    return text

def safe_html_cleaner(text: str) -> str:
    # Сначала конвертируем Markdown в HTML
    text = convert_markdown_to_html(text)
    
    # Заменяем HTML-списки на текстовые маркеры (Telegram HTML не поддерживает списки)
    text = re.sub(r'(?i)<li>', '• ', text)
    text = re.sub(r'(?i)</li>', '\n', text)
    text = re.sub(r'(?i)</?ul>', '', text)
    text = re.sub(r'(?i)</?ol>', '', text)
    
    # Заменяем <br> и <br/> на обычные новые строки
    text = re.sub(r'(?i)<br\s*/?>', '\n', text)
    text = re.sub(r'&(?!([a-zA-Z]+|#[0-9]+);)', '&amp;', text)
    allowed_tags = ['b', 'strong', 'i', 'em', 'u', 'ins', 's', 'strike', 'del', 'code', 'pre', 'blockquote', 'tg-spoiler', 'a']
    tag_re = re.compile(r'</?([a-zA-Z0-9_-]+)(?:\s+[^>]*)?>')
    
    pos = 0
    result = []
    for match in tag_re.finditer(text):
        chunk = text[pos:match.start()]
        chunk = chunk.replace('<', '&lt;').replace('>', '&gt;')
        result.append(chunk)
        
        tag_name = match.group(1).lower()
        full_tag = match.group(0)
        
        if tag_name in allowed_tags or (tag_name == 'blockquote' and 'expandable' in full_tag):
            result.append(full_tag)
        else:
            result.append(full_tag.replace('<', '&lt;').replace('>', '&gt;'))
            
        pos = match.end()
        
    chunk = text[pos:]
    chunk = chunk.replace('<', '&lt;').replace('>', '&gt;')
    result.append(chunk)
    
    return "".join(result)

async def describe_image(base64_image: str, model_name: str = None) -> str:
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/chat/completions"
    headers = {"Content-Type": "application/json"}
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
        
    prompt_text = (
        "Ты — помощник-визионер. Твоя задача — максимально подробно и детально описать присланное изображение на русском языке. "
        "Опиши людей, предметы, их действия, цвета, локацию, распознай и дословно выпиши весь текст на картинке, если он есть. "
        "Пиши структурировано и понятно, так как твой ответ будет использоваться для поиска информации по этой картинке в интернете."
    )
    
    # Запрашиваем список всех доступных в API моделей
    try:
        all_api_models = await get_cached_models()
    except Exception:
        all_api_models = []

    # Находим все мультимодальные модели на сервере
    vision_models_on_server = [m for m in all_api_models if is_model_multimodal(m)]
    
    # Формируем каскадную цепочку мультимодальных моделей
    fallback_chain = []
    
    # Если передан конкретный model_name и он есть на сервере, он идет первым
    if model_name and model_name in all_api_models:
        fallback_chain.append(model_name)
        
    # Добавляем все остальные мультимодальные модели с сервера
    for m in vision_models_on_server:
        if m not in fallback_chain:
            fallback_chain.append(m)
            
    # Если ничего не нашли, а пользователь передал модель, пробуем её
    if not fallback_chain and model_name:
        fallback_chain.append(model_name)
        
    # Если цепочка всё ещё пустая, используем OLLAMA_MODEL как крайний случай
    if not fallback_chain:
        fallback_chain = [OLLAMA_MODEL]
        
    start_model = fallback_chain[0]
    
    logger.info(f"[VISION] Запуск распознавания картинки. Стартовая модель: {start_model}, Цепочка: {fallback_chain}...")
    
    for active_model in fallback_chain:
        payload = {
            "model": active_model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt_text},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }
            ],
            "temperature": 0.2
        }
        
        # Пробуем до 2 раз на каждую модель
        max_attempts_per_model = 2
        for attempt in range(1, max_attempts_per_model + 1):
            if active_model != start_model or attempt > 1:
                await update_status(warning_text=f"Переключаемся на резервный визионер {active_model} (Попытка {attempt}/{max_attempts_per_model})...")
                
            async with httpx.AsyncClient(timeout=20.0) as client:
                try:
                    logger.info(f"[VISION] Запрос к {active_model}. Попытка {attempt}/{max_attempts_per_model}. POST на {url}...")
                    response = await client.post(url, json=payload, headers=headers)
                    logger.info(f"[VISION] Ответ {active_model}: {response.status_code}")
                    response.raise_for_status()
                    result = response.json()
                    content = result["choices"][0]["message"].get("content", "") or ""
                    logger.info(f"[VISION] Успешно получено описание от {active_model}. Длина: {len(content)} симв.")
                    return content.strip()
                except Exception as e:
                    logger.warning(f"[VISION] [СБОЙ] Модель {active_model} на попытке {attempt} вернула ошибку: {repr(e)}")
                    
    # Если все модели сбоят
    raise ValueError("Ни одна из доступных мультимодальных моделей (визионеров) не смогла распознать изображение.")


async def generate_document_summary(filename: str, file_text: str, model_name: str = None) -> str:
    prompt = (
        f"Ты — эксперт по анализу документов. Сделай краткую, но информативную выжимку прикрепленного файла '{filename}'.\n"
        f"Укажи основную тему документа, ключевые разделы, сущности, даты и суть содержимого.\n"
        f"Если это таблица Excel, опиши структуру листов, заголовки ключевых колонок и диапазон данных.\n\n"
        f"Текст документа:\n{file_text[:15000]}\n\n"
        f"Твой ответ должен быть на русском языке, лаконичным, структурированным и содержать не более 1500 символов."
    )
    logger.info(f"[SUMMARY] Генерация выжимки для файла: {filename}...")
    try:
        summary = await ask_ollama(prompt, model_name=model_name)
        logger.info(f"[SUMMARY] Выжимка успешно сгенерирована. Длина: {len(summary)} симв.")
        return summary.strip()
    except Exception as e:
        logger.error(f"[SUMMARY] [ОШИБКА] Не удалось сгенерировать выжимку: {e}")
        return f"Файл {filename}. Содержит текст: {file_text[:500]}..."

async def extract_text_from_document(file_path: str, file_name: str) -> str:
    ext = file_name.lower().split('.')[-1]
    
    if ext == 'pdf':
        if not pypdf:
            raise ImportError("Модуль 'pypdf' не установлен на сервере.")
        reader = pypdf.PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
        
    elif ext == 'docx':
        if not docx:
            raise ImportError("Модуль 'python-docx' не установлен на сервере.")
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
        
    elif ext == 'xlsx':
        if not openpyxl:
            raise ImportError("Модуль 'openpyxl' не установлен на сервере.")
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text_parts.append(f"--- Лист: {sheet} ---")
            for row in ws.iter_rows(values_only=True):
                row_str = " | ".join([str(val) for val in row if val is not None])
                if row_str.strip():
                    text_parts.append(row_str)
        return "\n".join(text_parts)
        
    elif ext == 'xls':
        if not xlrd:
            raise ImportError("Модуль 'xlrd' не установлен на сервере.")
        wb = xlrd.open_workbook(file_path)
        text_parts = []
        for sheet_idx in range(wb.nsheets):
            ws = wb.sheet_by_index(sheet_idx)
            text_parts.append(f"--- Лист: {ws.name} ---")
            for row_idx in range(ws.nrows):
                row = ws.row_values(row_idx)
                row_str = " | ".join([str(val) for val in row if val != ''])
                if row_str.strip():
                    text_parts.append(row_str)
        return "\n".join(text_parts)
        
    else:
        # Default to text decoding
        with open(file_path, 'rb') as f:
            raw_bytes = f.read()
        for encoding in ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']:
            try:
                return raw_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Не удалось определить кодировку текстового файла.")

def make_status_html(query: str, status_text: str, progress_percent: int, search_queries: list = None, attempt: int = 1, warning_text: str = None, thoughts: str = None, model_name: str = None) -> str:
    # Update ContextVars for tracking status
    if query is not None:
        current_query_var.set(query)
    if status_text is not None:
        current_status_text_var.set(status_text)
    if progress_percent is not None:
        current_progress_percent_var.set(progress_percent)
    if search_queries is not None:
        current_search_queries_var.set(search_queries)
    if attempt is not None:
        current_attempt_var.set(attempt)
    if thoughts is not None:
        current_thoughts_var.set(thoughts)

    if not model_name:
        model_name = current_model_var.get()
    display_query = query[:50] + "..." if len(query) > 50 else query
    
    filled_blocks = progress_percent // 10
    empty_blocks = 10 - filled_blocks
    progress_bar = f"📊 <code>[{'■' * filled_blocks}{'□' * empty_blocks}] {progress_percent}%</code>"
    
    attempt_str = f" (Попытка {attempt}/3)" if attempt > 1 else ""
    
    if progress_percent <= 30:
        circle = "🔴"
    elif progress_percent <= 50:
        circle = "🟠"
    elif progress_percent <= 75:
        circle = "🟡"
    elif progress_percent <= 90:
        circle = "🔵"
    else:
        circle = "🟢"
        
    model_str = f"🤖 <b>Модель:</b> <code>{model_name}</code>\n" if model_name else ""
    
    html = (
        f"🔍 <b>Запрос:</b> <code>{display_query}</code>\n"
        f"{progress_bar}\n"
        f"{model_str}"
        f"{circle} <b>Статус:</b> {status_text}{attempt_str}\n"
    )
    
    if thoughts:
        clean_thoughts = clean_markdown_from_thoughts(thoughts)
        if clean_thoughts:
            safe_thoughts = safe_html_cleaner(clean_thoughts)
            html += f"\n<blockquote expandable>{safe_thoughts}</blockquote>\n"
        
    if warning_text:
        html += f"\n⚠️ {warning_text}\n"
        
    if search_queries:
        html += "\n<i>Ищу по запросам:</i>"
        for item in search_queries:
            if isinstance(item, dict):
                q_text = item.get("query", "")
                html += f"\n🔍 <code>{q_text}</code>"
                for site in item.get("sites", []):
                    html += f"\n   └ 📄 <i>{site}</i>"
            else:
                html += f"\n🔍 <code>{item}</code>"
            
    return auto_close_html_tags(html)

def get_clean_fallback_query(query: str) -> str:
    """
    Очищает запрос от префиксов [Фото-запрос] или [Файл-запрос] с содержимым файла,
    возвращая только чистый текст вопроса пользователя.
    """
    if not query:
        return ""
        
    query = query.strip()
    
    # 1. Если это Фото-запрос
    if query.startswith("[Фото-запрос]"):
        prefix = "[Фото-запрос] Вопрос пользователя по фото:"
        if query.startswith(prefix):
            return query[len(prefix):].strip()
        return query.replace("[Фото-запрос]", "").strip()
        
    # 2. Если это Файл-запрос
    if query.startswith("[Файл-запрос]"):
        marker = "Вопрос пользователя по файлу:"
        idx = query.rfind(marker)
        if idx != -1:
            return query[idx + len(marker):].strip()
            
        cleaned = query.replace("[Файл-запрос]", "").strip()
        lines = cleaned.split("\n")
        if lines:
            return lines[0][:80].strip()
            
        return cleaned[:80].strip()
        
    return query

async def update_status(
    status_text: str = None,
    progress_percent: int = None,
    warning_text: str = None,
    search_queries: list = None,
    attempt: int = None,
    thoughts: str = None
):
    updater = current_status_updater_var.get()
    if not updater:
        return
        
    if status_text is not None:
        current_status_text_var.set(status_text)
    if progress_percent is not None:
        current_progress_percent_var.set(progress_percent)
    if search_queries is not None:
        current_search_queries_var.set(search_queries)
    if attempt is not None:
        current_attempt_var.set(attempt)
    if thoughts is not None:
        current_thoughts_var.set(thoughts)
        
    html = make_status_html(
        query=current_query_var.get(),
        status_text=current_status_text_var.get(),
        progress_percent=current_progress_percent_var.get(),
        search_queries=current_search_queries_var.get(),
        attempt=current_attempt_var.get(),
        warning_text=warning_text,
        thoughts=current_thoughts_var.get()
    )
    try:
        await updater(html)
    except Exception as e:
        logger.error(f"Ошибка при обновлении статуса: {e}")

async def fetch_page_content(url: str) -> str:
    lower_url = url.lower().split("?")[0].split("#")[0]
    if lower_url.endswith((".pdf", ".zip", ".tar.gz", ".rar", ".png", ".jpg", ".jpeg", ".gif", ".mp3", ".mp4", ".avi", ".mov", ".docx", ".xlsx", ".pptx", ".exe", ".dmg")):
        logger.warning(f"[FETCH] Пропускаем {url}, так как это бинарный файл по расширению")
        return ""

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
        "Accept-Language": "ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
        "Sec-Ch-Ua": '"Google Chrome";v="125", "Chromium";v="125", "Not.A/Brand";v="24"',
        "Sec-Ch-Ua-Mobile": "?0",
        "Sec-Ch-Ua-Platform": '"Windows"',
        "Sec-Fetch-Dest": "document",
        "Sec-Fetch-Mode": "navigate",
        "Sec-Fetch-Site": "none",
        "Sec-Fetch-User": "?1"
    }
    logger.info(f"[FETCH] Начинаем скачивание страницы: {url}")
    async with httpx.AsyncClient(timeout=10.0, follow_redirects=True) as client:
        try:
            response = await client.get(url, headers=headers)
            logger.info(f"[FETCH] Сервер вернул код {response.status_code} для {url}")
            if response.status_code == 200:
                content_type = response.headers.get("content-type", "").lower()
                if "text/html" not in content_type and "text/plain" not in content_type:
                    logger.warning(f"[FETCH] Пропускаем {url}, так как это не HTML/текстовый контент (Content-Type: {content_type})")
                    return ""
                soup = BeautifulSoup(response.text, "html.parser")
                
                # Вырезаем стандартные мусорные теги
                for element in soup(["script", "style", "header", "footer", "nav", "aside", "iframe", "noscript"]):
                    element.decompose()
                
                # Агрессивно удаляем рекламные блоки, куки, виджеты и меню по селекторам классов/id
                unwanted_selectors = [
                    "div[class*='cookie']", "div[id*='cookie']",
                    "div[class*='banner']", "div[id*='banner']",
                    "div[class*='popup']", "div[id*='popup']",
                    "div[class*='social']", "div[id*='social']",
                    "div[class*='share']", "div[id*='share']",
                    "div[class*='ad-']", "div[id*='ad-']",
                    "div[class*='ads']", "div[id*='ads']",
                    "div[class*='menu']", "div[id*='menu']",
                    "div[class*='sidebar']", "div[id*='sidebar']",
                    "div[class*='modal']", "div[id*='modal']",
                    "div[class*='widget']", "div[id*='widget']",
                    "span[class*='cookie']", "span[id*='cookie']"
                ]
                for selector in unwanted_selectors:
                    try:
                        for element in soup.select(selector):
                            element.decompose()
                    except Exception:
                        pass
                        
                text = soup.get_text(separator=" ")
                
                # Дополнительная очистка на уровне строк
                lines = []
                for line in text.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    # Игнорируем слишком короткие нетекстовые строки
                    if len(line) < 5 and not any(c.isalnum() for c in line):
                        continue
                    lower_line = line.lower()
                    # Игнорируем явный cookie/сервисный мусор
                    if any(phrase in lower_line for phrase in [
                        "cookie", "куки", "согласен на обработку", "политика конфиденциальности", 
                        "все права защищены", "согласиться и продолжить", "нажмите, чтобы",
                        "подписаться", "авторизоваться", "войти в аккаунт", "зарегистрироваться"
                    ]):
                        continue
                    lines.append(line)
                    
                clean_text = "\n".join(lines)
                logger.info(f"[FETCH] Успешно скачано {url}. Символов всего после очистки: {len(clean_text)}")
                return clean_text
            else:
                logger.warning(f"[FETCH] Неверный статус-код {response.status_code} для {url}")
                return ""
        except Exception as e:
            logger.error(f"[FETCH] [ОШИБКА] Исключение при скачивании {url}: {repr(e)}")
            return ""

async def search_searxng_raw(query: str, time_range: str = None) -> list:
    url = f"{SEARXNG_URL.rstrip('/')}/search"
    params = {
        "q": query,
        "format": "json",
        "pageno": 1
    }
    if time_range:
        params["time_range"] = time_range
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    
    logger.info(f"[SEARXNG] Запрос к SearxNG: {url} с q='{query}'")
    for attempt in range(1, 3):
        async with httpx.AsyncClient(timeout=15.0) as client:
            try:
                await asyncio.sleep(1.0)
                logger.info(f"[SEARXNG] Попытка {attempt}/2. Отправка GET...")
                response = await client.get(url, params=params, headers=headers)
                
                if response.status_code == 429:
                    logger.warning(f"[SEARXNG] [HTTP 429] Лимит запросов. Попытка {attempt}/2. Ждем 3.0 секунды...")
                    await update_status(warning_text="Поисковая система перегружена. Повторяем попытку через 3 секунды...")
                    await asyncio.sleep(3.0)
                    continue
                    
                response.raise_for_status()
                data = response.json()
                results = data.get("results", [])
                logger.info(f"[SEARXNG] Успешно получен ответ. Всего результатов: {len(results)}")
                return results
            except Exception as e:
                logger.error(f"[SEARXNG] [ОШИБКА] Исключение при запросе к SearxNG по '{query}' (попытка {attempt}): {repr(e)}")
                if attempt < 2:
                    await update_status(warning_text="Задержка ответа от поисковой системы. Повторяем попытку...")
                    await asyncio.sleep(2.0)
                    
    return []


async def ask_ollama(prompt_or_messages, system_prompt: str = None, model_name: str = None, timeout: float = 15.0) -> str:
    if isinstance(prompt_or_messages, list):
        messages = prompt_or_messages
    else:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt_or_messages}
        ]
        
    chunks = []
    try:
        async for chunk in ask_ollama_stream(messages, model_name=model_name):
            chunks.append(chunk)
    except Exception as e:
        logger.error(f"[OLLAMA] Ошибка при вызове ask_ollama через стрим: {e}")
        raise e
        
    result = "".join(chunks)
    if not result:
        raise ValueError("Получен пустой ответ от ИИ-модели.")
    return result


async def ask_ollama_stream(messages: list, model_name: str = None, use_cascade: bool = True):
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/chat/completions"
    headers = {
        "Content-Type": "application/json"
    }
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
        
    start_model = model_name or OLLAMA_MODEL
    
    if not use_cascade:
        fallback_chain = [start_model]
    else:
        # Запрашиваем список всех доступных в API моделей
        all_api_models = await get_cached_models()
        
        # Строим каскадную цепочку моделей (ограничиваем, чтобы избежать бесконечного перебора)
        fallback_chain = []
        if start_model:
            fallback_chain.append(start_model)
        
        extra_added = 0
        for am in all_api_models:
            if am not in fallback_chain:
                fallback_chain.append(am)
                extra_added += 1
                if extra_added >= 4:
                    break
                
        if not fallback_chain:
            fallback_chain = [start_model or OLLAMA_MODEL]
        
    logger.info(f"[OLLAMA-STREAM] Запуск ask_ollama_stream. Стартовая модель: {start_model}, Цепочка: {fallback_chain}, Сообщений в истории: {len(messages)}")
    
    yielded_any = False
    
    for active_model in fallback_chain:
        if yielded_any:
            break
            
        # Очищаем от картинок, если резервная модель текстовая
        current_messages = messages
        if not is_model_multimodal(active_model):
            current_messages = clean_multimodal_messages(messages)
            
        payload = {
            "model": active_model,
            "messages": current_messages,
            "temperature": 0.3,
            "stream": True
        }
        
        # Для стартовой модели пробуем 2 раза, для резервных — 1 раз
        max_attempts_per_model = 2 if active_model == start_model else 1
        for attempt in range(1, max_attempts_per_model + 1):
            if yielded_any:
                break
                
            if active_model != start_model or attempt > 1:
                await update_status(warning_text=f"Переключаемся на резервный ИИ-стрим {active_model} (Попытка {attempt}/{max_attempts_per_model})...")
                
            async with httpx.AsyncClient(timeout=60.0) as client:
                try:
                    logger.info(f"[OLLAMA-STREAM] Запрос к {active_model}. Попытка {attempt}/{max_attempts_per_model}. Установка соединения...")
                    
                    response = None
                    retry_delay = 2.0
                    stream_ctx = None
                    for r_attempt in range(4):  # 1 основная + 3 ретрая при 429
                        if r_attempt > 0:
                            logger.warning(f"[OLLAMA-STREAM] Повторное подключение к {active_model} после 429 через {retry_delay} сек...")
                            await update_status(warning_text=f"ИИ-стрим перегружен (429). Повтор через {retry_delay} сек...")
                            await asyncio.sleep(retry_delay)
                            retry_delay *= 2.0
                            
                        stream_ctx = client.stream("POST", url, json=payload, headers=headers)
                        response = await stream_ctx.__aenter__()
                        logger.info(f"[OLLAMA-STREAM] Статус соединения {active_model} (повтор {r_attempt}): {response.status_code}")
                        if response.status_code != 429:
                            break
                        await stream_ctx.__aexit__(None, None, None)
                        stream_ctx = None
                        
                    response.raise_for_status()
                    logger.info(f"[OLLAMA-STREAM] Успешно подключено к {active_model}, ожидаем первый токен...")
                    
                    # Обновляем текущую модель
                    current_model_var.set(active_model)
                    
                    in_reasoning = False
                    reasoning_ended = False
                    
                    lines_iter = response.aiter_lines()
                    
                    # Читаем до первой содержательной строки с таймаутом 10с
                    async def read_until_first_data():
                        async for line in lines_iter:
                            if not line.strip():
                                continue
                            if line.startswith("data: "):
                                data_str = line[6:]
                                if data_str.strip() == "[DONE]":
                                    continue
                                try:
                                    data_json = json.loads(data_str)
                                    delta = data_json["choices"][0]["delta"]
                                    if delta.get("content") or delta.get("reasoning") or delta.get("reasoning_content"):
                                        return line
                                except Exception:
                                    pass
                        return None

                    try:
                        first_line = await asyncio.wait_for(read_until_first_data(), timeout=10.0)
                        if not first_line:
                            raise ValueError("Стрим закрылся без токенов")
                    except asyncio.TimeoutError:
                        logger.warning(f"[OLLAMA-STREAM] Тайм-аут 10 сек ожидания первого токена от {active_model}")
                        raise asyncio.TimeoutError("Тайм-аут ожидания первого токена (10 сек)")

                    # Функция для обработки строки и возврата контента чанка
                    def process_line(line_str: str):
                        nonlocal in_reasoning, reasoning_ended, yielded_any
                        if not line_str.strip():
                            return None
                        if line_str.startswith("data: "):
                            data_str = line_str[6:]
                            if data_str.strip() == "[DONE]":
                                return None
                            try:
                                data_json = json.loads(data_str)
                                delta = data_json["choices"][0]["delta"]
                                content = delta.get("content", "") or ""
                                reasoning = delta.get("reasoning", "") or delta.get("reasoning_content", "") or ""
                                
                                chunk_res = ""
                                if reasoning:
                                    yielded_any = True
                                    if not in_reasoning and not reasoning_ended:
                                        in_reasoning = True
                                        chunk_res += "<think>\n" + reasoning
                                    else:
                                        chunk_res += reasoning
                                elif content:
                                    yielded_any = True
                                    if in_reasoning:
                                        in_reasoning = False
                                        reasoning_ended = True
                                        chunk_res += "\n</think>\n" + content
                                    else:
                                        chunk_res += content
                                return chunk_res
                            except Exception as parse_err:
                                logger.error(f"[OLLAMA-STREAM] Ошибка парсинга чанка: {parse_err}")
                        return None

                    # Отправляем первый чанк
                    first_chunk = process_line(first_line)
                    if first_chunk:
                        yield first_chunk

                    # Читаем оставшуюся часть стрима без таймаута на первый токен
                    try:
                        async for line in lines_iter:
                            chunk = process_line(line)
                            if chunk:
                                yield chunk
                                
                        # Закрываем тег рассуждений, если стрим закончился, а мы ещё внутри
                        if in_reasoning:
                            yield "\n</think>\n"
                            
                        # Если мы дошли до конца и отдали хоть один токен, генерация считается успешной
                        if yielded_any:
                            return
                    except Exception as stream_err:
                        logger.warning(f"[OLLAMA-STREAM] Ошибка при чтении стрима после первого токена: {stream_err}")
                        raise stream_err
                    finally:
                        if stream_ctx:
                            await stream_ctx.__aexit__(None, None, None)
                except Exception as e:
                    logger.warning(f"[OLLAMA-STREAM] [СБОЙ] Стрим модели {active_model} на попытке {attempt} вернул ошибку: {repr(e)}")
                    if yielded_any:
                        logger.error(f"[OLLAMA-STREAM] Ошибка возникла посреди генерации, прерываем стрим.")
                        yield f"\n\n[Ошибка стриминга: {repr(e)}]"
                        return
                        
    # Если мы дошли сюда и ничего не отдали
    if not yielded_any:
        raise ValueError("Ни одна из доступных ИИ-моделей не смогла ответить на запрос в режиме стриминга.")


def parse_thoughts_and_answer(raw_text: str) -> tuple[str, str, bool]:
    text = raw_text.strip()
    
    # Регулярные выражения для поиска открывающего и закрывающего тегов (поддерживаем think, thinking, thought, thoughts)
    start_match = re.search(r'<\s*(?:mm:)?(?:think|thinking|thought|thoughts)(?:\s+[^>]*)?>', text, re.IGNORECASE)
    end_match = re.search(r'</\s*(?:mm:)?(?:think|thinking|thought|thoughts)\s*>', text, re.IGNORECASE)
    
    if start_match:
        think_start = start_match.start()
        think_content_start = start_match.end()
        
        if end_match and end_match.start() > think_start:
            think_end = end_match.start()
            answer_start = end_match.end()
            
            thoughts = text[think_content_start:think_end].strip()
            answer = text[answer_start:].strip()
            return thoughts, answer, False
        else:
            thoughts = text[think_content_start:].strip()
            return thoughts, "", True
    else:
        if end_match:
            think_end = end_match.start()
            answer_start = end_match.end()
            thoughts = text[:think_end].strip()
            answer = text[answer_start:].strip()
            return thoughts, answer, False
        else:
            return "", text, False

def format_draft_html(thoughts: str, answer: str, is_still_thinking: bool, show_thoughts: bool = True) -> str:
    html = ""
    clean_thoughts = clean_markdown_from_thoughts(thoughts) if thoughts else ""
    safe_thoughts = safe_html_cleaner(clean_thoughts) if clean_thoughts else ""
    safe_answer = safe_html_cleaner(answer) if answer else ""
    
    if show_thoughts and safe_thoughts:
        html += f"<blockquote expandable>{safe_thoughts}</blockquote>"
    if safe_answer:
        if html:
            html += f"\n\n{safe_answer}"
        else:
            html = safe_answer
    return auto_close_html_tags(html)

async def stream_to_message(status_msg: types.Message, messages: list, model_name: str = None, is_fast: bool = False, start_time: float = None, num_sources: int = None) -> str:
    start_model = model_name or OLLAMA_MODEL
    max_attempts = 3
    accumulated_raw_text = ""
    success = False
    
    # 1. Пробуем стартовую модель до 3 раз
    for attempt in range(1, max_attempts + 1):
        accumulated_raw_text = ""
        last_update_ts = 0.0
        try:
            # use_cascade=False, чтобы опрашивать только эту конкретную модель
            async for chunk in ask_ollama_stream(messages, start_model, use_cascade=False):
                if not chunk:
                    continue
                accumulated_raw_text += chunk
                
                now = time.monotonic()
                if now - last_update_ts >= 0.4:
                    thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
                    draft_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
                    
                    if not draft_html.strip():
                        draft_html = "⏳ <i>Анализирую собранные данные и формулирую ответ...</i>"
                        
                    if status_msg:
                        await update_manager.request_update(status_msg, draft_html, is_final=False)
                    last_update_ts = now
            
            if not accumulated_raw_text.strip():
                raise ValueError("Получен пустой ответ от модели.")
            
            success = True
            break
            
        except Exception as e:
            logger.error(f"[RETRY-MSG] Ошибка стартовой модели {start_model} на попытке {attempt}/{max_attempts}: {e}")
            if attempt < max_attempts:
                if status_msg:
                    await update_manager.request_update(
                        status_msg, 
                        f"⚠️ <i>Сбой генерации ({start_model}). Повторная попытка ({attempt}/{max_attempts})...</i>", 
                        is_final=True
                    )
                await asyncio.sleep(2.0)
                
    # 2. Если стартовая модель не ответила после 3 попыток, запускаем каскад резервных моделей с задержкой в 5 секунд
    if not success:
        logger.warning(f"[CASCADE-MSG] Стартовая модель {start_model} не ответила за {max_attempts} попыток. Запуск каскада резервных моделей...")
        try:
            all_api_models = await get_cached_models()
        except Exception as ex:
            logger.error(f"Не удалось получить список моделей из API: {ex}")
            all_api_models = []
            
        tried_models = {start_model}
        full_remaining_chain = [m for m in all_api_models if m not in tried_models]
        
        for active_model in full_remaining_chain:
            logger.info(f"[CASCADE-MSG] Переключение на резервную модель {active_model} через 5 секунд...")
            if status_msg:
                await update_manager.request_update(
                    status_msg,
                    f"⏳ <i>ИИ-модель {start_model} перегружена. Подключаем резервный ИИ {active_model} через 5 сек...</i>",
                    is_final=True
                )
                
            await asyncio.sleep(5.0)
            
            accumulated_raw_text = ""
            last_update_ts = 0.0
            try:
                async for chunk in ask_ollama_stream(messages, active_model, use_cascade=False):
                    if not chunk:
                        continue
                    accumulated_raw_text += chunk
                    
                    now = time.monotonic()
                    if now - last_update_ts >= 0.4:
                        thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
                        draft_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
                        
                        if not draft_html.strip():
                            draft_html = f"⏳ <i>Генерирую ответ с помощью резервной модели {active_model}...</i>"
                            
                        if status_msg:
                            await update_manager.request_update(status_msg, draft_html, is_final=False)
                        last_update_ts = now
                
                if not accumulated_raw_text.strip():
                    raise ValueError(f"Резервная модель {active_model} вернула пустой ответ.")
                
                success = True
                logger.info(f"[CASCADE-MSG] Успешно получен ответ от резервной модели {active_model}")
                break
            except Exception as e:
                logger.error(f"[CASCADE-MSG] Ошибка резервной модели {active_model}: {e}")
                
    if not success:
        # Все попытки и все резервные модели исчерпаны
        final_html = (
            "❌ <b>Ошибка генерации ответа</b>\n\n"
            "К сожалению, все доступные ИИ-модели перегружены или недоступны. "
            "Пожалуйста, повторите ваш запрос позже."
        )
        if num_sources is not None:
            final_html += f"\n\n🔎 <b>Использовано источников:</b> {num_sources}"
        if status_msg:
            await update_manager.request_update(status_msg, final_html, is_final=True)
        return "Ошибка генерации ответа от ИИ-модели."

    thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
    final_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
    if not final_html.strip():
        if is_fast:
            final_html = safe_html_cleaner(clean_markdown_from_thoughts(thoughts))
        else:
            final_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=True)
            
    if num_sources is not None:
        final_html += f"\n\n🔎 <b>Использовано источников:</b> {num_sources}"
        
    if start_time:
        duration = time.monotonic() - start_time
        minutes = int(duration // 60)
        seconds = duration % 60
        if minutes > 0:
            time_str = f"{minutes} мин. {seconds:.1f} сек."
        else:
            time_str = f"{seconds:.1f} сек."
        if num_sources is not None:
            final_html += f"\n⏱ <i>Время выполнения: {time_str}</i>"
        else:
            final_html += f"\n\n⏱ <i>Время выполнения: {time_str}</i>"
    
    if status_msg:
        await update_manager.request_update(status_msg, final_html, is_final=True)
        
    return clean_model_answer(answer) if answer else clean_model_answer(thoughts)

async def stream_to_inline(inline_message_id: str, messages: list, model_name: str = None, is_fast: bool = False, start_time: float = None, num_sources: int = None) -> str:
    start_model = model_name or OLLAMA_MODEL
    max_attempts = 3
    accumulated_raw_text = ""
    success = False
    
    # 1. Пробуем стартовую модель до 3 раз
    for attempt in range(1, max_attempts + 1):
        accumulated_raw_text = ""
        last_update_ts = 0.0
        try:
            async for chunk in ask_ollama_stream(messages, start_model, use_cascade=False):
                if not chunk:
                    continue
                accumulated_raw_text += chunk
                
                now = time.monotonic()
                if now - last_update_ts >= 0.4:
                    thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
                    draft_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
                    
                    if not draft_html.strip():
                        draft_html = "⏳ <i>Анализирую собранные данные и формулирую ответ...</i>"
                        
                    if inline_message_id:
                        await update_manager.request_update(inline_message_id, draft_html, is_final=False)
                    last_update_ts = now
            
            if not accumulated_raw_text.strip():
                raise ValueError("Получен пустой ответ от модели.")
            
            success = True
            break
            
        except Exception as e:
            logger.error(f"[RETRY-INLINE] Ошибка стартовой модели {start_model} на попытке {attempt}/{max_attempts}: {e}")
            if attempt < max_attempts:
                if inline_message_id:
                    await update_manager.request_update(
                        inline_message_id, 
                        f"⚠️ <i>Сбой генерации ({start_model}). Повторная попытка ({attempt}/{max_attempts})...</i>", 
                        is_final=True
                    )
                await asyncio.sleep(2.0)
                
    # 2. Если стартовая модель не ответила после 3 попыток, запускаем каскад резервных моделей с задержкой в 5 секунд
    if not success:
        logger.warning(f"[CASCADE-INLINE] Стартовая модель {start_model} не ответила за {max_attempts} попыток. Запуск каскада резервных моделей...")
        try:
            all_api_models = await get_cached_models()
        except Exception as ex:
            logger.error(f"Не удалось получить список моделей из API: {ex}")
            all_api_models = []
            
        tried_models = {start_model}
        full_remaining_chain = [m for m in all_api_models if m not in tried_models]
        
        for active_model in full_remaining_chain:
            logger.info(f"[CASCADE-INLINE] Переключение на резервную модель {active_model} через 5 секунд...")
            if inline_message_id:
                await update_manager.request_update(
                    inline_message_id,
                    f"⏳ <i>ИИ-модель {start_model} перегружена. Подключаем резервный ИИ {active_model} через 5 сек...</i>",
                    is_final=True
                )
                
            await asyncio.sleep(5.0)
            
            accumulated_raw_text = ""
            last_update_ts = 0.0
            try:
                async for chunk in ask_ollama_stream(messages, active_model, use_cascade=False):
                    if not chunk:
                        continue
                    accumulated_raw_text += chunk
                    
                    now = time.monotonic()
                    if now - last_update_ts >= 0.4:
                        thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
                        draft_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
                        
                        if not draft_html.strip():
                            draft_html = f"⏳ <i>Генерирую ответ с помощью резервной модели {active_model}...</i>"
                            
                        if inline_message_id:
                            await update_manager.request_update(inline_message_id, draft_html, is_final=False)
                        last_update_ts = now
                
                if not accumulated_raw_text.strip():
                    raise ValueError(f"Резервная модель {active_model} вернула пустой ответ.")
                
                success = True
                logger.info(f"[CASCADE-INLINE] Успешно получен ответ от резервной модели {active_model}")
                break
            except Exception as e:
                logger.error(f"[CASCADE-INLINE] Ошибка резервной модели {active_model}: {e}")
                
    if not success:
        # Все попытки и все резервные модели исчерпаны
        final_html = (
            "❌ <b>Ошибка генерации ответа</b>\n\n"
            "К сожалению, все доступные ИИ-модели перегружены или недоступны. "
            "Пожалуйста, повторите ваш запрос позже."
        )
        if num_sources is not None:
            final_html += f"\n\n🔎 <b>Использовано источников:</b> {num_sources}"
        if inline_message_id:
            await update_manager.request_update(inline_message_id, final_html, is_final=True)
        return "Ошибка генерации ответа от ИИ-модели."

    thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
    final_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
    if not final_html.strip():
        if is_fast:
            final_html = safe_html_cleaner(clean_markdown_from_thoughts(thoughts))
        else:
            final_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=True)
            
    if num_sources is not None:
        final_html += f"\n\n🔎 <b>Использовано источников:</b> {num_sources}"
        
    if start_time:
        duration = time.monotonic() - start_time
        minutes = int(duration // 60)
        seconds = duration % 60
        if minutes > 0:
            time_str = f"{minutes} мин. {seconds:.1f} сек."
        else:
            time_str = f"{seconds:.1f} сек."
        if num_sources is not None:
            final_html += f"\n⏱ <i>Время выполнения: {time_str}</i>"
        else:
            final_html += f"\n\n⏱ <i>Время выполнения: {time_str}</i>"
    
    if inline_message_id:
        await update_manager.request_update(inline_message_id, final_html, is_final=True)
        
    return clean_model_answer(answer) if answer else clean_model_answer(thoughts)

def parse_queries_from_response(response: str, default_query: str) -> list[str]:
    queries = []
    for line in response.strip().split("\n"):
        line = line.strip().strip('"').strip("'").strip("-").strip("*").strip()
        if line and (line[0].isdigit() or line.startswith("•")):
            while line and (line[0].isdigit() or line[0] in [".", ")", "-", " ", "*", "•"]):
                line = line[1:].strip()
        if line:
            queries.append(line)
    if not queries:
        queries = [default_query]
    return queries

def validate_queries(queries: list[str]) -> tuple[bool, str]:
    if not queries:
        return False, "Список поисковых запросов пуст."
        
    for q in queries:
        q_strip = q.strip()
        if not q_strip:
            continue
            
        lower_q = q_strip.lower()
        if "think" in lower_q or "mm:" in lower_q:
            return False, f"Запрос '{q_strip}' содержит служебные слова рассуждений (think, mm:)."
            
        words = q_strip.split()
        if len(words) > 10 or len(q_strip) > 80:
            return False, f"Запрос '{q_strip[:30]}...' слишком длинный (более 10 слов или 80 символов) и похож на обычный текст рассуждений."
            
    return True, ""

async def generate_initial_queries(query: str, chat_history: list, is_fast: bool = False, model_name: str = None, base64_images: list[str] = None, context_summary: str = None) -> tuple[list[str], str]:
    dt = get_current_datetime_str()
    
    if is_fast:
        system_prompt = (
            f"Текущие дата и время: {dt}.\n"
            "Ты — эксперт по поиску информации. Твоя задача — проанализировать тему исследования (запрос пользователя) "
            "и сгенерировать ровно 3 различных поисковых запроса для поисковой системы SearxNG, чтобы собрать ОБЩУЮ информацию по этой теме.\n"
            "КРИТИЧЕСКИ ВАЖНО:\n"
            "1. Модель НЕ является собеседником в чате. С тобой никто не ведет переписку. Последняя реплика пользователя — это исключительно тема/запрос для исследования, а не сообщение, на которое нужно ответить или с которым нужно поспорить.\n"
            "2. Твоя единственная цель — сгенерировать 3 нейтральных, объективных и разносторонних поисковых запроса для сбора фактов, не принимая ничью сторону.\n"
            "3. Выбирай язык поисковых запросов правильно для нахождения наиболее качественной и релевантной информации: если тема/запрос касается технологий, IT, программирования, науки, мировых новостей, англоязычных продуктов, ПО или сервисов, формулируй запросы на английском языке (English), так как на нем доступно гораздо больше актуальных и полных данных. Для специфических локальных тем используй русский язык.\n"
            "4. ПРАВИЛО НЕПРЕДВЗЯТОГО ПОИСКА: При поиске информации о последних или актуальных версиях (например, версий программ, модельного ряда устройств, законов или спортивных событий), НЕ вписывай в запросы конкретные названия версий или моделей, известные тебе из обучения (например, не пиши конкретное название старой модели девайса, конкретную версию операционной системы или определенный год), так как это сместит выдачу к устаревшим результатам. Используй общие категориальные понятия (например, 'latest device model', 'newest software version') в сочетании с текущим годом (2026) и словами вроде 'latest', 'new', 'release', 'announcement'.\n"
            "5. Твой ответ должен состоять ТОЛЬКО из поисковых запросов. Запрещено использовать тег <think> и писать ход мыслей. Запрещено писать вводные фразы или любые другие пояснения. Просто выведи 3 поисковых запроса, каждый на новой строке, без кавычек и нумерации."
        )
    else:
        system_prompt = (
            f"Текущие дата и время: {dt}.\n"
            "Ты — эксперт по поиску информации. Твоя задача — проанализировать тему исследования (запрос пользователя) "
            "и сгенерировать ровно 3 различных поисковых запроса для поисковой системы SearxNG, чтобы собрать ОБЩУЮ информацию по этой теме.\n"
            "КРИТИЧЕСКИ ВАЖНО:\n"
            "1. Модель НЕ является собеседником в чате. С тобой никто не ведет переписку. Последняя реплика пользователя — это исключительно тема/запрос для исследования, а не сообщение, на которое нужно ответить или с которым нужно поспорить.\n"
            "2. Твоя единственная цель — сгенерировать 3 нейтральных, объективных и разносторонних поисковых запроса для сбора фактов, не принимая ничью сторону.\n"
            "3. Выбирай язык поисковых запросов правильно для нахождения наиболее качественной и релевантной информации: если тема/запрос касается технологий, IT, программирования, науки, мировых новостей, англоязычных продуктов, ПО или сервисов, формулируй запросы на английском языке (English), так как на нем доступно гораздо больше актуальных и полных данных. Для специфических локальных тем используй русский язык.\n"
            "4. ПРАВИЛО НЕПРЕДВЗЯТОГО ПОИСКА: При поиске информации о последних или актуальных версиях (например, версий программ, модельного ряда устройств, законов или спортивных событий), НЕ вписывай в запросы конкретные названия версий или моделей, известные тебе из обучения (например, не пиши конкретное название старой модели девайса, конкретную версию операционной системы или определенный год), так как это сместит выдачу к устаревшим результатам. Используй общие категориальные понятия (например, 'latest device model', 'newest software version') в сочетании с текущим годом (2026) и словами вроде 'latest', 'new', 'release', 'announcement'.\n"
            "5. Запрещено использовать в рассуждениях обращения к пользователю ('ты', 'вы') и фразы, имитирующие живой диалог.\n"
            "6. Твой ответ должен СТРОГО начинаться со своего хода мыслей, завернутого в тег <think>, например:\n"
            "<think>\nЗдесь твои подробные рассуждения о том, какие ключевые понятия темы нужно исследовать и почему нужны именно эти запросы. Пиши только чистый простой текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков).\n</think>\n"
            "После закрывающего тега </think> сразу напиши поисковые запросы, каждый на новой строке, без кавычек и нумерации."
        )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    if is_fast:
        prompt_text = (
            f"Тема исследования (запрос пользователя): {query}\n\n"
            "Сгенерируй поисковые запросы для сбора фактов по этой теме.\n"
            "КРИТИЧЕСКИ ВАЖНО:\n"
            "1. Не используй конкретные названия версий/моделей, известные тебе из обучения. Используй общие понятия (например, 'latest device version' вместо устаревшего названия) и год (2026).\n"
            "2. Не используй тег <think> и не пиши свои рассуждения.\n"
            "3. Выведи только поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек, без нумерации, без обращения к пользователю и без споров с утверждениями."
        )
    else:
        prompt_text = (
            f"Тема исследования (запрос пользователя): {query}\n\n"
            "Сгенерируй поисковые запросы для сбора фактов по этой теме.\n"
            "КРИТИЧЕСКИ ВАЖНО:\n"
            "1. Не используй конкретные названия версий/моделей, известные тебе из обучения. Используй общие понятия (например, 'latest device version' вместо устаревшего названия) и год (2026).\n"
            "2. Начни свой ответ строго со своего хода мыслей на русском языке, завернутого в тег <think>...</think>.\n"
            "3. Внутри тегов <think> проводи сухой анализ ключевых понятий темы, без обращения к пользователю и споров с ним.\n"
            "4. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и без нумерации."
        )
        
    if base64_images and is_model_multimodal(model_name):
        content_structure = [{"type": "text", "text": prompt_text}]
        for b64 in base64_images:
            content_structure.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})
        messages.append({"role": "user", "content": content_structure})
    else:
        if base64_images and context_summary:
            prompt_text = f"Контекст присланных пользователем фото (Описание от визионера):\n{context_summary}\n\n" + prompt_text
        messages.append({"role": "user", "content": prompt_text})
    
    saved_thoughts = ""
    all_api_models = await get_cached_models()
    
    for attempt in range(1, 4):
        active_model = get_model_for_attempt(model_name, attempt, all_api_models)
        try:
            if active_model != (model_name or OLLAMA_MODEL) or attempt > 1:
                await update_status(warning_text=f"Переключаемся на ИИ {active_model} для подбора запросов (Попытка {attempt}/3)...")
                
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (initial_queries, модель {active_model}, попытка {attempt}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not is_fast and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                error_msg = "В твоем ответе СТРОГО отсутствует обязательный ход мыслей внутри тегов <think>...</think>. Ты обязан сначала написать свои рассуждения в теге <think>, а затем запросы."
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
                
            logger.warning(f"Попытка {attempt} ({active_model}) - Ошибка валидации запросов: {error_msg}")
            messages.append({"role": "assistant", "content": response})
            if is_fast:
                user_err_msg = (
                    f"Ошибка валидации запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ровно 3 поисковых запроса (каждый на новой строке, без кавычек, без нумерации и без какого-либо хода мыслей)."
                )
            else:
                user_err_msg = (
                    f"Ошибка форматирования запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ответ строго по формату:\n"
                    "1. Ход мыслей СТРОГО внутри тегов <think>...</think>.\n"
                    "2. После закрывающего тега </think> СРАЗУ выведи только поисковые запросы (каждый на новой строке, без кавычек, без нумерации и без лишних рассуждений)."
                )
            messages.append({
                "role": "user",
                "content": user_err_msg
            })
        except Exception as e:
            logger.error(f"Ошибка на попытке {attempt} ({active_model}) при генерации первичных запросов: {repr(e)}")
            
    # Если 3 основные попытки провалились, перебираем все остальные доступные модели из API
    tried_models = {
        get_model_for_attempt(model_name, 1, all_api_models),
        get_model_for_attempt(model_name, 2, all_api_models),
        get_model_for_attempt(model_name, 3, all_api_models)
    }
    
    full_remaining_chain = [m for m in all_api_models if m not in tried_models]
    
    for active_model in full_remaining_chain:
        await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} для подбора запросов...")
        try:
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (каскад {active_model}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not is_fast and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
        except Exception as e:
            logger.warning(f"Каскадный опрос модели {active_model} завершился ошибкой: {repr(e)}")
            
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла сгенерировать валидные поисковые запросы.")

async def generate_refined_queries(query: str, chat_history: list, research_state: str, model_name: str = None, context_summary: str = None) -> tuple[list[str], str]:
    dt = get_current_datetime_str()
    clean_query = get_clean_fallback_query(query)
    system_prompt = (
        f"Текущие дата и время: {dt}.\n"
        "Ты — эксперт по глубокому поиску. Тебе предоставлен текущий лог исследования (Research State) с первыми результатами.\n"
        "Твоя задача — проанализировать лог исследования и тему пользователя, найти пробелы и подобрать 3 глубоких поисковых запроса.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Модель НЕ является собеседником в чате. С тобой никто не ведет переписку. Последняя реплика пользователя — это исключительно тема/запрос для исследования, а не сообщение, на которое нужно ответить или с которым нужно поспорить.\n"
        "2. Твоя единственная цель — сгенерировать 3 нейтральных, объективных и разносторонних поисковых запроса для сбора фактов, не принимая ничью сторону.\n"
        "3. Выбирай язык поисковых запросов правильно для нахождения наиболее качественной и релевантной информации: если тема/запрос касается технологий, IT, программирования, науки, мировых новостей, англоязычных продуктов, ПО или сервисов, формулируй запросы на английском языке (English), так как на нем доступно гораздо больше актуальных и полных данных. Для специфических локальных тем используй русский язык.\n"
        "4. ПРАВИЛО НЕПРЕДВЗЯТОГО ПОИСКА: При поиске информации о последних или актуальных версиях (например, версий программ, модельного ряда устройств, законов или спортивных событий), НЕ вписывай в запросы конкретные названия версий или моделей, известные тебе из обучения (например, не пиши конкретное название старой модели девайса, конкретную версию операционной системы или определенный год), так как это сместит выдачу к устаревшим результатам. Используй общие категориальные понятия (например, 'latest device model', 'newest software version') в сочетании с текущим годом (2026) и словами вроде 'latest', 'new', 'release', 'announcement'.\n"
        "5. Твой ответ должен СТРОГО начинаться со своего хода мыслей, завернутого в тег <think>...</think>.\n"
        "6. Внутри тегов <think> пиши исключительно чистый текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков) и без обращения к пользователю.\n"
        "7. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и нумерации.\n"
        "8. Каждый поисковый запрос должен быть предельно кратким (не более 6-8 слов) и состоять только из ключевых фраз по теме поиска, а не из длинных предложений.\n"
        "9. Любые вводные или завершающие фразы, вежливые обращения, приветствия и мета-размышления после тега </think> СТРОГО ЗАПРЕЩЕНЫ."
    )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
    
    prompt = (
        f"Тема исследования (запрос пользователя): {query}\n\n"
        f"Текущий лог исследования:\n{research_state}\n\n"
        "Сгенерируй 3 глубоких уточняющих запроса.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Не используй конкретные названия версий/моделей, известные тебе из обучения. Используй общие понятия (например, 'latest device version' вместо устаревшего названия) и год (2026).\n"
        "2. Начни свой ответ строго со своего хода мыслей на русском языке, завернутого в тег <think>...</think>.\n"
        "3. Внутри тегов <think> проводи сухой анализ ключевых понятий темы, без обращения к пользователю и споров с ним.\n"
        "4. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и без нумерации."
    )
    if context_summary:
        prompt = f"Контекст присланных пользователем файлов/фото (Краткая выжимка):\n{context_summary}\n\n" + prompt
        
    messages.append({"role": "user", "content": prompt})
    
    saved_thoughts = ""
    all_api_models = await get_cached_models()
    
    for attempt in range(1, 4):
        active_model = get_model_for_attempt(model_name, attempt, all_api_models)
        try:
            if active_model != (model_name or OLLAMA_MODEL) or attempt > 1:
                await update_status(warning_text=f"Переключаемся на ИИ {active_model} для подбора запросов (Попытка {attempt}/3)...")
                
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (refined_queries, модель {active_model}, попытка {attempt}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                error_msg = "В твоем ответе СТРОГО отсутствует обязательный ход мыслей внутри тегов <think>...</think>. Ты обязан сначала написать свои рассуждения в теге <think>, а затем запросы."
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
                
            logger.warning(f"Попытка {attempt} ({active_model}) - Ошибка валидации запросов: {error_msg}")
            await update_status(warning_text=f"Неверный формат запросов (Попытка {attempt}/3). Корректируем и генерируем заново...")
            messages.append({"role": "assistant", "content": response})
            messages.append({
                "role": "user",
                "content": (
                    f"Ошибка форматирования запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ответ строго по формату:\n"
                    "1. Ход мыслей СТРОГО внутри тегов <think>...</think>.\n"
                    "2. После закрывающего тега </think> СРАЗУ выведи только поисковые запросы (каждый на новой строке, без кавычек, без нумерации и без лишних рассуждений)."
                )
            })
        except Exception as e:
            logger.error(f"Ошибка на попытке {attempt} ({active_model}) при генерации уточняющих запросов: {repr(e)}")
            
    # Если 3 основные попытки провалились, перебираем все остальные доступные модели из API
    tried_models = {
        get_model_for_attempt(model_name, 1, all_api_models),
        get_model_for_attempt(model_name, 2, all_api_models),
        get_model_for_attempt(model_name, 3, all_api_models)
    }
    
    full_remaining_chain = [m for m in all_api_models if m not in tried_models]
    
    for active_model in full_remaining_chain:
        await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} для подбора запросов...")
        try:
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (каскад {active_model}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
        except Exception as e:
            logger.warning(f"Каскадный опрос модели {active_model} завершился ошибкой: {repr(e)}")
            
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла сгенерировать поисковые запросы.")

async def generate_opinion_queries(query: str, chat_history: list, research_state: str, model_name: str = None, context_summary: str = None) -> tuple[list[str], str]:
    dt = get_current_datetime_str()
    clean_query = get_clean_fallback_query(query)
    system_prompt = (
        f"Текущие дата и время: {dt}.\n"
        "Ты — эксперт по анализу общественного мнения и отзывов. Твоя задача — подобрать 3 запроса для поиска отзывов, опыта людей и мнений по теме исследования.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Модель НЕ является собеседником в чате. С тобой никто не ведет переписку. Последняя реплика пользователя — это исключительно тема/запрос для исследования, а не сообщение, на которое нужно ответить или с которым нужно поспорить.\n"
        "2. Твоя единственная цель — сгенерировать 3 нейтральных, объективных и разносторонних поисковых запроса для сбора отзывов и мнений, не принимая ничью сторону.\n"
        "3. Выбирай язык поисковых запросов правильно для нахождения наиболее качественной и релевантной информации: если тема/запрос касается технологий, IT, программирования, науки, мировых новостей, англоязычных продуктов, ПО или сервисов, формулируй запросы на английском языке (English), так как на нем доступно гораздо больше актуальных и полных данных. Для специфических локальных тем используй русский язык.\n"
        "4. ПРАВИЛО НЕПРЕДВЗЯТОГО ПОИСКА: При поиске информации о последних или актуальных версиях (например, версий программ, модельного ряда устройств, законов или спортивных событий), НЕ вписывай в запросы конкретные названия версий или моделей, известные тебе из обучения (например, не пиши конкретное название старой модели девайса, конкретную версию операционной системы или определенный год), так как это сместит выдачу к устаревшим результатам. Используй общие категориальные понятия (например, 'latest device model', 'newest software version') в сочетании с текущим годом (2026) и словами вроде 'latest', 'new', 'release', 'announcement'.\n"
        "5. Твой ответ должен СТРОГО начинаться со своего хода мыслей, завернутого в тег <think>...</think>.\n"
        "6. Внутри тегов <think> пиши исключительно чистый текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков) и без обращения к пользователю.\n"
        "7. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и нумерации.\n"
        "8. Каждый поисковый запрос должен быть предельно кратким (не более 6-8 слов) и состоять только из ключевых фраз по теме поиска, а не из длинных предложений.\n"
        "9. Любые вводные или завершающие фразы, вежливые обращения, приветствия и мета-размышления после тега </think> СТРОГО ЗАПРЕЩЕНЫ."
    )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    prompt = (
        f"Тема исследования (запрос пользователя): {query}\n\n"
        f"Текущий лог исследования:\n{research_state}\n\n"
        "Сгенерируй 3 запроса для сбора мнений и отзывов.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Не используй конкретные названия версий/моделей, известные тебе из обучения. Используй общие понятия (например, 'latest device version' вместо устаревшего названия) и год (2026).\n"
        "2. Начни свой ответ строго со своего хода мыслей на русском языке, завернутого в тег <think>...</think>.\n"
        "3. Внутри тегов <think> проводи сухой анализ ключевых понятий темы, без обращения к пользователю и споров с ним.\n"
        "4. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и без нумерации."
    )
    if context_summary:
        prompt = f"Контекст присланных пользователем файлов/фото (Краткая выжимка):\n{context_summary}\n\n" + prompt
        
    messages.append({"role": "user", "content": prompt})
    
    saved_thoughts = ""
    all_api_models = await get_cached_models()
    
    for attempt in range(1, 4):
        active_model = get_model_for_attempt(model_name, attempt, all_api_models)
        try:
            if active_model != (model_name or OLLAMA_MODEL) or attempt > 1:
                await update_status(warning_text=f"Переключаемся на ИИ {active_model} для подбора запросов (Попытка {attempt}/3)...")
                
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (opinion_queries, модель {active_model}, попытка {attempt}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                error_msg = "В твоем ответе СТРОГО отсутствует обязательный ход мыслей внутри тегов <think>...</think>. Ты обязан сначала написать свои рассуждения в теге <think>, а затем запросы."
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
                
            logger.warning(f"Попытка {attempt} ({active_model}) - Ошибка валидации запросов: {error_msg}")
            await update_status(warning_text=f"Неверный формат запросов (Попытка {attempt}/3). Корректируем и генерируем заново...")
            messages.append({"role": "assistant", "content": response})
            messages.append({
                "role": "user",
                "content": (
                    f"Ошибка форматирования запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ответ строго по формату:\n"
                    "1. Ход мыслей СТРОГО внутри тегов <think>...</think>.\n"
                    "2. После закрывающего тега </think> СРАЗУ выведи только поисковые запросы (каждый на новой строке, без кавычек, без нумерации и без лишних рассуждений)."
                )
            })
        except Exception as e:
            logger.error(f"Ошибка на попытке {attempt} ({active_model}) при генерации запросов мнений: {repr(e)}")
            
    # Если 3 основные попытки провалились, перебираем все остальные доступные модели из API
    tried_models = {
        get_model_for_attempt(model_name, 1, all_api_models),
        get_model_for_attempt(model_name, 2, all_api_models),
        get_model_for_attempt(model_name, 3, all_api_models)
    }
    
    full_remaining_chain = [m for m in all_api_models if m not in tried_models]
    
    for active_model in full_remaining_chain:
        await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} для подбора запросов...")
        try:
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (каскад {active_model}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
        except Exception as e:
            logger.warning(f"Каскадный опрос модели {active_model} завершился ошибкой: {repr(e)}")
            
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла сгенерировать поисковые запросы.")

async def generate_cross_queries(query: str, chat_history: list, research_state: str, model_name: str = None, context_summary: str = None) -> tuple[list[str], str]:
    dt = get_current_datetime_str()
    clean_query = get_clean_fallback_query(query)
    system_prompt = (
        f"Текущие дата и время: {dt}.\n"
        "Ты — эксперт по анализу данных. Твоя задача — сопоставить уже найденную общую информацию и собранные мнения/отзывы по теме исследования, "
        "выявить противоречия, сомнительные утверждения или пробелы, которые требуют дополнительной перепроверки в поисковике.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Модель НЕ является собеседником в чате. С тобой никто не ведет переписку. Последняя реплика пользователя — это исключительно тема/запрос для исследования, а не сообщение, на которое нужно ответить или с которым нужно поспорить.\n"
        "2. Твоя единственная цель — сгенерировать 3 нейтральных, объективных и разносторонних поисковых запроса для проверки противоречий, не принимая ничью сторону.\n"
        "3. Выбирай язык поисковых запросов правильно для нахождения наиболее качественной и релевантной информации: если тема/запрос касается технологий, IT, программирования, науки, мировых новостей, англоязычных продуктов, ПО или сервисов, формулируй запросы на английском языке (English), так как на нем доступно гораздо больше актуальных и полных данных. Для специфических локальных тем используй русский язык.\n"
        "4. ПРАВИЛО НЕПРЕДВЗЯТОГО ПОИСКА: При поиске информации о последних или актуальных версиях (например, версий программ, модельного ряда устройств, законов или спортивных событий), НЕ вписывай в запросы конкретные названия версий или моделей, известные тебе из обучения (например, не пиши конкретное название старой модели девайса, конкретную версию операционной системы или определенный год), так как это сместит выдачу к устаревшим результатам. Используй общие категориальные понятия (например, 'latest device model', 'newest software version') в сочетании с текущим годом (2026) и словами вроде 'latest', 'new', 'release', 'announcement'.\n"
        "5. Твой ответ должен СТРОГО начинаться со своего хода мыслей, завернутого в тег <think>...</think>.\n"
        "6. Внутри тегов <think> пиши исключительно чистый текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков) и без обращения к пользователю.\n"
        "7. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и нумерации.\n"
        "8. Каждый поисковый запрос должен быть предельно кратким (не более 6-8 слов) и состоять только из ключевых фраз по теме поиска, а не из длинных предложений.\n"
        "9. Любые вводные или завершающие фразы, вежливые обращения, приветствия и мета-размышления после тега </think> СТРОГО ЗАПРЕЩЕНЫ."
    )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    prompt = (
        f"Тема исследования (запрос пользователя): {query}\n\n"
        f"Текущие результаты исследования:\n{research_state}\n\n"
        "Сгенерируй 3 поисковых запроса для проверки противоречий и сомнительных фактов.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Не используй конкретные названия версий/моделей, известные тебе из обучения. Используй общие понятия (например, 'latest device version' вместо устаревшего названия) и год (2026).\n"
        "2. Начни свой ответ строго со своего хода мыслей на русском языке, завернутого в тег <think>...</think>.\n"
        "3. Внутри тегов <think> проводи сухой анализ ключевых понятий темы, без обращения к пользователю и споров с ним.\n"
        "4. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и без нумерации."
    )
    if context_summary:
        prompt = f"Контекст присланных пользователем файлов/фото (Краткая выжимка):\n{context_summary}\n\n" + prompt
        
    messages.append({"role": "user", "content": prompt})
    
    saved_thoughts = ""
    all_api_models = await get_cached_models()
    
    for attempt in range(1, 4):
        active_model = get_model_for_attempt(model_name, attempt, all_api_models)
        try:
            if active_model != (model_name or OLLAMA_MODEL) or attempt > 1:
                await update_status(warning_text=f"Переключаемся на ИИ {active_model} для подбора запросов (Попытка {attempt}/3)...")
                
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (cross_queries, модель {active_model}, попытка {attempt}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                error_msg = "В твоем ответе СТРОГО отсутствует обязательный ход мыслей внутри тегов <think>...</think>. Ты обязан сначала написать свои рассуждения в теге <think>, а затем запросы."
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
                
            logger.warning(f"Попытка {attempt} ({active_model}) - Ошибка валидации запросов: {error_msg}")
            await update_status(warning_text=f"Неверный формат запросов (Попытка {attempt}/3). Корректируем и генерируем заново...")
            messages.append({"role": "assistant", "content": response})
            messages.append({
                "role": "user",
                "content": (
                    f"Ошибка форматирования запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ответ строго по формату:\n"
                    "1. Ход мыслей СТРОГО внутри тегов <think>...</think>.\n"
                    "2. После закрывающего тега </think> СРАЗУ выведи только поисковые запросы (каждый на новой строке, без кавычек, без нумерации и без лишних рассуждений)."
                )
            })
        except Exception as e:
            logger.error(f"Ошибка на попытке {attempt} ({active_model}) при генерации перекрестных запросов: {repr(e)}")
            
    # Если 3 основные попытки провалились, перебираем все остальные доступные модели из API
    tried_models = {
        get_model_for_attempt(model_name, 1, all_api_models),
        get_model_for_attempt(model_name, 2, all_api_models),
        get_model_for_attempt(model_name, 3, all_api_models)
    }
    
    full_remaining_chain = [m for m in all_api_models if m not in tried_models]
    
    for active_model in full_remaining_chain:
        await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} для подбора запросов...")
        try:
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (каскад {active_model}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
        except Exception as e:
            logger.warning(f"Каскадный опрос модели {active_model} завершился ошибкой: {repr(e)}")
            
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла сгенерировать поисковые запросы.")


async def run_cross_verification(query: str, general_context: str, opinions_context: str, model_name: str = None) -> str:
    dt = get_current_datetime_str()
    system_prompt = (
        f"Текущие дата и время: {dt}.\n"
        "Ты — ведущий аналитик данных. Проведи перекрестный анализ собранной информации для ответа на вопрос пользователя.\n"
        "У тебя есть два блока данных:\n"
        "1. Общая справочная информация (результаты этапов 1.1 и 1.2).\n"
        "2. Мнения людей, экспертов, отзывы, клинические исследования или личный опыт использования.\n\n"
        "Проанализируй эти данные и составь структурированный отчет:\n"
        "- Что из найденного совпадает и подтверждается обоими блоками?\n"
        "- Какие есть противоречия (разница между теорией/официальными данными и реальными отзывами)?\n"
        "- Что является доказанным фактом, а что — субъективным мнением или мифом?\n"
        "Пиши строго без использования Markdown-разметки (без звездочек, решеток, жирного шрифта, списков и т.д.), только чистый структурированный текст. На русском языке."
    )
    
    prompt = (
        f"Вопрос пользователя: {query}\n\n"
        f"БЛОК 1 (Общая информация):\n{general_context}\n\n"
        f"БЛОК 2 (Мнения и отзывы):\n{opinions_context}\n\n"
        "Составь аналитический отчет перекрестной проверки."
    )
    
    try:
        return await ask_ollama(prompt, system_prompt, model_name=model_name)
    except Exception as e:
        logger.error(f"Ошибка при перекрестном анализе: {e}")
        return "Не удалось провести перекрестный анализ данных."


async def generate_alternative_query(
    original_query: str,
    reason: str,
    failed_urls: str,
    user_query: str,
    model_name: str = None
) -> str:
    system_prompt = (
        "Ты — эксперт по поисковой оптимизации и обходу блокировок при веб-парсинге.\n"
        "Твоя задача — переформулировать поисковый запрос пользователя, так как предыдущий запрос не привел к успеху.\n"
        "ТЕБЕ ПРЕДОСТАВЛЕНЫ:\n"
        "1. Предыдущий неудачный запрос.\n"
        "2. Причина неудачи (пустая выдача или блокировки защит сайтов).\n"
        "3. Список URL, которые не удалось скачать (их нужно избегать или искать альтернативные источники).\n"
        "4. Общая тема исследования пользователя.\n\n"
        "ПРАВИЛА ГЕНЕРАЦИИ:\n"
        "- Сформулируй новый, альтернативный поисковый запрос (не более 6-8 слов), состоящий только из ключевых слов.\n"
        "- Запрос должен быть нацелен на поиск той же информации, но с использованием синонимов, других ключевых слов или смежных источников (например, если официальный сайт заблокирован Cloudflare, ищи статьи, обзоры, зеркала или темы на форумах/Reddit/GitHub).\n"
        "- Избегай повторения слов из неудачного запроса, если они привели к заблокированным сайтам.\n"
        "- Выбирай язык запроса правильно: используй английский язык (English) для IT, технологий, программирования и науки, и русский для локальных тем.\n"
        "- ПРАВИЛО НЕПРЕДВЗЯТОГО ПОИСКА: При поиске информации о последних или актуальных версиях (например, версий программ, модельного ряда устройств, законов или спортивных событий), НЕ вписывай в запросы конкретные названия версий или моделей, известные тебе из обучения (например, не пиши конкретное название старой модели девайса, конкретную версию операционной системы или определенный год), так как это сместит выдачу к устаревшим результатам. Используй общие категориальные понятия (например, 'latest device model', 'newest software version') в сочетании с текущим годом (2026) и словами вроде 'latest', 'new', 'release', 'announcement'.\n"
        "- Ответь СТРОГО одним новым поисковым запросом. Запрещено использовать тег <think>, писать ход мыслей, вводные слова, кавычки и пояснения."
    )
    
    prompt = (
        f"Общая тема исследования: {user_query}\n"
        f"Предыдущий неудачный запрос: {original_query}\n"
        f"Причина неудачи: {reason}\n"
        f"Список сбоивших URL (избегай их или ищи альтернативные домены/платформы):\n{failed_urls or 'Нет'}\n\n"
        f"Сформулируй новый эффективный поисковый запрос:"
    )
    
    try:
        new_query = await ask_ollama(prompt, system_prompt=system_prompt, model_name=model_name, timeout=15.0)
        if "<think>" in new_query and "</think>" in new_query:
            new_query = new_query.split("</think>")[-1].strip()
        cleaned_query = new_query.strip().replace('"', '').replace("'", "")
        if cleaned_query and len(cleaned_query) < 100:
            return cleaned_query
    except Exception as e:
        logger.warning(f"[REFINE-QUERY] Ошибка при генерации альтернативного запроса: {e}")
        
    return original_query + " альтернативные источники"


async def ai_choose_best_sites(
    user_query: str,
    search_query: str,
    candidates: list[dict],
    downloaded_urls: list[str],
    failed_urls: list[str],
    needed_count: int,
    model_name: str = None
) -> list[int]:
    if not candidates:
        return []
        
    candidates_text = []
    for idx, cand in enumerate(candidates):
        status = ""
        if cand["url"] in downloaded_urls:
            status = " [УЖЕ СКАЧАН]"
        elif cand["url"] in failed_urls:
            status = " [ОШИБКА СКАЧИВАНИЯ]"
        candidates_text.append(f"Индекс: {idx}\nЗаголовок: {cand['title']}\nДомен: {cand['url'].split('/')[2] if '/' in cand['url'] else cand['url']}\nСниппет: {cand['snippet']}{status}\n---")
        
    candidates_block = "\n".join(candidates_text)
    
    system_prompt = (
        "Ты — ИИ-навигатор поисковых систем. Твоя задача — проанализировать результаты поиска и выбрать наиболее полезные и авторитетные ссылки для скачивания.\n"
        "ТЕБЕ ПРЕДОСТАВЛЕНЫ:\n"
        "1. Общая тема исследования пользователя и текущий поисковый запрос.\n"
        "2. Список кандидатов (результатов поиска) с их индексами, заголовками, доменами и краткими описаниями (сниппетами).\n"
        "3. Количество ссылок, которые нужно выбрать (needed_count).\n"
        "4. Маркеры [УЖЕ СКАЧАН] и [ОШИБКА СКАЧИВАНИЯ] для ссылок, которые выбирать нельзя.\n\n"
        "КРИТИЧЕСКИ ВАЖНЫЕ ПРАВИЛА:\n"
        f"- Выбери ровно до {needed_count} наиболее авторитетных, подходящих и информативных кандидатов (их индексов) из предоставленного списка, у которых НЕТ маркеров [УЖЕ СКАЧАН] или [ОШИБКА СКАЧИВАНИЯ].\n"
        "- Ответь строго в формате JSON-массива индексов, например: `[0, 2, 5]`.\n"
        "- Если в предоставленном списке кандидатов нет подходящих/релевантных сайтов, или если все оставшиеся сайты не подходят / заблокированы, верни пустой список `[]`.\n"
        "- Твой ответ должен быть СТРОГО валидным JSON-массивом (начинаться с [ и заканчиваться на ]). Никакого лишнего текста, рассуждений, тегов <think> и Markdown-оформления. Только JSON-массив индексов."
    )
    
    prompt = (
        f"Тема исследования: {user_query}\n"
        f"Текущий поисковый запрос: {search_query}\n"
        f"Нужно выбрать ссылок: {needed_count}\n\n"
        f"Результаты поиска:\n{candidates_block}\n\n"
        "Выбери лучшие индексы для скачивания:"
    )
    
    try:
        response = await ask_ollama(prompt, system_prompt=system_prompt, model_name=model_name, timeout=15.0)
        if "<think>" in response and "</think>" in response:
            response = response.split("</think>")[-1].strip()
        
        m = re.search(r"\[\s*\d*\s*(?:,\s*\d*\s*)*\]", response)
        if m:
            chosen_indexes = json.loads(m.group(0))
            if isinstance(chosen_indexes, list):
                valid_indexes = [int(i) for i in chosen_indexes if isinstance(i, (int, float)) and 0 <= int(i) < len(candidates)]
                return valid_indexes[:needed_count]
    except Exception as e:
        logger.warning(f"[AI-NAVIGATOR] Ошибка при выборе индексов: {e}. Ответ модели: {response if 'response' in locals() else 'нет'}")
        
    return []
async def quick_check_url(url: str) -> bool:
    lower_url = url.lower().split("?")[0].split("#")[0]
    if lower_url.endswith((".pdf", ".zip", ".tar.gz", ".rar", ".png", ".jpg", ".jpeg", ".gif", ".mp3", ".mp4", ".avi", ".mov", ".docx", ".xlsx", ".pptx", ".exe", ".dmg")):
        return False
        
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3",
    }
    try:
        async with httpx.AsyncClient(timeout=3.0, follow_redirects=True) as client:
            async with client.stream("GET", url, headers=headers) as response:
                if response.status_code != 200:
                    return False
                content_type = response.headers.get("content-type", "").lower()
                if "text/html" not in content_type and "text/plain" not in content_type:
                    return False
                return True
    except Exception:
        return False


async def process_single_query_parallel(
    query_str: str,
    fetched_urls_shared: set,
    limit: int,
    query_state_item: dict,
    mode: str = "deep",
    model_name: str = None,
    user_query: str = "",
    status_updater = None,
    search_queries_state: list = None,
    iteration: int = 1,
    current_thoughts: str = None,
    progress_percent: int = 30,
    time_range: str = None
) -> list[dict]:
    logger.info(f"[PARALLEL] Старт обработки запроса: '{query_str}'. Лимит скачивания: от 3 до 5 сайтов, time_range: {time_range}")
    
    target_count = 5
    current_query = query_str
    attempts_left = 3
    
    success_sites = []
    failed_urls = []
    downloaded_urls = []
    
    while attempts_left > 0:
        raw_results = await search_searxng_raw(current_query, time_range=time_range)
        
        candidates = []
        for res in raw_results:
            url = res.get("url")
            if not url or url in fetched_urls_shared:
                continue
            title = res.get("title", "Без названия").strip()
            snippet = res.get("content", "").strip()
            candidates.append({
                "url": url, 
                "title": title, 
                "snippet": snippet,
                "publish_date": res.get("publishedDate") or res.get("date")
            })
            
        logger.info(f"[PARALLEL] Для запроса '{current_query}' найдено {len(candidates)} потенциальных новых сайтов.")
        
        # Предварительно быстро проверяем кандидатов на доступность (скачиваем только заголовки)
        candidates_to_check = candidates[:10]
        valid_candidates = []
        
        if candidates_to_check:
            logger.info(f"[PARALLEL] Быстрая проверка доступности {len(candidates_to_check)} кандидатов...")
            
            if status_updater and search_queries_state:
                await status_updater(
                    make_status_html(
                        query=user_query,
                        status_text="Параллельный сбор информации..." if mode == "fast" else "Проверка доступности сайтов...",
                        progress_percent=progress_percent,
                        search_queries=search_queries_state,
                        attempt=iteration,
                        thoughts=current_thoughts
                    )
                )
                
            async def check_cand(cand):
                is_ok = await quick_check_url(cand["url"])
                if is_ok:
                    return cand
                return None

            tasks = [check_cand(c) for c in candidates_to_check]
            checked_results = await asyncio.gather(*tasks)
            valid_candidates = [res for res in checked_results if res is not None]
            
            # Все, которые не прошли быструю проверку, добавляем в failed_urls
            for c in candidates_to_check:
                if c["url"] not in [vc["url"] for vc in valid_candidates]:
                    failed_urls.append(c["url"])
                    
        # Проверяем количество успешно прошедших проверку кандидатов
        should_refine = False
        reason = ""
        
        if len(valid_candidates) < 3:
            logger.info(f"[PARALLEL] Доступных кандидатов после быстрой проверки меньше 3 ({len(valid_candidates)}). Автоматически переформулируем запрос.")
            should_refine = True
            reason = f"Найдено слишком мало доступных сайтов по этому запросу ({len(valid_candidates)} из 10 проверенных). Большинство сайтов заблокированы или недоступны."
        else:
            # Вызываем ИИ для выбора
            needed = target_count - len(success_sites)
            chosen_indexes = await ai_choose_best_sites(
                user_query=user_query,
                search_query=current_query,
                candidates=valid_candidates,
                downloaded_urls=[],
                failed_urls=[],
                needed_count=needed,
                model_name=model_name
            )
            
            if not chosen_indexes:
                logger.info(f"[PARALLEL] Модель ИИ не выбрала ни одного релевантного сайта.")
                should_refine = True
                reason = "ИИ-навигатор проанализировал доступные сайты, но не нашел среди них релевантных источников по теме исследования."
            else:
                # Скачиваем выбранные сайты полностью
                chosen_candidates = [valid_candidates[idx] for idx in chosen_indexes if idx < len(valid_candidates)]
                
                logger.info(f"[PARALLEL] Скачиваем {len(chosen_candidates)} выбранных ИИ сайтов полностью...")
                
                async def download_full(cand):
                    try:
                        content = await fetch_page_content(cand["url"])
                        if content and content.strip():
                            return {
                                "url": cand["url"],
                                "title": cand["title"],
                                "content": content,
                                "publish_date": cand["publish_date"]
                            }
                        else:
                            failed_urls.append(cand["url"])
                    except Exception as e:
                        logger.warning(f"[PARALLEL] Ошибка скачивания {cand['url']}: {e}")
                        failed_urls.append(cand["url"])
                    return None
                    
                download_tasks = [download_full(c) for c in chosen_candidates]
                downloaded_results = await asyncio.gather(*download_tasks)
                
                success_downloads = [res for res in downloaded_results if res is not None]
                
                for res in success_downloads:
                    fetched_urls_shared.add(res["url"])
                    
                    display_title = res["title"]
                    if len(display_title) > 40:
                        display_title = display_title[:37] + "..."
                        
                    if display_title not in query_state_item["sites"]:
                        query_state_item["sites"].append(display_title)
                        
                    success_sites.append({
                        "url": res["url"],
                        "title": display_title,
                        "content": res["content"],
                        "publish_date": res["publish_date"]
                    })
                
                if not success_sites:
                    logger.warning(f"[PARALLEL] Ни один из выбранных ИИ сайтов не удалось скачать.")
                    should_refine = True
                    reason = "Выбранные ИИ-моделью сайты не удалось скачать из-за непредвиденных ошибок сети или блокировок при полной загрузке."
                else:
                    query_state_item["sites"] = [res["title"] for res in success_sites]
                    logger.info(f"[PARALLEL] Успешно собрано {len(success_sites)} сайтов по запросу '{current_query}'. Завершаем поиск.")
                    return success_sites
                
        # Если нужно переформулировать
        if should_refine:
            attempts_left -= 1
            if attempts_left == 0:
                logger.warning(f"[PARALLEL] Лимит попыток исчерпан. Возвращаем {len(success_sites)} сайтов по запросу '{current_query}'.")
                query_state_item["sites"] = [res["title"] for res in success_sites]
                return success_sites
                
            failed_urls_str = "\n".join([f"- {url}" for url in failed_urls])
            
            if search_queries_state:
                for item in search_queries_state:
                    if item["query"] == current_query:
                        item["query"] = f"{current_query} 🔄 переформулируется..."
                        break
                        
            if status_updater and search_queries_state:
                await status_updater(
                    make_status_html(
                        query=user_query,
                        status_text=f"Запрос '{current_query}' не дал результатов. Переформулируем...",
                        progress_percent=progress_percent,
                        search_queries=search_queries_state,
                        attempt=iteration,
                        thoughts=current_thoughts
                    )
                )
                
            new_query = await generate_alternative_query(
                original_query=current_query,
                reason=reason,
                failed_urls=failed_urls_str,
                user_query=user_query,
                model_name=model_name
            )
            
            logger.info(f"[PARALLEL] Переформулировали '{current_query}' -> '{new_query}'. Осталось попыток: {attempts_left}")
            
            if search_queries_state:
                for item in search_queries_state:
                    if item["query"].startswith(current_query):
                        item["query"] = new_query
                        item["sites"] = []
                        break
                        
            current_query = new_query


def parse_date(date_str: str) -> datetime:
    if not date_str:
        return None
    # Пробуем разные форматы
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d.%m.%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(date_str[:19].strip(), fmt)
        except ValueError:
            continue
    # Попробуем вытащить год-месяц-день регулярным выражением
    m = re.search(r"(\d{4})[-./](\d{2})[-./](\d{2})", date_str)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except ValueError:
            pass
    m2 = re.search(r"(\d{2})[-./](\d{2})[-./](\d{4})", date_str)
    if m2:
        try:
            return datetime(int(m2.group(3)), int(m2.group(2)), int(m2.group(1)))
        except ValueError:
            pass
    return None

async def check_relevance(query: str, research_state: str, model_name: str = None) -> bool:
    system_prompt = (
        "Ты — эксперт-валидатор актуальности информации.\n"
        "Твоя задача — проанализировать собранные материалы исследования и определить, содержат ли они актуальную, свежую и полную информацию для ответа на запрос пользователя.\n"
        "Обрати внимание на даты публикаций, версии программного обеспечения, упоминания последних событий.\n"
        "Ответь строго в формате JSON:\n"
        "{\n"
        "  \"is_relevant\": true/false,\n"
        "  \"reason\": \"краткое объяснение, почему данные актуальны или устарели/неполны\"\n"
        "}\n"
        "Где:\n"
        "- is_relevant = true, если собранная информация актуальна на текущий момент (или в материалах есть самые последние данные) и повторный поиск не требуется.\n"
        "- is_relevant = false, если информация явно устарела (например, обсуждаются старые версии, старые модели ИИ, или в материалах отсутствуют свежие новости за последний год, хотя запрос подразумевает актуальность)."
    )
    
    prompt = (
        f"Запрос пользователя: {query}\n\n"
        f"Материалы исследования:\n{research_state}\n\n"
        "Проведи анализ актуальности и верни JSON с ключами is_relevant и reason."
    )
    
    try:
        response = await ask_ollama(prompt, system_prompt=system_prompt, model_name=model_name, timeout=45.0)
        clean_response = response
        if "<think>" in response and "</think>" in response:
            clean_response = response.split("</think>")[-1].strip()
            
        json_match = re.search(r"\{.*\}", clean_response, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group(0))
            is_relevant = data.get("is_relevant", True)
            logger.info(f"[RELEVANCE-CHECK] Результат проверки: {is_relevant}. Причина: {data.get('reason', 'нет')}")
            return is_relevant
    except Exception as e:
        logger.error(f"[RELEVANCE-CHECK] Ошибка проверки актуальности: {repr(e)}. По умолчанию считаем актуальным.")
    
    return True


async def run_multistep_search(query: str, user_id: int, status_updater, is_inline: bool, chat_id: int = None, status_msg: types.Message = None, inline_message_id: str = None, mode: str = "deep", base64_images: list[str] = None, context_summary: str = None) -> str:
    # 0. Если есть картинки, но нет описания - выполняем Vision-анализ в фоне
    if base64_images and not context_summary:
        await status_updater(
            make_status_html(
                query=query,
                status_text="Анализ картинки...",
                progress_percent=5,
                attempt=1
            )
        )
        try:
            descriptions = []
            for b64 in base64_images:
                desc = await describe_image(b64, model_name=None)
                descriptions.append(desc)
            
            if len(descriptions) == 1:
                context_summary = descriptions[0]
            else:
                context_summary = ""
                for idx, desc in enumerate(descriptions, 1):
                    context_summary += f"Изображение {idx}:\n{desc}\n\n"
                context_summary = context_summary.strip()
                
            logger.info(f"[VISION-FALLBACK] Успешно сгенерировано описание для {len(base64_images)} картинок во время поиска.")
        except Exception as ve:
            logger.error(f"[VISION-FALLBACK] Ошибка распознавания изображений во время поиска: {ve}")
            context_summary = "Ошибка распознавания изображения визионером."

    start_time = time.monotonic()
    active_session_ids.pop(user_id, None)
    chat_history = user_histories.get(user_id, [])
    fetched_urls = set()
    
    research_state = ""
    if chat_history and mode == "deep":
        history_parts = ["--- Предыдущий контекст диалога ---"]
        for msg in chat_history:
            role_label = "Пользователь" if msg["role"] == "user" else "Ассистент"
            history_parts.append(f"{role_label}: {msg['content']}")
        history_parts.append("-----------------------------------\n")
        research_state += "\n".join(history_parts) + "\n\n"
    dt = get_current_datetime_str()
    current_thoughts = None
    model_name = user_models.get(user_id, OLLAMA_MODEL)
    
    logger.info(f"[SEARCH-FLOW] Старт поиска. Запрос: '{query}', Режим: {mode}, Модель: {model_name}, Юзер ID: {user_id}, История сообщений: {len(chat_history)}, Картинок: {len(base64_images) if base64_images else 0}")
    current_model_var.set(model_name)
    current_status_updater_var.set(status_updater)
    current_query_var.set(query)
    current_status_text_var.set("Инициализация поиска...")
    current_progress_percent_var.set(0)
    current_search_queries_var.set(None)
    current_attempt_var.set(1)
    current_thoughts_var.set(None)
    
    log_chat_id = chat_id if chat_id is not None else user_id
    if mode == "deep":
        clear_research_log(log_chat_id)

    # 0. Классификация типа вопроса
    question_type = "COMPREHENSIVE"
    try:
        class_system_prompt = (
            "Ты — быстрый классификатор запросов.\n"
            "Определи, является ли вопрос пользователя точечным фактологическим (FACT) или широким/комплексным (COMPREHENSIVE).\n"
            "Примеры FACT:\n"
            "- 'какая высота Эвереста?'\n"
            "- 'в каком году родился Пушкин?'\n"
            "- 'версия python 3.13 дата выхода'\n"
            "- 'какая последняя модель у OpenAI?'\n"
            "Примеры COMPREHENSIVE:\n"
            "- 'как работает квантовый компьютер?'\n"
            "- 'сравни react и vue'\n"
            "- 'расскажи про историю искусственного интеллекта'\n"
            "Ответь строго одним словом: либо FACT, либо COMPREHENSIVE."
        )
        class_res = await ask_ollama(f"Запрос пользователя: {query}", system_prompt=class_system_prompt, model_name=model_name, timeout=10.0)
        class_res_clean = class_res.strip().upper()
        if "FACT" in class_res_clean:
            question_type = "FACT"
        elif "COMPREHENSIVE" in class_res_clean:
            question_type = "COMPREHENSIVE"
        logger.info(f"[CLASSIFICATION] Запрос '{query}' классифицирован как {question_type}")
    except Exception as e:
        logger.error(f"[CLASSIFICATION] Ошибка классификации запроса: {repr(e)}. Используем COMPREHENSIVE по умолчанию.")
        
    if mode == "fast":
        sites_limit = 10 if question_type == "FACT" else 5
    else:
        # Для глубокого поиска лимит на один запрос жестко от 3 до 5 сайтов
        sites_limit = 5
    
    for iteration in range(1, 4):
        current_thoughts = None
        await status_updater(
            make_status_html(
                query=query,
                status_text="Первичный сбор информации...",
                progress_percent=20,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        try:
            initial_queries, initial_thoughts = await generate_initial_queries(
                query=query,
                chat_history=chat_history,
                is_fast=(mode == "fast"),
                model_name=model_name,
                base64_images=base64_images,
                context_summary=context_summary
            )
        except Exception as e:
            err_text = "❌ К сожалению, все доступные ИИ-модели сейчас перегружены или недоступны. Пожалуйста, повторите запрос позже."
            logger.error(f"[SEARCH-FLOW] Критическая ошибка на этапе 1.1 (первичная генерация): {repr(e)}")
            await status_updater(err_text)
            raise ValueError(err_text)
        logger.info(f"Итерация {iteration} - Первичные запросы: {initial_queries}")
        if initial_thoughts and mode != "fast":
            current_thoughts = initial_thoughts
            
        search_queries_state = [{"query": q, "sites": []} for q in initial_queries]
        
        await status_updater(
            make_status_html(
                query=query,
                status_text="Первичный сбор информации...",
                progress_percent=30 if mode == "deep" else 50,
                search_queries=search_queries_state,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        async def delayed_task(item, delay):
            await asyncio.sleep(delay)
            return await process_single_query_parallel(
                query_str=item["query"],
                fetched_urls_shared=fetched_urls_shared,
                limit=sites_limit,
                query_state_item=item,
                mode=mode,
                model_name=model_name,
                user_query=query,
                status_updater=status_updater,
                search_queries_state=search_queries_state,
                iteration=iteration,
                current_thoughts=current_thoughts,
                progress_percent=30 if mode == "deep" else 50
            )
            
        fetched_urls_shared = set(fetched_urls)
        tasks = [
            delayed_task(item, i * 0.8)
            for i, item in enumerate(search_queries_state)
        ]
        
        await status_updater(
            make_status_html(
                query=query,
                status_text=f"Параллельный сбор информации по {len(initial_queries)} запросам...",
                progress_percent=30 if mode == "deep" else 50,
                search_queries=search_queries_state,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        results_lists = await asyncio.gather(*tasks)
        
        initial_sites = []
        for sites_list in results_lists:
            initial_sites.extend(sites_list)
            for s in sites_list:
                fetched_urls.add(s["url"])
                
        await status_updater(
            make_status_html(
                query=query,
                status_text="Первичный сбор информации успешно завершен.",
                progress_percent=35 if mode == "deep" else 60,
                search_queries=search_queries_state,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        # Конспектирование первичных материалов
        if mode == "deep":
            summary_tasks = []
            summary_states = {q_str: "⏳ ожидает..." for q_str in initial_queries}
            for idx, q_str in enumerate(initial_queries):
                q_sites = results_lists[idx] if idx < len(results_lists) else []
                summary_tasks.append(
                    generate_query_summary(
                        query_str=q_str,
                        user_query=query,
                        sites=q_sites,
                        model_name=model_name,
                        stage_name="Первичный сбор информации",
                        is_opinion=False,
                        chat_id=log_chat_id,
                        status_updater=status_updater,
                        summary_states=summary_states,
                        progress_percent=35,
                        iteration=iteration,
                        current_thoughts=current_thoughts,
                        search_queries_state=search_queries_state
                    )
                )
            
            # Сразу показываем статус ожидания/подготовки конспектирования
            status_text = make_summary_status_text("Первичный сбор информации", summary_states)
            await status_updater(
                make_status_html(
                    query=query,
                    status_text=status_text,
                    progress_percent=35,
                    search_queries=search_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            summaries = await asyncio.gather(*summary_tasks)
            initial_summary = "\n\n".join(summaries)
        else:
            # Для быстрого поиска просто объединяем очищенные тексты всех найденных сайтов
            parts = []
            for s in initial_sites:
                parts.append(f"--- Источник: {s['url']} ---\n{s['content']}\n")
            initial_summary = "\n".join(parts)
        
        research_state += f"\n--- Попытка {iteration} ---\n"
        research_state += f"Первичные запросы: {', '.join(initial_queries)}\n"
        research_state += f"Выжимка первичных материалов:\n{initial_summary}\n"
        
        if mode == "fast":
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Синтез финального ответа...",
                    progress_percent=80,
                    search_queries=search_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            break
        
        # --- ФАЗА 1.2: Уточняющий сбор ---
        current_thoughts = None
        await status_updater(
            make_status_html(
                query=query,
                status_text="Анализ и уточняющий сбор...",
                progress_percent=40,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        refined_sites = []
        refined_queries = []
        refined_queries_state = []
        try:
            accumulated_research_log = read_research_log(log_chat_id) if mode == "deep" else research_state
            refined_queries, refined_thoughts = await generate_refined_queries(
                query=query,
                chat_history=chat_history,
                research_state=accumulated_research_log,
                model_name=model_name,
                context_summary=context_summary
            )
            logger.info(f"Итерация {iteration} - Уточняющие запросы: {refined_queries}")
            if refined_thoughts:
                current_thoughts = refined_thoughts
                
            refined_queries_state = [{"query": q, "sites": []} for q in refined_queries]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Анализ и уточняющий сбор...",
                    progress_percent=50,
                    search_queries=refined_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            async def delayed_task_refined(item, delay):
                await asyncio.sleep(delay)
                return await process_single_query_parallel(
                    query_str=item["query"],
                    fetched_urls_shared=fetched_urls_shared,
                    limit=sites_limit,
                    query_state_item=item,
                    mode=mode,
                    model_name=model_name,
                    user_query=query,
                    status_updater=status_updater,
                    search_queries_state=refined_queries_state,
                    iteration=iteration,
                    current_thoughts=current_thoughts,
                    progress_percent=50
                )
                
            fetched_urls_shared = set(fetched_urls)
            tasks = [
                delayed_task_refined(item, i * 0.8)
                for i, item in enumerate(refined_queries_state)
            ]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Параллельный сбор уточняющей информации...",
                    progress_percent=50,
                    search_queries=refined_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            results_lists = await asyncio.gather(*tasks)
            
            for sites_list in results_lists:
                refined_sites.extend(sites_list)
                for s in sites_list:
                    fetched_urls.add(s["url"])
                    
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Уточняющий сбор информации успешно завершен.",
                    progress_percent=55,
                    search_queries=refined_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
        except Exception as e:
            logger.error(f"[SEARCH-FLOW] Пропуск уточняющего сбора из-за сбоя ИИ: {repr(e)}")
            results_lists = []
            
        # Конспектирование уточняющих материалов
        if mode == "deep":
            summary_tasks = []
            summary_states = {q_str: "⏳ ожидает..." for q_str in refined_queries}
            for idx, q_str in enumerate(refined_queries):
                q_sites = results_lists[idx] if idx < len(results_lists) else []
                summary_tasks.append(
                    generate_query_summary(
                        query_str=q_str,
                        user_query=query,
                        sites=q_sites,
                        model_name=model_name,
                        stage_name="Уточняющий сбор информации",
                        is_opinion=False,
                        chat_id=log_chat_id,
                        status_updater=status_updater,
                        summary_states=summary_states,
                        progress_percent=55,
                        iteration=iteration,
                        current_thoughts=current_thoughts,
                        search_queries_state=refined_queries_state
                    )
                )
            
            # Сразу показываем статус ожидания/подготовки конспектирования
            status_text = make_summary_status_text("Уточняющий сбор информации", summary_states)
            await status_updater(
                make_status_html(
                    query=query,
                    status_text=status_text,
                    progress_percent=55,
                    search_queries=refined_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            summaries = await asyncio.gather(*summary_tasks)
            refined_summary = "\n\n".join(summaries)
        else:
            refined_summary = ""
        
        if refined_queries:
            research_state += f"Уточняющие запросы: {', '.join(refined_queries)}\n"
            research_state += f"Выжимка уточняющих материалов:\n{refined_summary}\n"
        
        # --- ЭТАП 2.1: Сбор мнений ---
        current_thoughts = None
        await status_updater(
            make_status_html(
                query=query,
                status_text="Сбор мнений, отзывов и опыта...",
                progress_percent=60,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        opinion_sites = []
        opinion_queries = []
        opinion_queries_state = []
        try:
            accumulated_research_log = read_research_log(log_chat_id) if mode == "deep" else research_state
            opinion_queries, opinion_thoughts = await generate_opinion_queries(
                query=query,
                chat_history=chat_history,
                research_state=accumulated_research_log,
                model_name=model_name,
                context_summary=context_summary
            )
            logger.info(f"Итерация {iteration} - Запросы мнений: {opinion_queries}")
            if opinion_thoughts:
                current_thoughts = opinion_thoughts
                
            opinion_queries_state = [{"query": q, "sites": []} for q in opinion_queries]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Сбор мнений, отзывов и опыта...",
                    progress_percent=70,
                    search_queries=opinion_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            async def delayed_task_opinion(item, delay):
                await asyncio.sleep(delay)
                return await process_single_query_parallel(
                    query_str=item["query"],
                    fetched_urls_shared=fetched_urls_shared,
                    limit=sites_limit,
                    query_state_item=item,
                    mode=mode,
                    model_name=model_name,
                    user_query=query,
                    status_updater=status_updater,
                    search_queries_state=opinion_queries_state,
                    iteration=iteration,
                    current_thoughts=current_thoughts,
                    progress_percent=70
                )
                
            fetched_urls_shared = set(fetched_urls)
            tasks = [
                delayed_task_opinion(item, i * 0.8)
                for i, item in enumerate(opinion_queries_state)
            ]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Параллельный сбор мнений, отзывов и опыта...",
                    progress_percent=70,
                    search_queries=opinion_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            results_lists = await asyncio.gather(*tasks)
            
            for sites_list in results_lists:
                opinion_sites.extend(sites_list)
                for s in sites_list:
                    fetched_urls.add(s["url"])
                    
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Сбор мнений успешно завершен.",
                    progress_percent=75,
                    search_queries=opinion_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
        except Exception as e:
            logger.error(f"[SEARCH-FLOW] Пропуск сбора мнений из-за сбоя ИИ: {repr(e)}")
            results_lists = []
            
        # Конспектирование мнений и отзывов
        if mode == "deep":
            summary_tasks = []
            summary_states = {q_str: "⏳ ожидает..." for q_str in opinion_queries}
            for idx, q_str in enumerate(opinion_queries):
                q_sites = results_lists[idx] if idx < len(results_lists) else []
                summary_tasks.append(
                    generate_query_summary(
                        query_str=q_str,
                        user_query=query,
                        sites=q_sites,
                        model_name=model_name,
                        stage_name="Сбор мнений и отзывов",
                        is_opinion=True,
                        chat_id=log_chat_id,
                        status_updater=status_updater,
                        summary_states=summary_states,
                        progress_percent=75,
                        iteration=iteration,
                        current_thoughts=current_thoughts,
                        search_queries_state=opinion_queries_state
                    )
                )
            
            # Сразу показываем статус ожидания/подготовки конспектирования
            status_text = make_summary_status_text("Сбор мнений и отзывов", summary_states)
            await status_updater(
                make_status_html(
                    query=query,
                    status_text=status_text,
                    progress_percent=75,
                    search_queries=opinion_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            summaries = await asyncio.gather(*summary_tasks)
            opinion_summary = "\n\n".join(summaries)
        else:
            opinion_summary = ""
        
        if opinion_queries:
            research_state += f"Запросы мнений: {', '.join(opinion_queries)}\n"
            research_state += f"Выжимка мнений/отзывов:\n{opinion_summary}\n"
        
        # --- ЭТАП 2.2: Перекрестный анализ и поиск ---
        cross_sites = []
        cross_queries = []
        cross_queries_state = []
        try:
            current_thoughts = None
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Анализ противоречий и подготовка проверочных запросов...",
                    progress_percent=76,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            accumulated_research_log = read_research_log(log_chat_id) if mode == "deep" else research_state
            cross_queries, cross_thoughts = await generate_cross_queries(
                query=query,
                chat_history=chat_history,
                research_state=accumulated_research_log,
                model_name=model_name,
                context_summary=context_summary
            )
            logger.info(f"Итерация {iteration} - Проверочные запросы: {cross_queries}")
            if cross_thoughts:
                current_thoughts = cross_thoughts
                
            cross_queries_state = [{"query": q, "sites": []} for q in cross_queries]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Параллельный сбор данных для перекрестной проверки...",
                    progress_percent=80,
                    search_queries=cross_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            async def delayed_task_cross(item, delay):
                await asyncio.sleep(delay)
                return await process_single_query_parallel(
                    query_str=item["query"],
                    fetched_urls_shared=fetched_urls_shared,
                    limit=sites_limit,
                    query_state_item=item,
                    mode=mode,
                    model_name=model_name,
                    user_query=query,
                    status_updater=status_updater,
                    search_queries_state=cross_queries_state,
                    iteration=iteration,
                    current_thoughts=current_thoughts,
                    progress_percent=80,
                    time_range="month"
                )
                
            fetched_urls_shared = set(fetched_urls)
            tasks = [
                delayed_task_cross(item, i * 0.8)
                for i, item in enumerate(cross_queries_state)
            ]
            
            results_lists = await asyncio.gather(*tasks)
            
            for sites_list in results_lists:
                cross_sites.extend(sites_list)
                for s in sites_list:
                    fetched_urls.add(s["url"])
                    
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Сбор проверочных материалов успешно завершен.",
                    progress_percent=85,
                    search_queries=cross_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
        except Exception as e:
            logger.error(f"[SEARCH-FLOW] Пропуск перекрестного сбора из-за сбоя ИИ: {repr(e)}")
            results_lists = []
            
        # Конспектирование проверочных материалов
        if mode == "deep":
            summary_tasks = []
            summary_states = {q_str: "⏳ ожидает..." for q_str in cross_queries}
            for idx, q_str in enumerate(cross_queries):
                q_sites = results_lists[idx] if idx < len(results_lists) else []
                summary_tasks.append(
                    generate_query_summary(
                        query_str=q_str,
                        user_query=query,
                        sites=q_sites,
                        model_name=model_name,
                        stage_name="Перекрестная проверка материалов",
                        is_opinion=False,
                        chat_id=log_chat_id,
                        status_updater=status_updater,
                        summary_states=summary_states,
                        progress_percent=85,
                        iteration=iteration,
                        current_thoughts=current_thoughts,
                        search_queries_state=cross_queries_state
                    )
                )
            
            # Сразу показываем статус ожидания/подготовки конспектирования
            status_text = make_summary_status_text("Перекрестная проверка материалов", summary_states)
            await status_updater(
                make_status_html(
                    query=query,
                    status_text=status_text,
                    progress_percent=85,
                    search_queries=cross_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            summaries = await asyncio.gather(*summary_tasks)
            cross_summary = "\n\n".join(summaries)
        else:
            cross_summary = ""
        
        if cross_queries:
            research_state += f"Проверочные запросы перекрестного анализа: {', '.join(cross_queries)}\n"
            research_state += f"Выжимка материалов перекрестной проверки:\n{cross_summary}\n"

        if mode == "deep":
            general_combined_context = read_research_log(log_chat_id)
            opinions_context = general_combined_context
        else:
            # Для быстрого поиска просто берем весь накопленный лог из research_state
            general_combined_context = research_state
            opinions_context = research_state
        
        logger.info(f"[SEARCH-FLOW] Запуск перекрестного анализа. Символов фактов/свежести: {len(general_combined_context)}, мнений: {len(opinions_context)}")
        await status_updater(
            make_status_html(
                query=query,
                status_text="Выполняем перекрестный анализ найденных материалов...",
                progress_percent=88,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        try:
            cross_report = await run_cross_verification(query, general_combined_context, opinions_context, model_name=model_name)
            logger.info(f"[SEARCH-FLOW] Успешно завершен перекрестный анализ. Размер отчета: {len(cross_report)} символов.")
            research_state += f"Результат перекрестного анализа:\n{cross_report}\n"
        except Exception as e:
            logger.error(f"[SEARCH-FLOW] Ошибка кросс-верификации: {repr(e)}")
            research_state += "Результат перекрестного анализа: ошибка обработки.\n"
            
        logger.info(f"[SEARCH-FLOW] Запуск проверки актуальности. Общая длина накопленного лога: {len(research_state)} символов.")
        await status_updater(
            make_status_html(
                query=query,
                status_text="Проверяем актуальность и свежесть информации...",
                progress_percent=93,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        try:
            is_current = await check_relevance(query, research_state, model_name=model_name)
            logger.info(f"[SEARCH-FLOW] Результат проверки актуальности: {is_current}")
        except Exception as e:
            logger.error(f"[SEARCH-FLOW] Ошибка проверки актуальности: {repr(e)}")
            is_current = True
            
        if is_current:
            current_thoughts = None
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Синтез финального ответа...",
                    progress_percent=100,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            break
        else:
            research_state += "Результат проверки актуальности: НЕАКТУАЛЬНО. Необходим повторный поиск с уточнением временных параметров.\n"
            if iteration < 3:
                await status_updater(
                    make_status_html(
                        query=query,
                        status_text="Данные устарели. Запуск повторного поиска с уточнением...",
                        progress_percent=95,
                        attempt=iteration,
                        warning_text="Обнаружены устаревшие данные. Уточняем временные рамки...",
                        thoughts=current_thoughts
                    )
                )
                await asyncio.sleep(3.0)
            else:
                current_thoughts = None
                await status_updater(
                    make_status_html(
                        query=query,
                        status_text="Синтез финального ответа...",
                        progress_percent=100,
                        attempt=iteration,
                        warning_text="Достигнут лимит попыток. Модель ответит на основе имеющихся данных.",
                        thoughts=current_thoughts
                    )
                )
                
    # --- СИНТЕЗ ФИНАЛЬНОГО ОТВЕТА (СТРИМИНГ С КРАСИВЫМ HTML-ФОРМАТИРОВАНИЕМ) ---
    inline_brief_instruction = (
        "КРИТИЧЕСКИ ВАЖНО: Ты отвечаешь в инлайн-режиме Telegram. Твой ответ должен быть КОРОТКИМ и ЛАКОНИЧНЫМ (не пиши длинные подробные тексты, давай только самую суть). При этом сохраняй стандартное форматирование в виде абзацев и маркированных списков, как указано в правилах ниже, но пиши строго по существу и очень кратко.\n"
    ) if is_inline else ""

    if is_inline:
        structure_rules_fast = (
            "ФОРМАТИРОВАНИЕ И СТРУКТУРА ОТВЕТА (ДИЗАЙН-КОД):\n"
            "Пиши ответ с использованием стандартного Markdown-форматирования (наш парсер автоматически переведет его в красивый Telegram HTML). Поскольку ты отвечаешь в компактном инлайн-режиме, твой ответ должен быть МАКСИМАЛЬНО КРАТКИМ и емким. Твой ответ должен быть отформатирован в виде 1-3 коротких абзацев (допускается один небольшой список), без раздувания объема и без длинных перечислений:\n"
            "1. ГЛАВНАЯ ВЫЖИМКА: Начни ответ с краткого прямого ответа на вопрос в 1-2 предложения, обернутого в цитату Markdown (напиши знак `>` в начале строк цитаты).\n"
            "2. ДЕТАЛИ: Приведи ключевые детали в виде 1-2 коротких абзацев или небольшого списка. Запрещено использовать заголовки `###`.\n"
            "3. АКЦЕНТЫ: Обязательно выделяй ключевые термины, важные имена и важные цифры жирным шрифтом (`**жирный**`).\n"
            "4. ОТСТУПЫ И ЧИТАЕМОСТЬ: Обязательно делай пустые строки (двойной перенос строки) между отдельными абзацами текста, цитатой-выжимкой и списком, чтобы текст визуально не слипался.\n"
            "5. ЛАКОНИЧНОСТЬ: Пиши строго по существу, общим объемом не более 80-100 слов. Никакой воды, общих введений и предысторий.\n"
            "Не используй цветные эмодзи (смайлики). Для разделителей используй строгие символы Юникода (например, ───, ✦, ▪, ❖, ➔)."
        )
        structure_rules_deep = (
            "ФОРМАТИРОВАНИЕ И СТРУКТУРА ОТВЕТА (ДИЗАЙН-КОД):\n"
            "Пиши финальный ответ (после </think>) с использованием стандартного Markdown-форматирования (наш парсер автоматически переведет его в красивый Telegram HTML). Поскольку ты отвечаешь в компактном инлайн-режиме, твой ответ должен быть МАКСИМАЛЬНО КРАТКИМ и емким. Твой ответ должен быть отформатирован в виде 1-3 коротких абзацев (допускается один небольшой список), без раздувания объема и без длинных перечислений:\n"
            "1. ГЛАВНАЯ ВЫЖИМКА: Начни ответ с краткого прямого ответа на вопрос в 1-2 предложения, обернутого в цитату Markdown (напиши знак `>` в начале строк цитаты).\n"
            "2. ДЕТАЛИ: Приведи ключевые детали в виде 1-2 коротких абзацев или небольшого списка. Запрещено использовать заголовки `###`.\n"
            "3. АКЦЕНТЫ: Обязательно выделяй ключевые термины, важные имена и важные цифры жирным шрифтом (`**жирный**`).\n"
            "4. ОТСТУПЫ И ЧИТАЕМОСТЬ: Обязательно делай пустые строки (двойной перенос строки) между отдельными абзацами текста, цитатой-выжимкой и списком, чтобы текст визуально не слипался.\n"
            "5. ЛАКОНИЧНОСТЬ: Пиши строго по существу, общим объемом не более 80-100 слов. Никакой воды, общих введений и предысторий.\n"
            "Не используй цветные эмодзи (смайлики). Для разделителей используй строгие символы Юникода (например, ───, ✦, ▪, ❖, ➔)."
        )
    else:
        structure_rules_fast = (
            "ФОРМАТИРОВАНИЕ И СТРУКТУРА ОТВЕТА (ДИЗАЙН-КОД):\n"
            "Пиши ответ с использованием стандартного Markdown-форматирования (наш парсер автоматически переведет его в красивый Telegram HTML). Твой ответ ОБЯЗАТЕЛЬНО должен быть структурированным и красиво оформленным:\n"
            "1. ГЛАВНАЯ ВЫЖИМКА: Начни ответ с краткого прямого ответа на вопрос в 1-2 предложения, обернутого в цитату Markdown (напиши знак `>` в начале строк цитаты), чтобы пользователь мог понять суть за секунду.\n"
            "2. ДЕТАЛИ ПО СУЩЕСТВУ: Далее приведи сжатые подробности, разбивая их на маркированные списки (через дефис `-` или звездочку `*`) и небольшие блоки с четкими заголовками (`### заголовок`).\n"
            "3. АКЦЕНТЫ: Обязательно выделяй ключевые термины, важные имена и важные цифры жирным шрифтом (`**жирный**`), чтобы текст легко сканировался взглядом.\n"
            "4. ОТСТУПЫ И ЧИТАЕМОСТЬ: Обязательно делай пустые строки (двойной перенос строки) между отдельными абзацами текста, цитатой-выжимкой, заголовками, списками и отдельными логическими блоками информации, чтобы текст визуально не слипался и легко читался.\n"
            "5. ЛАКОНИЧНОСТЬ: Будь предельно кратким и конкретным. Никакой воды, общих введений, рассуждений и академической перегрузки.\n"
            "Не используй цветные эмодзи (смайлики). Для разделителей используй строгие символы Юникода (например, ───, ✦, ▪, ❖, ➔)."
        )
        structure_rules_deep = (
            "ФОРМАТИРОВАНИЕ И СТРУКТУРА ОТВЕТА (ДИЗАЙН-КОД):\n"
            "Пиши финальный ответ (после </think>) с использованием стандартного Markdown-форматирования (наш парсер автоматически переведет его в красивый Telegram HTML). Твой ответ ОБЯЗАТЕЛЬНО должен быть структурированным и красиво оформленным:\n"
            "1. ГЛАВНАЯ ВЫЖИМКА: Начни ответ с краткого прямого ответа на вопрос в 1-2 предложения, обернутого в цитату Markdown (напиши знак `>` в начале строк цитаты), чтобы пользователь мог понять суть за секунду.\n"
            "2. ДЕТАЛИ ПО СУЩЕСТВУ: Далее приведи сжатые подробности, разбивая их на маркированные списки (через дефис `-` или звездочку `*`) и небольшие блоки с четкими заголовками (`### заголовок`).\n"
            "3. АКЦЕНТЫ: Обязательно выделяй ключевые термины, важные имена и важные цифры жирным шрифтом (`**жирный**`), чтобы текст легко сканировался взглядом.\n"
            "4. ОТСТУПЫ И ЧИТАЕМОСТЬ: Обязательно делай пустые строки (двойной перенос строки) между отдельными абзацами текста, цитатой-выжимкой, заголовками, списками и отдельными логическими блоками информации, чтобы текст визуально не слипался и легко читался.\n"
            "5. ЛАКОНИЧНОСТЬ: Будь предельно кратким и конкретным. Никакой воды, общих введений, рассуждений и академической перегрузки.\n"
            "Не используй цветные эмодзи (смайлики). Для разделителей используй строгие символы Юникода (например, ───, ✦, ▪, ❖, ➔)."
        )

    if mode == "fast":
        system_prompt = (
            f"Текущие дата и время: {dt}.\n"
            "Ты — полезный ИИ-помощник, который отвечает на вопрос пользователя на основе проведенного исследования.\n"
            "КРИТИЧЕСКИ ВАЖНО: Твой финальный ответ должен быть СТРОГО НА РУССКОМ ЯЗЫКЕ (in Russian), независимо от того, на каком языке был запрос пользователя или на каком языке написаны найденные в поиске источники. Даже если все источники англоязычные, переведи и изложи всю суть исключительно на русском языке.\n"
            f"{inline_brief_instruction}"
            "АНАЛИЗ КОНКРЕТНОСТИ ВОПРОСА:\n"
            "- Если вопрос конкретный, точечный и узкий (например, требует точное имя, дату, число, команду, ответ «да/нет» или один точный факт) и в нем нет просьбы расписать подробно, отвечай максимально КРАТКО И ПРЯМО (буквально в одно предложение или даже одно слово/число), без какого-либо расширения темы, предыстории или лишнего контекста.\n"
            "- Если вопрос широкий, открытый или размытый (например, «как работает X», «почему Y», «объясни Z»), твой ответ должен быть максимально емким, сфокусированным строго на ответе по существу, без раздувания объема текста. Давай только сухую выжимку фактов (активно используя списки), полностью исключив лирические отступления и воду.\n"
            "КРИТИЧЕСКИ ВАЖНО: Сразу пиши финальный ответ. Запрещено использовать тег <think> и писать ход мыслей/рассуждения.\n"
            "ПРАВИЛО ИСКЛЮЧЕНИЯ УСТАРЕВШИХ ДАННЫХ И НЕПРЕДВЗЯТОСТИ:\n"
            "- Помни, что сейчас на дворе 2026 год. Если в отчете исследования упоминаются устаревшие версии или старые модели продуктов как «последние», но из контекста вопроса очевидно, что ищется самая новая актуальная информация, НЕ утверждай, что эти старые модели являются последними на сегодняшний день.\n"
            "- Если в найденных материалах нет явных подтверждений о релизах новых версий, прямо укажи, что точное название новой модели в источниках за 2026 год не приведено, и перечисли только достоверные факты, избегая домыслов.\n"
            f"{structure_rules_fast}"
        )
    else:
        system_prompt = (
            f"Текущие дата и время: {dt}.\n"
            "Ты — полезный ИИ-помощник, который отвечает на вопрос пользователя на основе проведенного исследования.\n"
            "КРИТИЧЕСКИ ВАЖНО: Твой финальный ответ должен быть СТРОГО НА РУССКОМ ЯЗЫКЕ (in Russian), независимо от того, на каком языке был запрос пользователя или на каком языке написаны найденные в поиске источники. Даже если все источники англоязычные, переведи и изложи всю суть исключительно на русском языке.\n"
            f"{inline_brief_instruction}"
            "КРИТИЧЕСКИ ВАЖНО: Начни свой ответ со своего хода мыслей, завернутого в тег <think>, например:\n"
            "<think>\nЗдесь твои подробные рассуждения и логика на русском языке. Пиши только чистый простой текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков).\n</think>\n"
            "После закрывающего тега </think> сразу напиши свой финальный ответ для пользователя. КРИТИЧЕСКИ ВАЖНО: Никаких вводных слов, пояснений, комментариев вроде 'Вот финальный ответ:', 'Хорошо, я понял задачу...', 'Я проанализировал...' и других мета-сообщений быть НЕ должно. Сразу выводи структурированную полезную суть ответа.\n"
            "АНАЛИЗ КОНКРЕТНОСТИ ВОПРОСА И ЛАКОНИЧНОСТЬ:\n"
            "- Твой ответ в глубоком поиске (deep mode) должен быть направлен строго по существу, без раздувания объемов текста. По умолчанию делай ответ лаконичным, емким и сжатым. Давай только сухую выжимку фактов (активно используя маркированные списки для перечислений), полностью исключив лирические отступления, предыстории и воду. Запрещено искусственно увеличивать длину ответа.\n"
            "- Если вопрос конкретный, точечный и узкий (например, требует точное имя, дату, число, команду, ответ «да/нет» или один точный факт) и в нем нет просьбы расписать подробно, отвечай максимально КРАТКО И ПРЯМО (буквально в одно предложение или даже одним словом/числом), без какого-либо расширения темы, предыстории или лишнего контекста.\n"
            "ПРАВИЛО ИСКЛЮЧЕНИЯ УСТАРЕВШИХ ДАННЫХ И НЕПРЕДВЗЯТОСТИ:\n"
            "- Помни, что сейчас на дворе 2026 год. Если в отчете исследования упоминаются устаревшие версии или старые модели продуктов как «последние», но из контекста вопроса очевидно, что ищется самая новая актуальная информация, НЕ утверждай, что эти старые модели являются последними на сегодняшний день.\n"
            "- Если в найденных материалах нет явных подтверждений о релизах новых версий, прямо укажи, что точное название новой модели в источниках за 2026 год не приведено, и перечисли только достоверные факты, избегая домыслов.\n"
            f"{structure_rules_deep}\n\n"
            "Правила безопасности HTML:\n"
            "- ЗАПРЕЩЕНО использовать технические заголовки этапов исследования и термины методологии в тексте финального ответа (например, писать заголовки или фразы 'Перекрестный анализ', 'Сбор мнений', 'Этап 1', 'Результаты проверки'). Текст должен быть цельным, естественным и сфокусированным исключительно на ответе по существу.\n"
            "- Если в ходе исследования обнаружены явные расхождения между официальными фактами и мнениями/отзывами, аккуратно и органично вплети это в плавное повествование (например: 'Хотя официально заявлено X, на практике пользователи часто сталкиваются с Y'), не создавая для этого искусственных разделов. Если расхождений нет, отвечай строго по фактам, не упоминая отсутствие отзывов или мнений."
        )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    if mode == "fast":
        prompt = (
            f"Вопрос пользователя: {query}\n\n"
            f"Полный отчет о проведенном исследовании:\n{research_state}\n\n"
            "Сформулируй финальный ответ для пользователя с использованием стандартного Markdown-форматирования согласно правилам в системном промпте."
        )
    else:
        prompt = (
            f"Вопрос пользователя: {query}\n\n"
            f"Полный отчет о проведенном исследовании и результаты перекрестного анализа:\n{research_state}\n\n"
            "Сформулируй финальный ответ для пользователя с использованием стандартного Markdown-форматирования согласно правилам в системном промпте."
        )
        
    if base64_images and is_model_multimodal(model_name):
        content_structure = [{"type": "text", "text": prompt}]
        for b64 in base64_images:
            content_structure.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})
        messages.append({"role": "user", "content": content_structure})
    else:
        if base64_images and context_summary:
            prompt = f"Контекст присланных пользователем фото (Описание от визионера):\n{context_summary}\n\n" + prompt
        messages.append({"role": "user", "content": prompt})
        
    logger.info(f"[SEARCH-FLOW] Запуск финального синтеза ответа. Общее время поиска до этого момента: {time.monotonic() - start_time:.2f} сек. Размер лога исследования: {len(research_state)} симв.")
    num_sources = len(fetched_urls)
    if is_inline:
        answer_text = await stream_to_inline(inline_message_id, messages, model_name=model_name, is_fast=(mode == "fast"), start_time=start_time, num_sources=num_sources)
    else:
        answer_text = await stream_to_message(status_msg, messages, model_name=model_name, is_fast=(mode == "fast"), start_time=start_time, num_sources=num_sources)
    logger.info(f"[SEARCH-FLOW] Финальный ответ успешно отправлен и сохранен. Длина ответа: {len(answer_text)} симв. Полное время: {time.monotonic() - start_time:.2f} сек.")
        
    # Сохраняем историю
    if user_id not in user_histories:
        user_histories[user_id] = []
        
    if base64_images:
        history_query = f"[Фото: {context_summary or 'Изображения'}] {query}"
    elif context_summary and "Содержимое прикрепленного файла" in query:
        history_query = f"[Файл] {query.split('Вопрос пользователя по файлу:')[-1].strip() if 'Вопрос пользователя по файлу:' in query else query[:100]}"
    else:
        history_query = query
        
    user_histories[user_id].append({"role": "user", "content": history_query})
    user_histories[user_id].append({"role": "assistant", "content": answer_text})
    
    # Ограничиваем историю до 10 реплик (5 пар диалогов)
    user_histories[user_id] = user_histories[user_id][-10:]
    save_user_histories()
    
    return answer_text

@dp.message(Command("start"))
async def start_cmd(message: types.Message):
    try:
        await message.delete()
    except Exception:
        pass
    await message.answer(
        "👋 Привет! Я инлайн-агент для поиска информации в интернете.\n\n"
        "Для работы со мной начните писать в любом чате: <code>@имя_этого_бота ваш запрос</code> "
        "и выберите предложенный вариант. Сообщение будет отправлено и динамически обновлено на основе результатов поиска и анализа ИИ."
    )

async def clear_cmd(message: types.Message):
    try:
        await message.delete()
    except Exception:
        pass
        
    user_id = message.from_user.id
    chat_id = message.chat.id
    
    # 1. Проверяем, есть ли что сохранять
    history = user_histories.get(user_id, [])
    research_log = read_research_log(chat_id)
    
    # Если сессия была просто восстановлена и в неё не писалось ничего нового
    is_restored_without_changes = (user_id in active_session_ids)
    
    saved_title = None
    if (history or research_log) and not is_restored_without_changes:
        status_save = await message.answer("⏳ Анализирую диалог и сохраняю сессию в архив...")
        try:
            user_model = user_models.get(user_id) or OLLAMA_MODEL
            title = await generate_session_title(history, model_name=user_model)
            await save_session_to_db(
                user_id=user_id,
                chat_id=chat_id,
                title=title,
                research_log=research_log,
                chat_history=history
            )
            saved_title = title
        except Exception as e:
            logger.error(f"Ошибка сохранения сессии при очистке: {e}")
        finally:
            try:
                await status_save.delete()
            except Exception:
                pass

    # 2. Очищаем историю и логи
    if user_id in user_histories:
        user_histories[user_id] = []
        save_user_histories()
        
    active_session_ids.pop(user_id, None)
    clear_research_log(chat_id)
    logger.info(f"История диалога и лог исследования для пользователя {user_id} очищены.")
        
    current_id = message.message_id
    deleted_count = 0
    status = await message.answer("⏳ Очищаю историю сообщений в чате (до 100 недавних)...")
    
    for msg_id in range(current_id, current_id - 100, -1):
        if msg_id == status.message_id:
            continue
        try:
            await bot.delete_message(chat_id=chat_id, message_id=msg_id)
            deleted_count += 1
        except Exception:
            continue
            
    try:
        bot_user = await bot.get_me()
        bot_username = bot_user.username or "SearchXNG_Bot"
        
        info_text = "🧹 <b>История сообщений успешно очищена! Бот готов к работе.</b>\n\n"
        if saved_title:
            info_text += f"📦 Текущая сессия сохранена в архив под названием: <i>«{saved_title}»</i>.\n\n"
            
        info_text += (
            "Вы можете отправить мне новый поисковый запрос прямо в этом чате или использовать меня в инлайн-режиме:\n"
            f"🔍 <code>@{bot_username} ваш запрос</code>"
        )
        
        await status.edit_text(info_text, parse_mode=ParseMode.HTML)
        last_clear_message_ids[user_id] = status.message_id
    except Exception as e:
        logger.error(f"Ошибка при обновлении статуса очистки: {e}")

async def sessions_cmd(message: types.Message):
    try:
        await message.delete()
    except Exception:
        pass
    user_id = message.from_user.id
    await show_sessions_page(message.chat.id, user_id, page=0)

SESSIONS_PER_PAGE = 5

async def show_sessions_page(chat_id: int, user_id: int, page: int, edit_message_id: int = None):
    offset = page * SESSIONS_PER_PAGE
    sessions = get_user_sessions(user_id, limit=SESSIONS_PER_PAGE, offset=offset)
    total_count = get_user_sessions_count(user_id)
    total_pages = max(1, math.ceil(total_count / SESSIONS_PER_PAGE))
    
    text = "🗄 <b>Архив ваших поисковых сессий:</b>\n"
    if total_count == 0:
        text += "\nУ вас пока нет сохраненных сессий. Они автоматически сохраняются при очистке диалога с помощью /clear."
        keyboard = types.InlineKeyboardMarkup(inline_keyboard=[])
    else:
        text += f"*(Всего сохранено: {total_count}, страница {page+1} из {total_pages})*"
        
        inline_keyboard = []
        for idx, (s_id, title, created_at) in enumerate(sessions):
            try:
                dt = datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S")
                date_str = dt.strftime("%d.%m.%y %H:%M")
            except Exception:
                date_str = created_at
            
            button_text = f"{offset + idx + 1}. {title} ({date_str})"
            if len(button_text) > 42:
                button_text = button_text[:39] + "..."
                
            inline_keyboard.append([
                types.InlineKeyboardButton(text=button_text, callback_data=f"sess_view:{s_id}:{page}")
            ])
            
        nav_row = []
        if page > 0:
            nav_row.append(types.InlineKeyboardButton(text="◀️ Назад", callback_data=f"sess_page:{page-1}"))
        else:
            nav_row.append(types.InlineKeyboardButton(text="⏹", callback_data="sess_noop"))
            
        nav_row.append(types.InlineKeyboardButton(text=f"{page+1}/{total_pages}", callback_data="sess_noop"))
        
        if page < total_pages - 1:
            nav_row.append(types.InlineKeyboardButton(text="Вперед ▶️", callback_data=f"sess_page:{page+1}"))
        else:
            nav_row.append(types.InlineKeyboardButton(text="⏹", callback_data="sess_noop"))
            
        inline_keyboard.append(nav_row)
        inline_keyboard.append([
            types.InlineKeyboardButton(text="❌ Закрыть архив", callback_data="sess_close")
        ])
        
        keyboard = types.InlineKeyboardMarkup(inline_keyboard=inline_keyboard)
        
    try:
        if edit_message_id:
            await bot.edit_message_text(
                text=text,
                chat_id=chat_id,
                message_id=edit_message_id,
                parse_mode=ParseMode.HTML,
                reply_markup=keyboard
            )
        else:
            await bot.send_message(
                chat_id=chat_id,
                text=text,
                parse_mode=ParseMode.HTML,
                reply_markup=keyboard
            )
    except Exception as e:
        logger.error(f"Ошибка при отображении страницы сессий: {e}")

api_models_cache = {"models": [], "expiry": 0}

async def get_available_models_from_api() -> list[str]:
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/models"
    headers = {}
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            data = response.json()
            models = [m["id"] for m in data.get("data", [])]
            models.sort()
            
            # Проверяем новые модели ленивым образом
            new_models = [m for m in models if m not in multimodal_cache]
            if new_models:
                logger.info(f"Обнаружены новые модели для проверки мультимодальности: {new_models}")
                tasks = [check_model_multimodal(m_id) for m_id in new_models]
                results = await asyncio.gather(*tasks, return_exceptions=True)
                
                cache_updated = False
                for m_id, res in zip(new_models, results):
                    if isinstance(res, bool):
                        multimodal_cache[m_id] = res
                        cache_updated = True
                    else:
                        logger.warning(f"Ошибка проверки модели {m_id}: {res}")
                
                if cache_updated:
                    save_multimodal_cache()
                    
            return models
    except Exception as e:
        logger.error(f"Не удалось получить список моделей из API: {e}")
        return [OLLAMA_MODEL]

async def get_cached_models() -> list[str]:
    now = time.time()
    if not api_models_cache["models"] or now > api_models_cache["expiry"]:
        models = await get_available_models_from_api()
        if models:
            api_models_cache["models"] = models
            api_models_cache["expiry"] = now + 120
    return api_models_cache["models"]

def get_model_pagination_keyboard(models: list[str], current_model: str, page: int = 0) -> types.InlineKeyboardMarkup:
    PAGE_SIZE = 8
    total_pages = (len(models) + PAGE_SIZE - 1) // PAGE_SIZE
    
    if page < 0:
        page = 0
    if page >= total_pages:
        page = total_pages - 1
        
    start_idx = page * PAGE_SIZE
    end_idx = start_idx + PAGE_SIZE
    page_models = models[start_idx:end_idx]
    
    buttons = []
    row = []
    for m_id in page_models:
        is_multimodal = is_model_multimodal(m_id)
        icon = "🖼️" if is_multimodal else "💬"
        is_active = current_model == m_id
        display_name = f"✅ {icon} {m_id}" if is_active else f"{icon} {m_id}"
        row.append(types.InlineKeyboardButton(text=display_name, callback_data=f"setmodel:{m_id}:{page}"))
        if len(row) == 2:
            buttons.append(row)
            row = []
    if row:
        buttons.append(row)
        
    nav_row = []
    if page > 0:
        nav_row.append(types.InlineKeyboardButton(text="⬅️ Назад", callback_data=f"modelpage:{page-1}"))
    else:
        nav_row.append(types.InlineKeyboardButton(text="❌", callback_data="modelpage_noop"))
        
    nav_row.append(types.InlineKeyboardButton(text=f"📄 {page+1} / {total_pages}", callback_data="modelpage_noop"))
    
    if page < total_pages - 1:
        nav_row.append(types.InlineKeyboardButton(text="Вперед ➡️", callback_data=f"modelpage:{page+1}"))
    else:
        nav_row.append(types.InlineKeyboardButton(text="❌", callback_data="modelpage_noop"))
        
    buttons.append(nav_row)
    close_row = [types.InlineKeyboardButton(text="❌ Закрыть меню", callback_data="close_model_menu")]
    buttons.append(close_row)
    return types.InlineKeyboardMarkup(inline_keyboard=buttons)

@dp.message(Command("model"))
async def model_cmd(message: types.Message):
    try:
        await message.delete()
    except Exception:
        pass
    user_id = message.from_user.id
    current_model = user_models.get(user_id, OLLAMA_MODEL)
    
    models = await get_cached_models()
    keyboard = get_model_pagination_keyboard(models, current_model, 0)
    
    await message.answer(
        text=(
            "🤖 <b>Выбор языковой модели ИИ</b>\n\n"
            f"Текущая активная модель: <code>{current_model}</code>\n\n"
            "<b>Обозначения моделей:</b>\n"
            "🖼️ — поддержка работы с изображениями (мультимодальные)\n"
            "💬 — только текстовые запросы\n\n"
            "Выберите модель из списка ниже для использования в ваших поисках:"
        ),
        parse_mode=ParseMode.HTML,
        reply_markup=keyboard
    )

@dp.inline_query()
async def inline_query_handler(inline_query: types.InlineQuery):
    query = inline_query.query.strip()
    if not query:
        return
    
    # Генерируем уникальный короткий ключ для кэширования запроса (Telegram ID limit = 64 chars)
    cache_key = uuid.uuid4().hex[:8]
    inline_queries_cache[cache_key] = {"query": query, "timestamp": time.time()}
    
    keyboard = types.InlineKeyboardMarkup(
        inline_keyboard=[[
            types.InlineKeyboardButton(text="⏳ Идет поиск...", callback_data="loading_dummy")
        ]]
    )
    
    results = [
        types.InlineQueryResultArticle(
            id=f"fast:{cache_key}",
            title=f"⚡ Быстрый поиск: {query[:35]}...",
            description="Полноценный сбор за 1 шаг (в пределах 2 минут)",
            input_message_content=types.InputTextMessageContent(
                message_text=(
                    f"🔍 <b>Запрос:</b> <code>{query}</code>\n"
                    f"📊 <code>[■■□□□□□□□□] 20%</code>\n\n"
                    f"🧠 <b>Статус:</b> Инициализация быстрого поиска..."
                ),
                parse_mode=ParseMode.HTML
            ),
            reply_markup=keyboard
        ),
        types.InlineQueryResultArticle(
            id=f"deep:{cache_key}",
            title=f"🔍 Глубокий поиск: {query[:35]}...",
            description="Глубокое многошаговое исследование (от 5 до 10 минут)",
            input_message_content=types.InputTextMessageContent(
                message_text=(
                    f"🔍 <b>Запрос:</b> <code>{query}</code>\n"
                    f"📊 <code>[■■□□□□□□□□] 20%</code>\n\n"
                    f"🧠 <b>Статус:</b> Инициализация глубокого поиска..."
                ),
                parse_mode=ParseMode.HTML
            ),
            reply_markup=keyboard
        )
    ]
    
    try:
        await inline_query.answer(results, cache_time=0, is_personal=True)
    except Exception as e:
        logger.error(f"Ошибка отправки инлайн ответа: {e}")

@dp.chosen_inline_result()
async def chosen_inline_result_handler(chosen_inline_result: types.ChosenInlineResult):
    result_id = chosen_inline_result.result_id
    inline_message_id = chosen_inline_result.inline_message_id
    user_id = chosen_inline_result.from_user.id
    if not inline_message_id:
        logger.warning("inline_message_id отсутствует в ChosenInlineResult")
        return
        
    mode = "deep"
    query = chosen_inline_result.query.strip()
    
    if result_id and ":" in result_id:
        try:
            mode_part, cache_key = result_id.split(":", 1)
            if mode_part in ["fast", "deep"]:
                mode = mode_part
                cached_data = inline_queries_cache.pop(cache_key, None)
                if cached_data:
                    query = cached_data["query"]
        except Exception as e:
            logger.error(f"Ошибка извлечения данных из инлайн-кэша: {e}")
    
    logger.info(f"Пользователь {user_id} выбрал режим {mode} для запроса: '{query}'. ID сообщения: {inline_message_id}")
    asyncio.create_task(process_inline_search(inline_message_id, query, user_id, mode))

async def process_inline_search(inline_message_id: str, query: str, user_id: int, mode: str = "deep"):
    task_key = f"inline:{inline_message_id}"
    active_search_tasks[task_key] = asyncio.current_task()
    try:
        async def update_status_func(text: str):
            try:
                keyboard = types.InlineKeyboardMarkup(
                    inline_keyboard=[[
                        types.InlineKeyboardButton(text="🛑 Остановить исследование", callback_data=f"stop:{task_key}")
                    ]]
                )
                await bot.edit_message_text(
                    text=text,
                    inline_message_id=inline_message_id,
                    parse_mode=ParseMode.HTML,
                    reply_markup=keyboard
                )
            except Exception as e:
                logger.error(f"Не удалось обновить инлайн статус: {e}")
                
        await run_multistep_search(
            query=query,
            user_id=user_id,
            status_updater=update_status_func,
            is_inline=True,
            inline_message_id=inline_message_id,
            mode=mode
        )
        
    except asyncio.CancelledError:
        logger.info(f"Инлайн-поиск {task_key} был отменен пользователем.")
        try:
            await bot.edit_message_text(
                text="🚫 <i>Исследование остановлено пользователем.</i>",
                inline_message_id=inline_message_id,
                parse_mode=ParseMode.HTML,
                reply_markup=None
            )
        except Exception as edit_err:
            logger.error(f"Не удалось отправить статус отмены в инлайн: {edit_err}")
        raise
    except Exception as e:
        logger.error(f"Ошибка при обработке инлайн-поиска: {e}")
        try:
            await bot.edit_message_text(
                text=f"❌ Произошла ошибка при обработке запроса: <i>{query}</i>\n\nДетали ошибки: {e}",
                inline_message_id=inline_message_id,
                parse_mode=ParseMode.HTML,
                reply_markup=None
            )
        except Exception as edit_err:
            logger.error(f"Не удалось обновить статус ошибки в инлайн сообщении: {edit_err}")
    finally:
        active_search_tasks.pop(task_key, None)

async def handle_photos_processing(photos: list, caption: str, chat_id: int, user_id: int):
    try:
        base64_images = []
        
        async def process_photo(photo_size):
            try:
                file = await bot.get_file(photo_size.file_id)
                file_io = BytesIO()
                await bot.download_file(file.file_path, file_io)
                b64_data = base64.b64encode(file_io.getvalue()).decode("utf-8")
                return b64_data
            except Exception as fe:
                logger.error(f"Не удалось скачать файл изображения: {fe}")
                return None
                
        tasks = [process_photo(photo) for photo in photos]
        results = await asyncio.gather(*tasks)
        base64_images = [r for r in results if r is not None]
        
        if not base64_images:
            raise ValueError("Не удалось загрузить ни одно изображение.")
            
        user_query = caption if caption else "Проанализируй присланные изображения."
        
        cache_key = uuid.uuid4().hex[:8]
        private_queries_cache[cache_key] = {
            "query": f"[Фото-запрос] Вопрос пользователя по фото: {user_query}",
            "base64_images": base64_images,
            "context_summary": None,  # Будет сгенерировано во время поиска
            "timestamp": time.time()
        }
        
        if not caption:
            # Картинка прислана без подписи - запрашиваем текст
            awaiting_photo_text[user_id] = cache_key
            
            keyboard = types.InlineKeyboardMarkup(
                inline_keyboard=[[
                    types.InlineKeyboardButton(text="Продолжить без текста", callback_data=f"continue_no_text:{cache_key}")
                ]]
            )
            
            await bot.send_message(
                chat_id=chat_id,
                text=(
                    f"✍️ <b>Будет ли какой-то текст или вопрос к этому фото?</b>\n"
                    f"Отправьте ваш вопрос следующим сообщением или нажмите кнопку ниже, чтобы продолжить без текста."
                ),
                parse_mode=ParseMode.HTML,
                reply_markup=keyboard
            )
        else:
            # Подпись есть - выводим стандартные кнопки
            keyboard = types.InlineKeyboardMarkup(
                inline_keyboard=[[
                    types.InlineKeyboardButton(text="⚡ Быстрый поиск", callback_data=f"fast:{cache_key}"),
                    types.InlineKeyboardButton(text="🔍 Глубокий поиск", callback_data=f"deep:{cache_key}")
                ]]
            )
            
            await bot.send_message(
                chat_id=chat_id,
                text=(
                    f"❓ <b>Запрос:</b> <code>{user_query[:50]}...</code>\n\n"
                    f"Выберите режим исследования:"
                ),
                parse_mode=ParseMode.HTML,
                reply_markup=keyboard
            )
    except Exception as e:
        logger.error(f"Ошибка при обработке фото: {e}")
        await bot.send_message(
            chat_id=chat_id,
            text=f"❌ Произошла ошибка при обработке фото: {e}",
            parse_mode=ParseMode.HTML
        )

async def process_media_group_after_delay(media_group_id: str, chat_id: int, user_id: int):
    await asyncio.sleep(1.5)
    data = media_groups_data.pop(media_group_id, None)
    if not data:
        return
    messages = data["messages"]
    messages.sort(key=lambda m: m.message_id)
    
    photos = []
    caption = ""
    for msg in messages:
        if msg.photo:
            photos.append(msg.photo[-1])
        if msg.caption and not caption:
            caption = msg.caption.strip()
            
    await handle_photos_processing(photos, caption, chat_id, user_id)

@dp.message(F.photo)
async def photo_message_handler(message: types.Message):
    logger.info(f"Получено фото: ID={message.message_id}, MediaGroup={message.media_group_id}")
    media_group_id = message.media_group_id
    if media_group_id:
        if media_group_id not in media_groups_data:
            media_groups_data[media_group_id] = {
                "messages": [],
                "timer": asyncio.create_task(process_media_group_after_delay(media_group_id, message.chat.id, message.from_user.id))
            }
        media_groups_data[media_group_id]["messages"].append(message)
    else:
        await handle_photos_processing([message.photo[-1]], message.caption.strip() if message.caption else "", message.chat.id, message.from_user.id)

@dp.message(F.document)
async def document_message_handler(message: types.Message):
    file_id = message.document.file_id
    file_name = message.document.file_name or "document"
    logger.info(f"Получен документ: ID={message.message_id}, Имя={file_name}")
    
    status_msg = await bot.send_message(
        chat_id=message.chat.id,
        text=f"📥 <b>Скачивание и анализ документа '{file_name}'...</b>",
        parse_mode=ParseMode.HTML
    )
    
    try:
        file_info = await bot.get_file(file_id)
        
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, file_name)
            await bot.download_file(file_info.file_path, temp_file_path)
            file_text = await extract_text_from_document(temp_file_path, file_name)
            
        if not file_text.strip():
            raise ValueError("Документ пуст или не содержит распознаваемого текста.")
            
        file_text = file_text[:20000]
        
        model_name = user_models.get(message.from_user.id, OLLAMA_MODEL)
        summary = await generate_document_summary(file_name, file_text, model_name=model_name)
        
        user_query = message.caption.strip() if message.caption else "Проанализируй содержимое прикрепленного файла."
        
        cache_key = uuid.uuid4().hex[:8]
        private_queries_cache[cache_key] = {
            "query": f"[Файл-запрос] Содержимое прикрепленного файла '{file_name}':\n\n{file_text}\n\nВопрос пользователя по файлу: {user_query}",
            "base64_images": None,
            "context_summary": summary,
            "timestamp": time.time()
        }
        
        preview = summary[:300] + "..." if len(summary) > 300 else summary
        
        keyboard = types.InlineKeyboardMarkup(
            inline_keyboard=[[
                types.InlineKeyboardButton(text="⚡ Быстрый поиск", callback_data=f"fast:{cache_key}"),
                types.InlineKeyboardButton(text="🔍 Глубокий поиск", callback_data=f"deep:{cache_key}")
            ]]
        )
        
        await status_msg.edit_text(
            text=(
                f"📄 <b>Файл '{file_name}' успешно обработан.</b>\n"
                f"📝 <b>Выжимка:</b> <i>{preview}</i>\n"
                f"❓ <b>Запрос:</b> <code>{user_query[:50]}...</code>\n\n"
                f"Выберите режим исследования:"
            ),
            parse_mode=ParseMode.HTML,
            reply_markup=keyboard
        )
        
    except Exception as e:
        logger.error(f"Ошибка при обработке документа: {e}")
        await status_msg.edit_text(
            text=f"❌ Произошла ошибка при обработке документа: {e}",
            parse_mode=ParseMode.HTML
        )

@dp.message()
async def private_message_handler(message: types.Message):
    print(f"DEBUG LOG: Получено сообщение! ID: {message.message_id}, Текст: {message.text}, Тип чата: {message.chat.type}", flush=True)
    
    if message.text and message.text.startswith("/"):
        if message.text.startswith("/start"):
            await start_cmd(message)
        elif message.text.startswith("/clear"):
            await clear_cmd(message)
        elif message.text.startswith("/model"):
            await model_cmd(message)
        elif message.text.startswith("/sessions"):
            await sessions_cmd(message)
        return
        
    user_id = message.from_user.id
    
    query = message.text.strip() if message.text else ""
    if not query:
        print("DEBUG LOG: Пустой запрос, игнорируем.", flush=True)
        return
        
    # Проверяем, ожидает ли бот текстовый запрос для ранее отправленного фото
    if user_id in awaiting_photo_text:
        cache_key = awaiting_photo_text.pop(user_id)
        if cache_key in private_queries_cache:
            private_queries_cache[cache_key]["query"] = f"[Фото-запрос] Вопрос пользователя по фото: {query}"
            
            keyboard = types.InlineKeyboardMarkup(
                inline_keyboard=[[
                    types.InlineKeyboardButton(text="⚡ Быстрый поиск", callback_data=f"fast:{cache_key}"),
                    types.InlineKeyboardButton(text="🔍 Глубокий поиск", callback_data=f"deep:{cache_key}")
                ]]
            )
            
            await message.answer(
                text=(
                    f"❓ <b>Запрос к фото установлен:</b> <code>{query[:50]}...</code>\n\n"
                    f"Выберите режим исследования:"
                ),
                parse_mode=ParseMode.HTML,
                reply_markup=keyboard
            )
            return
    
    if user_id in last_clear_message_ids:
        clear_msg_id = last_clear_message_ids.pop(user_id)
        try:
            await bot.delete_message(chat_id=message.chat.id, message_id=clear_msg_id)
            logger.info(f"Сообщение об очистке {clear_msg_id} удалено перед новым запросом.")
        except Exception as e:
            logger.warning(f"Не удалось удалить сообщение об очистке {clear_msg_id}: {e}")
            
    # Генерируем уникальный ключ и сохраняем запрос в кэш
    cache_key = uuid.uuid4().hex[:8]
    private_queries_cache[cache_key] = {"query": query, "timestamp": time.time()}
    
    keyboard = types.InlineKeyboardMarkup(
        inline_keyboard=[[
            types.InlineKeyboardButton(text="⚡ Быстрый поиск", callback_data=f"fast:{cache_key}"),
            types.InlineKeyboardButton(text="🔍 Глубокий поиск", callback_data=f"deep:{cache_key}")
        ]]
    )
    
    await message.answer(
        text=(
            f"❓ <b>Запрос:</b> <code>{query[:50]}...</code>\n\n"
            "Выберите режим исследования:"
        ),
        parse_mode=ParseMode.HTML,
        reply_markup=keyboard
    )

@dp.callback_query()
async def callback_query_handler(callback_query: types.CallbackQuery):
    user_id = callback_query.from_user.id
    data = callback_query.data or ""
    
    if data.startswith("continue_no_text:"):
        try:
            cache_key = data.split(":", 1)[1]
            awaiting_photo_text.pop(user_id, None)
            
            query_data = private_queries_cache.get(cache_key)
            if query_data:
                user_query = "Проанализируй присланные изображения."
                
                keyboard = types.InlineKeyboardMarkup(
                    inline_keyboard=[[
                        types.InlineKeyboardButton(text="⚡ Быстрый поиск", callback_data=f"fast:{cache_key}"),
                        types.InlineKeyboardButton(text="🔍 Глубокий поиск", callback_data=f"deep:{cache_key}")
                    ]]
                )
                
                combined_description = query_data.get("context_summary") or ""
                preview = combined_description[:300] + "..." if len(combined_description) > 300 else combined_description
                
                text_parts = []
                if preview:
                    text_parts.append(f"📝 <b>Описание:</b> <i>{preview}</i>")
                text_parts.append(f"❓ <b>Запрос:</b> <code>{user_query}</code>\n")
                text_parts.append("Выберите режим исследования:")
                
                try:
                    await callback_query.message.edit_text(
                        text="\n".join(text_parts),
                        parse_mode=ParseMode.HTML,
                        reply_markup=keyboard
                    )
                except Exception as e:
                    logger.error(f"Ошибка при обновлении сообщения после отмены текста: {e}")
        except Exception as outer_err:
            logger.error(f"Критическая ошибка в обработчике continue_no_text: {outer_err}")
        finally:
            await callback_query.answer()
        return
        
    # Обработка callback-запросов для архива сессий
    if data.startswith("sess_page:"):
        _, target_page = data.split(":")
        await show_sessions_page(
            chat_id=callback_query.message.chat.id,
            user_id=user_id,
            page=int(target_page),
            edit_message_id=callback_query.message.message_id
        )
        await callback_query.answer()
        return
        
    elif data == "sess_noop":
        await callback_query.answer()
        return
        
    elif data == "sess_close":
        try:
            await callback_query.message.delete()
        except Exception:
            pass
        await callback_query.answer()
        return
        
    elif data.startswith("sess_view:"):
        parts = data.split(":")
        s_id = int(parts[1])
        page = int(parts[2])
        
        session = get_session_by_id(s_id)
        if not session:
            await callback_query.answer("⚠️ Сессия не найдена.", show_alert=True)
            await show_sessions_page(callback_query.message.chat.id, user_id, page, callback_query.message.message_id)
            return
            
        _, _, _, title, created_at, _, _ = session
        try:
            dt = datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S")
            date_str = dt.strftime("%d.%m.%Y %H:%M")
        except Exception:
            date_str = created_at
            
        text = (
            f"📂 <b>Сессия:</b> «{title}»\n"
            f"📅 <b>Создана:</b> {date_str}\n\n"
            f"Выберите необходимое действие ниже:"
        )
        
        keyboard = types.InlineKeyboardMarkup(inline_keyboard=[
            [types.InlineKeyboardButton(text="🔄 Восстановить сессию в чат", callback_data=f"sess_act_restore:{s_id}:{page}")],
            [types.InlineKeyboardButton(text="📖 Показать финальный ответ", callback_data=f"sess_act_ans:{s_id}:{page}")],
            [types.InlineKeyboardButton(text="📄 Скачать конспект исследования", callback_data=f"sess_act_log:{s_id}:{page}")],
            [types.InlineKeyboardButton(text="🗑 Удалить из архива", callback_data=f"sess_act_del:{s_id}:{page}")],
            [types.InlineKeyboardButton(text="🔙 Назад к списку", callback_data=f"sess_back:{page}")]
        ])
        
        await bot.edit_message_text(
            text=text,
            chat_id=callback_query.message.chat.id,
            message_id=callback_query.message.message_id,
            parse_mode=ParseMode.HTML,
            reply_markup=keyboard
        )
        await callback_query.answer()
        return
        
    elif data.startswith("sess_back:"):
        _, page = data.split(":")
        await show_sessions_page(
            chat_id=callback_query.message.chat.id,
            user_id=user_id,
            page=int(page),
            edit_message_id=callback_query.message.message_id
        )
        await callback_query.answer()
        return
        
    elif data.startswith("sess_act_restore:"):
        parts = data.split(":")
        s_id = int(parts[1])
        page = int(parts[2])
        chat_id = callback_query.message.chat.id
        
        session = get_session_by_id(s_id)
        if not session:
            await callback_query.answer("⚠️ Сессия не найдена.", show_alert=True)
            return
            
        _, _, _, title, _, saved_log, saved_history_json = session
        
        # 1. Автосохранение текущей сессии перед загрузкой старой
        current_history = user_histories.get(user_id, [])
        current_log = read_research_log(chat_id)
        
        if current_history or current_log:
            try:
                user_model = user_models.get(user_id) or OLLAMA_MODEL
                current_title = await generate_session_title(current_history, model_name=user_model)
                await save_session_to_db(
                    user_id=user_id,
                    chat_id=chat_id,
                    title=current_title,
                    research_log=current_log,
                    chat_history=current_history
                )
            except Exception as e:
                logger.error(f"Ошибка автосохранения текущей сессии перед восстановлением: {e}")
                
        # 2. Восстанавливаем сохраненную сессию
        try:
            restored_history = json.loads(saved_history_json)
            user_histories[user_id] = restored_history
            active_session_ids[user_id] = s_id
            save_user_histories()
            
            file_path = f"data/research_log_{chat_id}.txt"
            os.makedirs("data", exist_ok=True)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(saved_log or "")
                
            try:
                await callback_query.message.delete()
            except Exception:
                pass
                
            last_answer = "Контекст пуст."
            for msg in reversed(restored_history):
                if msg.get("role") == "assistant":
                    last_answer = msg.get("content", "")
                    break
                    
            await bot.send_message(
                chat_id=chat_id,
                text=(
                    f"🔄 <b>Сессия «{title}» успешно восстановлена!</b>\n\n"
                    f"ИИ снова помнит контекст исследования. Вы можете продолжить общение с этого места.\n\n"
                    f"💬 <b>Последний ответ ИИ:</b>\n{last_answer[:600]}..." if len(last_answer) > 600 else f"💬 <b>Последний ответ ИИ:</b>\n{last_answer}"
                ),
                parse_mode=ParseMode.HTML
            )
            await callback_query.answer("Сессия восстановлена в чат.")
        except Exception as e:
            logger.error(f"Ошибка восстановления сессии ID {s_id}: {e}")
            await callback_query.answer("❌ Ошибка при восстановлении сессии.", show_alert=True)
        return

    elif data.startswith("sess_act_ans:"):
        parts = data.split(":")
        s_id = int(parts[1])
        page = int(parts[2])
        
        session = get_session_by_id(s_id)
        if not session:
            await callback_query.answer("⚠️ Сессия не найдена.", show_alert=True)
            return
            
        _, _, _, title, _, _, history_json = session
        
        final_answer = "В этой сессии нет ответов бота."
        try:
            history = json.loads(history_json)
            for msg in reversed(history):
                if msg.get("role") == "assistant":
                    final_answer = msg.get("content", "")
                    break
        except Exception:
            pass
            
        await bot.send_message(
            chat_id=callback_query.message.chat.id,
            text=(
                f"📖 <b>Финальный ответ сессии «{title}»:</b>\n\n"
                f"{final_answer}"
            ),
            parse_mode=ParseMode.HTML
        )
        await callback_query.answer("Ответ отправлен в чат.")
        return
        
    elif data.startswith("sess_act_log:"):
        parts = data.split(":")
        s_id = int(parts[1])
        page = int(parts[2])
        
        session = get_session_by_id(s_id)
        if not session:
            await callback_query.answer("⚠️ Сессия не найдена.", show_alert=True)
            return
            
        _, _, _, title, _, log_text, _ = session
        
        if not log_text or log_text.strip() == "":
            await callback_query.answer("⚠️ Лог исследования пуст.", show_alert=True)
            return
            
        try:
            safe_title = re.sub(r'[\\/*?:"<>|]', "", title)[:30].strip().replace(" ", "_")
            file_name = f"research_log_{safe_title}.txt"
            
            with tempfile.NamedTemporaryFile(mode="w+", suffix=".txt", delete=False, encoding="utf-8") as temp_file:
                temp_file.write(log_text)
                temp_file_path = temp_file.name
                
            input_file = types.FSInputFile(temp_file_path, filename=file_name)
            await bot.send_document(
                chat_id=callback_query.message.chat.id,
                document=input_file,
                caption=f"📄 Полный лог исследования для сессии «{title}»."
            )
            
            os.remove(temp_file_path)
            await callback_query.answer("Файл отправлен.")
        except Exception as e:
            logger.error(f"Ошибка при отправке файла лога исследования: {e}")
            await callback_query.answer("❌ Не удалось отправить файл.", show_alert=True)
        return
        
    elif data.startswith("sess_act_del:"):
        parts = data.split(":")
        s_id = int(parts[1])
        page = int(parts[2])
        
        delete_session_by_id(s_id)
        await callback_query.answer("🗑 Сессия удалена из архива.", show_alert=True)
        
        total_count = get_user_sessions_count(user_id)
        max_page = max(0, math.ceil(total_count / SESSIONS_PER_PAGE) - 1)
        target_page = min(page, max_page)
        
        await show_sessions_page(
            chat_id=callback_query.message.chat.id,
            user_id=user_id,
            page=target_page,
            edit_message_id=callback_query.message.message_id
        )
        return

    if callback_query.data == "loading_dummy":
        await callback_query.answer()
        return
        
    if callback_query.data == "close_model_menu":
        try:
            await callback_query.message.delete()
        except Exception:
            try:
                await callback_query.message.edit_reply_markup(reply_markup=None)
            except Exception:
                pass
        await callback_query.answer()
        return
        
    data_str = callback_query.data or ""
    if ":" not in data_str:
        await callback_query.answer()
        return
        
    prefix, payload_part = data_str.split(":", 1)

    if prefix == "stop":
        task_key = payload_part
        task = active_search_tasks.get(task_key)
        if task and not task.done():
            task.cancel()
            logger.info(f"[CANCEL] Задача {task_key} была успешно отменена.")
            await callback_query.answer("Исследование остановлено.", show_alert=False)
        else:
            await callback_query.answer("Исследование уже завершено или не найдено.", show_alert=False)
            
        # Удаляем сообщение для приватных чатов или обновляем в инлайне
        if task_key.startswith("private:"):
            try:
                _, chat_id, message_id = task_key.split(":")
                await bot.delete_message(chat_id=int(chat_id), message_id=int(message_id))
            except Exception as e:
                logger.error(f"Не удалось удалить сообщение после остановки в приватном чате: {e}")
        elif task_key.startswith("inline:"):
            try:
                _, inline_message_id = task_key.split(":", 1)
                await bot.edit_message_text(
                    text="🚫 <i>Исследование остановлено пользователем.</i>",
                    inline_message_id=inline_message_id,
                    parse_mode=ParseMode.HTML,
                    reply_markup=None
                )
            except Exception as e:
                logger.error(f"Не удалось отредактировать сообщение после остановки в инлайне: {e}")
        return
    
    if prefix == "setmodel":
        user_id = callback_query.from_user.id
        if ":" in payload_part:
            model_id, page_str = payload_part.rsplit(":", 1)
            try:
                page = int(page_str)
            except ValueError:
                page = 0
        else:
            model_id = payload_part
            page = 0
            
        models = await get_cached_models()
        if model_id in models:
            user_models[user_id] = model_id
            save_user_models()
            logger.info(f"Пользователь {user_id} переключил модель на '{model_id}' и настройки были сохранены.")
            
            try:
                await callback_query.message.delete()
            except Exception:
                try:
                    await callback_query.message.edit_text(
                        text=f"🤖 <b>Выбор языковой модели ИИ</b>\n\n✅ Модель изменена на: <code>{model_id}</code>",
                        parse_mode=ParseMode.HTML,
                        reply_markup=None
                    )
                except Exception:
                    pass
                
            await callback_query.answer(f"Модель изменена на {model_id}!", show_alert=False)
        else:
            await callback_query.answer("⚠️ Неизвестная модель ИИ.", show_alert=True)
        return
        
    elif prefix == "modelpage":
        user_id = callback_query.from_user.id
        current_model = user_models.get(user_id, OLLAMA_MODEL)
        try:
            page = int(payload_part)
        except ValueError:
            page = 0
            
        models = await get_cached_models()
        keyboard = get_model_pagination_keyboard(models, current_model, page)
        try:
            await callback_query.message.edit_reply_markup(reply_markup=keyboard)
        except Exception as e:
            logger.error(f"Не удалось переключить страницу моделей: {e}")
        await callback_query.answer()
        return
        
    if prefix not in ["fast", "deep"]:
        await callback_query.answer()
        return
        
    mode = prefix
    cache_key = payload_part
    
    query_data = private_queries_cache.pop(cache_key, None)
    if not query_data:
        await callback_query.answer("⚠️ Запрос устарел. Пожалуйста, отправьте его заново.", show_alert=True)
        try:
            await callback_query.message.delete()
        except Exception:
            pass
        return
        
    query = query_data["query"]
    base64_images = query_data.get("base64_images")
    context_summary = query_data.get("context_summary")
    status_text = "Инициализация быстрого поиска..." if mode == "fast" else "Инициализация глубокого поиска..."
    status_msg = callback_query.message
    
    try:
        await status_msg.edit_text(
            text=(
                f"🔍 <b>Запрос:</b> <code>{query[:50]}...</code>\n"
                f"📊 <code>[■■□□□□□□□□] 20%</code>\n\n"
                f"🧠 <b>Статус:</b> {status_text}"
            ),
            parse_mode=ParseMode.HTML,
            reply_markup=None
        )
    except Exception as e:
        logger.error(f"Не удалось переключить сообщение в статус поиска: {e}")
        await callback_query.answer("Произошла ошибка инициализации.", show_alert=True)
        return
        
    task_key = f"private:{status_msg.chat.id}:{status_msg.message_id}"
    async def run_search_task():
        active_search_tasks[task_key] = asyncio.current_task()
        try:
            async def update_status_func(text: str):
                try:
                    keyboard = types.InlineKeyboardMarkup(
                        inline_keyboard=[[
                            types.InlineKeyboardButton(text="🛑 Остановить исследование", callback_data=f"stop:{task_key}")
                        ]]
                    )
                    await status_msg.edit_text(text=text, parse_mode=ParseMode.HTML, reply_markup=keyboard)
                except Exception as e:
                    logger.error(f"Не удалось обновить статус: {e}")
                    
            async with ChatActionSender.typing(bot=bot, chat_id=status_msg.chat.id):
                await run_multistep_search(
                    query=query,
                    user_id=callback_query.from_user.id,
                    status_updater=update_status_func,
                    is_inline=False,
                    chat_id=status_msg.chat.id,
                    status_msg=status_msg,
                    mode=mode,
                    base64_images=base64_images,
                    context_summary=context_summary
                )
        except asyncio.CancelledError:
            logger.info(f"Исследование {task_key} было отменено пользователем.")
            try:
                await bot.delete_message(chat_id=status_msg.chat.id, message_id=status_msg.message_id)
            except Exception as delete_err:
                logger.error(f"Не удалось удалить сообщение после отмены: {delete_err}")
            raise
        except Exception as e:
            logger.error(f"Ошибка при обработке ЛС в асинхронном таске: {e}")
            try:
                await status_msg.answer(
                    text=f"❌ Произошла ошибка при обработке запроса: <i>{query}</i>\n\nДетали ошибки: {e}",
                    parse_mode=ParseMode.HTML
                )
            except Exception as edit_err:
                logger.error(f"Не удалось отправить сообщение об ошибке: {edit_err}")
        finally:
            active_search_tasks.pop(task_key, None)

    asyncio.create_task(run_search_task())
    await callback_query.answer()

async def clean_expired_caches():
    logger.info("[CACHE] Запущена фоновая задача очистки устаревших кэшей.")
    while True:
        await asyncio.sleep(600)  # Каждые 10 минут
        try:
            now = time.time()
            expiry_time = 3600  # 1 час
            
            # Очищаем inline_queries_cache
            expired_inline = [
                k for k, v in inline_queries_cache.items()
                if isinstance(v, dict) and now - v.get("timestamp", 0) > expiry_time
            ]
            for k in expired_inline:
                inline_queries_cache.pop(k, None)
                
            # Очищаем private_queries_cache
            expired_private = [
                k for k, v in private_queries_cache.items()
                if isinstance(v, dict) and now - v.get("timestamp", 0) > expiry_time
            ]
            for k in expired_private:
                private_queries_cache.pop(k, None)
                
            if expired_inline or expired_private:
                logger.info(f"[CACHE] Очищено устаревших записей: inline={len(expired_inline)}, private={len(expired_private)}")
        except Exception as e:
            logger.error(f"[CACHE] Ошибка при очистке кэшей: {e}")

async def main():
    init_db()
    logger.info("Регистрация команд в Telegram...")
    await bot.set_my_commands([
        types.BotCommand(command="start", description="Запустить бота"),
        types.BotCommand(command="model", description="Выбрать модель ИИ"),
        types.BotCommand(command="clear", description="Очистить историю сообщений"),
        types.BotCommand(command="sessions", description="Архив прошлых поисковых сессий")
    ])
    
    # Прогреваем кэш доступных моделей при старте и проверяем их мультимодальность
    try:
        logger.info("Прогрев кэша моделей и ленивая проверка мультимодальности...")
        await get_cached_models()
    except Exception as e:
        logger.error(f"Ошибка при прогреве кэша моделей: {e}")
        
    logger.info("Запуск фоновых задач...")
    update_manager.start()
    asyncio.create_task(clean_expired_caches())
    logger.info("Запуск Telegram-бота...")
    await dp.start_polling(bot, allowed_updates=["message", "inline_query", "chosen_inline_result", "callback_query"])

if __name__ == "__main__":
    asyncio.run(main())
