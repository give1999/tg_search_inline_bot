import os
import logging
import json
import re
from aiogram import Bot, Dispatcher, types, F, BaseMiddleware
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode, ChatType
from aiogram.filters import Command
from aiogram.utils.chat_action import ChatActionSender
from aiogram.types import TelegramObject
import httpx
import asyncio
from bs4 import BeautifulSoup
from datetime import datetime
import time
import uuid
import base64
from io import BytesIO
from contextvars import ContextVar
import tempfile

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import xlrd
except ImportError:
    xlrd = None

current_model_var = ContextVar('current_model_var', default=None)
current_status_updater_var = ContextVar('current_status_updater_var', default=None)
current_query_var = ContextVar('current_query_var', default="")
current_progress_percent_var = ContextVar('current_progress_percent_var', default=0)
current_search_queries_var = ContextVar('current_search_queries_var', default=None)
current_attempt_var = ContextVar('current_attempt_var', default=1)
current_thoughts_var = ContextVar('current_thoughts_var', default=None)
current_status_text_var = ContextVar('current_status_text_var', default="")


# Кэш мультимодальных моделей
multimodal_cache = {}
MULTIMODAL_CACHE_FILE = "/app/data/multimodal_cache.json"

def load_multimodal_cache():
    global multimodal_cache
    path = MULTIMODAL_CACHE_FILE
    if not os.path.exists(path):
        path = "./multimodal_cache.json"
    
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                multimodal_cache = json.load(f)
                logger.info(f"Успешно загружен кэш мультимодальных моделей: {len(multimodal_cache)} записей.")
        except Exception as e:
            logger.error(f"Ошибка при загрузке multimodal_cache из {path}: {e}")
            multimodal_cache = {}
    else:
        multimodal_cache = {}

def save_multimodal_cache():
    path = MULTIMODAL_CACHE_FILE
    dir_path = os.path.dirname(path)
    if dir_path and not os.path.exists(dir_path):
        path = "./multimodal_cache.json"
    
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(multimodal_cache, f, ensure_ascii=False, indent=4)
        logger.info(f"Кэш мультимодальных моделей сохранен в {path}")
    except Exception as e:
        logger.error(f"Ошибка при сохранении multimodal_cache в {path}: {e}")

multimodal_check_semaphore = asyncio.Semaphore(2)

async def check_model_multimodal(model_name: str) -> bool:
    sample_jpeg_b64 = (
        "/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAgGBgcGBQgHBwcJCQgKDBQNDAsLDBkSEw8UHRofH"
        "h0aHBwgJC4nICIsIxwcKDcpLDAxNDQ0Hyc5PTgyPC4zNDL/2wBDAQkJCQwLDBgNDRgyIRwhMj"
        "IyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjL/wAARCABA"
        "AEADASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgED"
        "AwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcY"
        "GRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJ"
        "ipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo"
        "6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgEC"
        "BAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl"
        "8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaH"
        "iImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn"
        "6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwDi6KKK+ZP3EKKKKACiiigAooooAKKKKACiiigA"
        "ooooA//Z"
    )
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/chat/completions"
    headers = {"Content-Type": "application/json"}
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
        
    payload = {
        "model": model_name,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Is this image red?"},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{sample_jpeg_b64}"}}
                ]
            }
        ],
        "temperature": 0.1
    }
    
    async with multimodal_check_semaphore:
        retry_delay = 2.0
        for attempt in range(1, 5):
            try:
                async with httpx.AsyncClient(timeout=15.0) as client:
                    logger.info(f"[MULTIMODAL-CHECK] Проверка модели {model_name} (Попытка {attempt}/4)...")
                    resp = await client.post(url, json=payload, headers=headers)
                    if resp.status_code == 200:
                        logger.info(f"[MULTIMODAL-CHECK] Модель {model_name} поддерживает изображения (HTTP 200).")
                        return True
                    elif resp.status_code == 429:
                        logger.warning(f"[MULTIMODAL-CHECK] Модель {model_name} вернула 429. Повтор через {retry_delay} сек...")
                        await asyncio.sleep(retry_delay)
                        retry_delay *= 2.0
                        continue
                    else:
                        logger.info(f"[MULTIMODAL-CHECK] Модель {model_name} НЕ поддерживает изображения (HTTP {resp.status_code}).")
                        return False
            except Exception as e:
                logger.warning(f"[MULTIMODAL-CHECK] Сетевая ошибка при проверке {model_name} (Попытка {attempt}/4): {e}")
                if attempt == 4:
                    raise e
                await asyncio.sleep(retry_delay)
                retry_delay *= 2.0
        return False

def is_model_multimodal(model_name: str) -> bool:
    return multimodal_cache.get(model_name, False)

# Для обратной совместимости
MULTIMODAL_MODELS = []
PREFERRED_BACKUP_MODELS = []
PREFERRED_VISION_MODELS = []

def clean_multimodal_messages(messages: list) -> list:
    cleaned = []
    for msg in messages:
        role = msg.get("role")
        content = msg.get("content")
        if isinstance(content, list):
            text_parts = []
            for part in content:
                if isinstance(part, dict):
                    if part.get("type") == "text":
                        text_parts.append(part.get("text", ""))
                else:
                    text_parts.append(str(part))
            cleaned.append({"role": role, "content": " ".join(text_parts).strip()})
        else:
            cleaned.append({"role": role, "content": content})
    return cleaned


def get_model_for_attempt(start_model: str, attempt: int, all_api_models: list[str]) -> str:
    if attempt == 1:
        return start_model or OLLAMA_MODEL
        
    fallback_chain = []
    for am in all_api_models:
        if am != start_model:
            fallback_chain.append(am)
            
    if not fallback_chain:
        return start_model or OLLAMA_MODEL
        
    idx = attempt - 2
    if idx < len(fallback_chain):
        return fallback_chain[idx]
    return fallback_chain[-1]


media_groups_data = {}


# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Инициализация переменных окружения
BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3")
OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY", "")
SEARXNG_URL = os.getenv("SEARXNG_URL", "http://127.0.0.1:8080")

if not BOT_TOKEN:
    raise ValueError("TELEGRAM_BOT_TOKEN не задан в переменных окружения!")

bot = Bot(token=BOT_TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))
dp = Dispatcher()

# Глобальный словарь для хранения неограниченной истории диалогов
user_histories = {}
# Глобальный словарь для отслеживания сообщений очистки чата
last_clear_message_ids = {}

# Кэши для обхода лимита в 64 символа у Telegram API
inline_queries_cache = {}   # {cache_key: query}
private_queries_cache = {}  # {cache_key: {"query": query}}

# Глобальный словарь моделей пользователей {user_id: model_name}
user_models = {}
USER_MODELS_FILE = "/app/data/user_models.json"

# Семафор для предотвращения параллельного ранжирования сайтов и перегрузки ИИ-шлюза
ranking_semaphore = asyncio.Semaphore(1)

def load_user_models():
    global user_models
    path = USER_MODELS_FILE
    if not os.path.exists(path):
        # Попробуем локальный путь, если мы вне докера
        path = "./user_models.json"
    
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                # Приводим ключи к int для совместимости в runtime
                user_models = {int(k): v for k, v in data.items() if k.isdigit()}
                logger.info(f"Успешно загружены настройки моделей для {len(user_models)} пользователей.")
        except Exception as e:
            logger.error(f"Ошибка при загрузке user_models из {path}: {e}")
            user_models = {}
    else:
        user_models = {}

def save_user_models():
    path = USER_MODELS_FILE
    # Если директория /app/data не существует, сохраняем в текущую
    dir_path = os.path.dirname(path)
    if dir_path and not os.path.exists(dir_path):
        path = "./user_models.json"
    
    try:
        # Для сохранения приводим ключи к строкам (JSON поддерживает только строковые ключи)
        data = {str(k): v for k, v in user_models.items()}
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        logger.info(f"Настройки моделей пользователей сохранены в {path}")
    except Exception as e:
        logger.error(f"Ошибка при сохранении user_models в {path}: {e}")

# Инициализируем при запуске
load_user_models()
load_multimodal_cache()

class LoggingMiddleware(BaseMiddleware):
    async def __call__(self, handler, event: TelegramObject, data):
        print(f"RAW UPDATE PRINT: {event.model_dump_json(exclude_none=True)}", flush=True)
        return await handler(event, data)

dp.message.outer_middleware(LoggingMiddleware())
dp.inline_query.outer_middleware(LoggingMiddleware())
dp.chosen_inline_result.outer_middleware(LoggingMiddleware())

def get_current_datetime_str() -> str:
    return datetime.now().strftime("%d.%m.%Y %H:%M:%S")

def clean_model_answer(answer: str) -> str:
    text = answer.strip()
    prefixes = [
        "ответ на запрос:",
        "ответ на вопрос:",
        "ответ пользователя:",
        "ответ:",
        "запрос:",
        "вывод:"
    ]
    changed = True
    while changed:
        changed = False
        lower_text = text.lower()
        for prefix in prefixes:
            if lower_text.startswith(prefix):
                text = text[len(prefix):].strip()
                changed = True
                break
    if text.startswith('"') and text.endswith('"'):
        text = text[1:-1].strip()
    elif text.startswith("'") and text.endswith("'"):
        text = text[1:-1].strip()
    return text

def clean_markdown_from_thoughts(text: str) -> str:
    if not text:
        return ""
    # Удаляем заголовки (символы # в начале строк)
    t = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    # Удаляем ссылки вида [text](url) -> оставляем только text
    t = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', t)
    # Удаляем жирный шрифт, курсив, зачеркивание, моноширинный, спойлеры: **, *, __, _, `, ~, ||
    t = re.sub(r'\*\*|__|\*|_|`|~|\|\|', '', t)
    # Удаляем маркеры списков в начале строк (- , * , + , •)
    t = re.sub(r'^\s*[-*+•]\s+', '', t, flags=re.MULTILINE)
    # Схлопываем лишние переводы строк
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()

def auto_close_html_tags(html: str) -> str:
    tag_re = re.compile(r'</?([a-zA-Z0-9_-]+)(?:\s+[^>]*)?>')
    open_tags = []
    
    for match in tag_re.finditer(html):
        full_tag = match.group(0)
        tag_name = match.group(1).lower()
        
        if full_tag.startswith('</'):
            if open_tags and open_tags[-1] == tag_name:
                open_tags.pop()
        else:
            open_tags.append(tag_name)
            
    closed_html = html
    for tag_name in reversed(open_tags):
        closed_html += f"</{tag_name}>"
        
    return closed_html

def safe_html_cleaner(text: str) -> str:
    text = re.sub(r'&(?!([a-zA-Z]+|#[0-9]+);)', '&amp;', text)
    allowed_tags = ['b', 'strong', 'i', 'em', 'u', 'ins', 's', 'strike', 'del', 'code', 'pre', 'blockquote']
    tag_re = re.compile(r'</?([a-zA-Z0-9_-]+)(?:\s+[^>]*)?>')
    
    pos = 0
    result = []
    for match in tag_re.finditer(text):
        chunk = text[pos:match.start()]
        chunk = chunk.replace('<', '&lt;').replace('>', '&gt;')
        result.append(chunk)
        
        tag_name = match.group(1).lower()
        full_tag = match.group(0)
        
        if tag_name in allowed_tags or (tag_name == 'blockquote' and 'expandable' in full_tag):
            result.append(full_tag)
        else:
            result.append(full_tag.replace('<', '&lt;').replace('>', '&gt;'))
            
        pos = match.end()
        
    chunk = text[pos:]
    chunk = chunk.replace('<', '&lt;').replace('>', '&gt;')
    result.append(chunk)
    
    return "".join(result)

async def describe_image(base64_image: str, model_name: str = None) -> str:
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/chat/completions"
    headers = {"Content-Type": "application/json"}
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
        
    prompt_text = (
        "Ты — помощник-визионер. Твоя задача — максимально подробно и детально описать присланное изображение на русском языке. "
        "Опиши людей, предметы, их действия, цвета, локацию, распознай и дословно выпиши весь текст на картинке, если он есть. "
        "Пиши структурировано и понятно, так как твой ответ будет использоваться для поиска информации по этой картинке в интернете."
    )
    
    start_model = model_name or OLLAMA_MODEL
    
    # Формируем каскадную цепочку мультимодальных моделей
    all_api_models = await get_cached_models()
    vision_models_on_server = [m for m in all_api_models if is_model_multimodal(m)]
    
    fallback_chain = []
    if start_model:
        fallback_chain.append(start_model)
    for vm in vision_models_on_server:
        if vm not in fallback_chain:
            fallback_chain.append(vm)
            
    if not fallback_chain:
        fallback_chain = [OLLAMA_MODEL]
            
    logger.info(f"[VISION] Запуск распознавания картинки. Стартовая модель: {start_model}, Цепочка: {fallback_chain}...")
    
    for active_model in fallback_chain:
        payload = {
            "model": active_model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt_text},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }
            ],
            "temperature": 0.2
        }
        
        # Пробуем до 2 раз на каждую модель
        max_attempts_per_model = 2
        for attempt in range(1, max_attempts_per_model + 1):
            if active_model != start_model or attempt > 1:
                await update_status(warning_text=f"Переключаемся на резервный визионер {active_model} (Попытка {attempt}/{max_attempts_per_model})...")
                
            async with httpx.AsyncClient(timeout=20.0) as client:
                try:
                    logger.info(f"[VISION] Запрос к {active_model}. Попытка {attempt}/{max_attempts_per_model}. POST на {url}...")
                    response = await client.post(url, json=payload, headers=headers)
                    logger.info(f"[VISION] Ответ {active_model}: {response.status_code}")
                    response.raise_for_status()
                    result = response.json()
                    content = result["choices"][0]["message"].get("content", "") or ""
                    logger.info(f"[VISION] Успешно получено описание от {active_model}. Длина: {len(content)} симв.")
                    return content.strip()
                except Exception as e:
                    logger.warning(f"[VISION] [СБОЙ] Модель {active_model} на попытке {attempt} вернула ошибку: {repr(e)}")
                    
    # Если все модели сбоят
    raise ValueError("Ни одна из доступных мультимодальных моделей (визионеров) не смогла распознать изображение.")


async def generate_document_summary(filename: str, file_text: str, model_name: str = None) -> str:
    prompt = (
        f"Ты — эксперт по анализу документов. Сделай краткую, но информативную выжимку прикрепленного файла '{filename}'.\n"
        f"Укажи основную тему документа, ключевые разделы, сущности, даты и суть содержимого.\n"
        f"Если это таблица Excel, опиши структуру листов, заголовки ключевых колонок и диапазон данных.\n\n"
        f"Текст документа:\n{file_text[:15000]}\n\n"
        f"Твой ответ должен быть на русском языке, лаконичным, структурированным и содержать не более 1500 символов."
    )
    logger.info(f"[SUMMARY] Генерация выжимки для файла: {filename}...")
    try:
        summary = await ask_ollama(prompt, model_name=model_name)
        logger.info(f"[SUMMARY] Выжимка успешно сгенерирована. Длина: {len(summary)} симв.")
        return summary.strip()
    except Exception as e:
        logger.error(f"[SUMMARY] [ОШИБКА] Не удалось сгенерировать выжимку: {e}")
        return f"Файл {filename}. Содержит текст: {file_text[:500]}..."

async def extract_text_from_document(file_path: str, file_name: str) -> str:
    ext = file_name.lower().split('.')[-1]
    
    if ext == 'pdf':
        if not pypdf:
            raise ImportError("Модуль 'pypdf' не установлен на сервере.")
        reader = pypdf.PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
        
    elif ext == 'docx':
        if not docx:
            raise ImportError("Модуль 'python-docx' не установлен на сервере.")
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
        
    elif ext == 'xlsx':
        if not openpyxl:
            raise ImportError("Модуль 'openpyxl' не установлен на сервере.")
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text_parts.append(f"--- Лист: {sheet} ---")
            for row in ws.iter_rows(values_only=True):
                row_str = " | ".join([str(val) for val in row if val is not None])
                if row_str.strip():
                    text_parts.append(row_str)
        return "\n".join(text_parts)
        
    elif ext == 'xls':
        if not xlrd:
            raise ImportError("Модуль 'xlrd' не установлен на сервере.")
        wb = xlrd.open_workbook(file_path)
        text_parts = []
        for sheet_idx in range(wb.nsheets):
            ws = wb.sheet_by_index(sheet_idx)
            text_parts.append(f"--- Лист: {ws.name} ---")
            for row_idx in range(ws.nrows):
                row = ws.row_values(row_idx)
                row_str = " | ".join([str(val) for val in row if val != ''])
                if row_str.strip():
                    text_parts.append(row_str)
        return "\n".join(text_parts)
        
    else:
        # Default to text decoding
        with open(file_path, 'rb') as f:
            raw_bytes = f.read()
        for encoding in ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']:
            try:
                return raw_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Не удалось определить кодировку текстового файла.")

def make_status_html(query: str, status_text: str, progress_percent: int, search_queries: list = None, attempt: int = 1, warning_text: str = None, thoughts: str = None, model_name: str = None) -> str:
    # Update ContextVars for tracking status
    if query is not None:
        current_query_var.set(query)
    if status_text is not None:
        current_status_text_var.set(status_text)
    if progress_percent is not None:
        current_progress_percent_var.set(progress_percent)
    if search_queries is not None:
        current_search_queries_var.set(search_queries)
    if attempt is not None:
        current_attempt_var.set(attempt)
    if thoughts is not None:
        current_thoughts_var.set(thoughts)

    if not model_name:
        model_name = current_model_var.get()
    display_query = query[:50] + "..." if len(query) > 50 else query
    
    filled_blocks = progress_percent // 10
    empty_blocks = 10 - filled_blocks
    progress_bar = f"📊 <code>[{'■' * filled_blocks}{'□' * empty_blocks}] {progress_percent}%</code>"
    
    attempt_str = f" (Попытка {attempt}/3)" if attempt > 1 else ""
    
    if progress_percent <= 30:
        circle = "🔴"
    elif progress_percent <= 50:
        circle = "🟠"
    elif progress_percent <= 75:
        circle = "🟡"
    elif progress_percent <= 90:
        circle = "🔵"
    else:
        circle = "🟢"
        
    model_str = f"🤖 <b>Модель:</b> <code>{model_name}</code>\n" if model_name else ""
        
    html = (
        f"🔍 <b>Запрос:</b> <code>{display_query}</code>\n"
        f"{progress_bar}\n"
        f"{model_str}\n"
        f"{circle} <b>Статус:</b> {status_text}{attempt_str}\n"
    )
    
    if thoughts:
        clean_thoughts = clean_markdown_from_thoughts(thoughts)
        if clean_thoughts:
            safe_thoughts = safe_html_cleaner(clean_thoughts)
            html += f"\n<blockquote expandable>{safe_thoughts}</blockquote>\n"
        
    if warning_text:
        html += f"\n⚠️ {warning_text}\n"
        
    if search_queries:
        html += "\n<i>Ищу по запросам:</i>"
        for item in search_queries:
            if isinstance(item, dict):
                q_text = item.get("query", "")
                html += f"\n🔍 <code>{q_text}</code>"
                for site in item.get("sites", []):
                    html += f"\n   └ 📄 <i>{site}</i>"
            else:
                html += f"\n🔍 <code>{item}</code>"
            
    return auto_close_html_tags(html)

def get_clean_fallback_query(query: str) -> str:
    """
    Очищает запрос от префиксов [Фото-запрос] или [Файл-запрос] с содержимым файла,
    возвращая только чистый текст вопроса пользователя.
    """
    if not query:
        return ""
        
    query = query.strip()
    
    # 1. Если это Фото-запрос
    if query.startswith("[Фото-запрос]"):
        prefix = "[Фото-запрос] Вопрос пользователя по фото:"
        if query.startswith(prefix):
            return query[len(prefix):].strip()
        return query.replace("[Фото-запрос]", "").strip()
        
    # 2. Если это Файл-запрос
    if query.startswith("[Файл-запрос]"):
        marker = "Вопрос пользователя по файлу:"
        idx = query.rfind(marker)
        if idx != -1:
            return query[idx + len(marker):].strip()
            
        cleaned = query.replace("[Файл-запрос]", "").strip()
        lines = cleaned.split("\n")
        if lines:
            return lines[0][:80].strip()
            
        return cleaned[:80].strip()
        
    return query

async def update_status(
    status_text: str = None,
    progress_percent: int = None,
    warning_text: str = None,
    search_queries: list = None,
    attempt: int = None,
    thoughts: str = None
):
    updater = current_status_updater_var.get()
    if not updater:
        return
        
    if status_text is not None:
        current_status_text_var.set(status_text)
    if progress_percent is not None:
        current_progress_percent_var.set(progress_percent)
    if search_queries is not None:
        current_search_queries_var.set(search_queries)
    if attempt is not None:
        current_attempt_var.set(attempt)
    if thoughts is not None:
        current_thoughts_var.set(thoughts)
        
    html = make_status_html(
        query=current_query_var.get(),
        status_text=current_status_text_var.get(),
        progress_percent=current_progress_percent_var.get(),
        search_queries=current_search_queries_var.get(),
        attempt=current_attempt_var.get(),
        warning_text=warning_text,
        thoughts=current_thoughts_var.get()
    )
    try:
        await updater(html)
    except Exception as e:
        logger.error(f"Ошибка при обновлении статуса: {e}")

async def fetch_page_content(url: str) -> str:
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    logger.info(f"[FETCH] Начинаем скачивание страницы: {url}")
    async with httpx.AsyncClient(timeout=15.0, follow_redirects=True) as client:
        try:
            response = await client.get(url, headers=headers)
            logger.info(f"[FETCH] Сервер вернул код {response.status_code} для {url}")
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, "html.parser")
                for element in soup(["script", "style", "header", "footer", "nav", "aside", "iframe", "noscript"]):
                    element.decompose()
                text = soup.get_text(separator=" ")
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                clean_text = "\n".join(lines)
                truncated_text = clean_text[:20000]
                logger.info(f"[FETCH] Успешно скачано {url}. Символов всего: {len(clean_text)}, сохранено: {len(truncated_text)}")
                return truncated_text
            else:
                logger.warning(f"[FETCH] Неверный статус-код {response.status_code} для {url}")
                return ""
        except Exception as e:
            logger.error(f"[FETCH] [ОШИБКА] Исключение при скачивании {url}: {repr(e)}")
            return ""

async def search_searxng_raw(query: str) -> list:
    url = f"{SEARXNG_URL.rstrip('/')}/search"
    params = {
        "q": query,
        "format": "json",
        "pageno": 1
    }
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    
    logger.info(f"[SEARXNG] Запрос к SearxNG: {url} с q='{query}'")
    for attempt in range(1, 3):
        async with httpx.AsyncClient(timeout=15.0) as client:
            try:
                await asyncio.sleep(1.0)
                logger.info(f"[SEARXNG] Попытка {attempt}/2. Отправка GET...")
                response = await client.get(url, params=params, headers=headers)
                
                if response.status_code == 429:
                    logger.warning(f"[SEARXNG] [HTTP 429] Лимит запросов. Попытка {attempt}/2. Ждем 3.0 секунды...")
                    await update_status(warning_text="Поисковая система перегружена. Повторяем попытку через 3 секунды...")
                    await asyncio.sleep(3.0)
                    continue
                    
                response.raise_for_status()
                data = response.json()
                results = data.get("results", [])
                logger.info(f"[SEARXNG] Успешно получен ответ. Всего результатов: {len(results)}")
                return results
            except Exception as e:
                logger.error(f"[SEARXNG] [ОШИБКА] Исключение при запросе к SearxNG по '{query}' (попытка {attempt}): {repr(e)}")
                if attempt < 2:
                    await update_status(warning_text="Задержка ответа от поисковой системы. Повторяем попытку...")
                    await asyncio.sleep(2.0)
                    
    return []


async def ask_ollama(prompt_or_messages, system_prompt: str = None, model_name: str = None) -> str:
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/chat/completions"
    headers = {
        "Content-Type": "application/json"
    }
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
    
    if isinstance(prompt_or_messages, list):
        messages = prompt_or_messages
    else:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt_or_messages}
        ]
        
    start_model = model_name or OLLAMA_MODEL
    
    # Запрашиваем список всех доступных в API моделей
    all_api_models = await get_cached_models()
    
    # Строим каскадную цепочку моделей (ограничиваем, чтобы избежать бесконечного перебора)
    fallback_chain = []
    if start_model:
        fallback_chain.append(start_model)
    
    extra_added = 0
    for am in all_api_models:
        if am not in fallback_chain:
            fallback_chain.append(am)
            extra_added += 1
            if extra_added >= 4:
                break
            
    if not fallback_chain:
        fallback_chain = [start_model or OLLAMA_MODEL]
        
    logger.info(f"[OLLAMA] Запуск ask_ollama. Стартовая модель: {start_model}, Цепочка: {fallback_chain}, Сообщений: {len(messages)}")
    
    for active_model in fallback_chain:
        # Если модель не мультимодальная, очищаем сообщения от картинок
        current_messages = messages
        if not is_model_multimodal(active_model):
            current_messages = clean_multimodal_messages(messages)
            
        payload = {
            "model": active_model,
            "messages": current_messages,
            "temperature": 0.3
        }
        
        # Для стартовой модели пробуем 2 раза, для резервных — 1 раз
        max_attempts_per_model = 2 if active_model == start_model else 1
        for attempt in range(1, max_attempts_per_model + 1):
            # Если это не первая модель в цепочке или не первая попытка, выводим статус
            if active_model != start_model or attempt > 1:
                await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} (Попытка {attempt}/{max_attempts_per_model})...")
                
            async with httpx.AsyncClient(timeout=15.0) as client:
                try:
                    logger.info(f"[OLLAMA] Запрос к {active_model}. Попытка {attempt}/{max_attempts_per_model}. POST на {url}...")
                    
                    response = None
                    retry_delay = 2.0
                    for r_attempt in range(4):  # 1 основная + 3 ретрая при 429
                        if r_attempt > 0:
                            logger.warning(f"[OLLAMA] Повторный запрос к {active_model} после 429 через {retry_delay} сек...")
                            await update_status(warning_text=f"ИИ-сервер перегружен (429). Повтор через {retry_delay} сек...")
                            await asyncio.sleep(retry_delay)
                            retry_delay *= 2.0
                            
                        response = await client.post(url, json=payload, headers=headers)
                        logger.info(f"[OLLAMA] Ответ {active_model} (повтор {r_attempt}): {response.status_code}")
                        if response.status_code != 429:
                            break
                            
                    response.raise_for_status()
                    result = response.json()
                    message_obj = result["choices"][0]["message"]
                    content = message_obj.get("content", "") or ""
                    reasoning = message_obj.get("reasoning", "") or message_obj.get("reasoning_content", "") or ""
                    if reasoning.strip():
                        content = f"<think>\n{reasoning}\n</think>\n{content}"
                    logger.info(f"[OLLAMA] Успешная генерация на модели {active_model}. Длина: {len(content)} симв.")
                    
                    # Обновляем глобальный контекст текущей модели
                    current_model_var.set(active_model)
                    return content
                except Exception as e:
                    logger.warning(f"[OLLAMA] [СБОЙ] Модель {active_model} на попытке {attempt} вернула ошибку: {repr(e)}")
                    # Переходим к следующей попытке или следующей модели
                    
    # Если мы дошли сюда, значит ни одна модель не сработала
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла ответить на запрос (сетевые таймауты или ошибки сервера).")

async def ask_ollama_stream(messages: list, model_name: str = None):
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/chat/completions"
    headers = {
        "Content-Type": "application/json"
    }
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
        
    start_model = model_name or OLLAMA_MODEL
    
    # Запрашиваем список всех доступных в API моделей
    all_api_models = await get_cached_models()
    
    # Строим каскадную цепочку моделей (ограничиваем, чтобы избежать бесконечного перебора)
    fallback_chain = []
    if start_model:
        fallback_chain.append(start_model)
    
    extra_added = 0
    for am in all_api_models:
        if am not in fallback_chain:
            fallback_chain.append(am)
            extra_added += 1
            if extra_added >= 4:
                break
            
    if not fallback_chain:
        fallback_chain = [start_model or OLLAMA_MODEL]
        
    logger.info(f"[OLLAMA-STREAM] Запуск ask_ollama_stream. Стартовая модель: {start_model}, Цепочка: {fallback_chain}, Сообщений в истории: {len(messages)}")
    
    yielded_any = False
    
    for active_model in fallback_chain:
        if yielded_any:
            break
            
        # Очищаем от картинок, если резервная модель текстовая
        current_messages = messages
        if not is_model_multimodal(active_model):
            current_messages = clean_multimodal_messages(messages)
            
        payload = {
            "model": active_model,
            "messages": current_messages,
            "temperature": 0.3,
            "stream": True
        }
        
        # Для стартовой модели пробуем 2 раза, для резервных — 1 раз
        max_attempts_per_model = 2 if active_model == start_model else 1
        for attempt in range(1, max_attempts_per_model + 1):
            if yielded_any:
                break
                
            if active_model != start_model or attempt > 1:
                await update_status(warning_text=f"Переключаемся на резервный ИИ-стрим {active_model} (Попытка {attempt}/{max_attempts_per_model})...")
                
            async with httpx.AsyncClient(timeout=45.0) as client:
                try:
                    logger.info(f"[OLLAMA-STREAM] Запрос к {active_model}. Попытка {attempt}/{max_attempts_per_model}. Установка соединения...")
                    
                    response = None
                    retry_delay = 2.0
                    stream_ctx = None
                    for r_attempt in range(4):  # 1 основная + 3 ретрая при 429
                        if r_attempt > 0:
                            logger.warning(f"[OLLAMA-STREAM] Повторное подключение к {active_model} после 429 через {retry_delay} сек...")
                            await update_status(warning_text=f"ИИ-стрим перегружен (429). Повтор через {retry_delay} сек...")
                            await asyncio.sleep(retry_delay)
                            retry_delay *= 2.0
                            
                        stream_ctx = client.stream("POST", url, json=payload, headers=headers)
                        response = await stream_ctx.__aenter__()
                        logger.info(f"[OLLAMA-STREAM] Статус соединения {active_model} (повтор {r_attempt}): {response.status_code}")
                        if response.status_code != 429:
                            break
                        await stream_ctx.__aexit__(None, None, None)
                        stream_ctx = None
                        
                    response.raise_for_status()
                    logger.info(f"[OLLAMA-STREAM] Успешно подключено к {active_model}, читаем чанки...")
                    
                    # Обновляем текущую модель
                    current_model_var.set(active_model)
                    
                    in_reasoning = False
                    reasoning_ended = False
                    try:
                        async for line in response.aiter_lines():
                            if not line.strip():
                                continue
                            if line.startswith("data: "):
                                data_str = line[6:]
                                if data_str.strip() == "[DONE]":
                                    break
                                try:
                                    data_json = json.loads(data_str)
                                    delta = data_json["choices"][0]["delta"]
                                    content = delta.get("content", "") or ""
                                    reasoning = delta.get("reasoning", "") or delta.get("reasoning_content", "") or ""
                                    
                                    # Обработка мыслей/рассуждений на лету
                                    if reasoning:
                                        yielded_any = True
                                        if not in_reasoning and not reasoning_ended:
                                            in_reasoning = True
                                            yield "<think>\n" + reasoning
                                        else:
                                            yield reasoning
                                    elif content:
                                        yielded_any = True
                                        if in_reasoning:
                                            in_reasoning = False
                                            reasoning_ended = True
                                            yield "\n</think>\n" + content
                                        else:
                                            yield content
                                except Exception as e:
                                    logger.error(f"[OLLAMA-STREAM] Ошибка парсинга чанка: {e}")
                                    
                        # Закрываем тег рассуждений, если стрим закончился, а мы ещё внутри
                        if in_reasoning:
                            yield "\n</think>\n"
                            
                        # Если мы дошли до конца и отдали хоть один токен, генерация считается успешной
                        if yielded_any:
                            return
                    finally:
                        if stream_ctx:
                            await stream_ctx.__aexit__(None, None, None)
                except Exception as e:
                    logger.warning(f"[OLLAMA-STREAM] [СБОЙ] Стрим модели {active_model} на попытке {attempt} вернул ошибку: {repr(e)}")
                    if yielded_any:
                        logger.error(f"[OLLAMA-STREAM] Ошибка возникла посреди генерации, прерываем стрим.")
                        yield f"\n\n[Ошибка стриминга: {repr(e)}]"
                        return
                        
    # Если мы дошли сюда и ничего не отдали
    if not yielded_any:
        raise ValueError("Ни одна из доступных ИИ-моделей не смогла ответить на запрос в режиме стриминга.")


def parse_thoughts_and_answer(raw_text: str) -> tuple[str, str, bool]:
    text = raw_text.strip()
    
    # Регулярные выражения для поиска открывающего и закрывающего тегов (поддерживаем think, thinking, thought, thoughts)
    start_match = re.search(r'<\s*(?:mm:)?(?:think|thinking|thought|thoughts)(?:\s+[^>]*)?>', text, re.IGNORECASE)
    end_match = re.search(r'</\s*(?:mm:)?(?:think|thinking|thought|thoughts)\s*>', text, re.IGNORECASE)
    
    if start_match:
        think_start = start_match.start()
        think_content_start = start_match.end()
        
        if end_match and end_match.start() > think_start:
            think_end = end_match.start()
            answer_start = end_match.end()
            
            thoughts = text[think_content_start:think_end].strip()
            answer = text[answer_start:].strip()
            return thoughts, answer, False
        else:
            thoughts = text[think_content_start:].strip()
            return thoughts, "", True
    else:
        if end_match:
            think_end = end_match.start()
            answer_start = end_match.end()
            thoughts = text[:think_end].strip()
            answer = text[answer_start:].strip()
            return thoughts, answer, False
        else:
            return "", text, False

def format_draft_html(thoughts: str, answer: str, is_still_thinking: bool, show_thoughts: bool = True) -> str:
    html = ""
    clean_thoughts = clean_markdown_from_thoughts(thoughts) if thoughts else ""
    safe_thoughts = safe_html_cleaner(clean_thoughts) if clean_thoughts else ""
    safe_answer = safe_html_cleaner(answer) if answer else ""
    
    if show_thoughts and safe_thoughts:
        html += f"<blockquote expandable>{safe_thoughts}</blockquote>"
    if safe_answer:
        if html:
            html += f"\n\n{safe_answer}"
        else:
            html = safe_answer
    return auto_close_html_tags(html)

async def stream_to_message(status_msg: types.Message, messages: list, model_name: str = None, is_fast: bool = False, start_time: float = None) -> str:
    accumulated_raw_text = ""
    last_update_ts = 0.0
    
    async for chunk in ask_ollama_stream(messages, model_name):
        if not chunk:
            continue
        accumulated_raw_text += chunk
        
        now = time.monotonic()
        if now - last_update_ts >= 1.2:
            thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
            draft_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
            
            if not draft_html.strip():
                draft_html = "⏳ <i>Анализирую собранные данные и формулирую ответ...</i>"
                
            try:
                await status_msg.edit_text(
                    text=draft_html,
                    parse_mode=ParseMode.HTML,
                    disable_web_page_preview=True
                )
            except Exception as e:
                logger.error(f"Ошибка стриминга в сообщение: {e}")
            last_update_ts = now
            
    thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
    final_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
    if not final_html.strip():
        if is_fast:
            final_html = safe_html_cleaner(clean_markdown_from_thoughts(thoughts))
        else:
            final_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=True)
            
    if start_time:
        duration = time.monotonic() - start_time
        minutes = int(duration // 60)
        seconds = duration % 60
        if minutes > 0:
            time_str = f"{minutes} мин. {seconds:.1f} сек."
        else:
            time_str = f"{seconds:.1f} сек."
        final_html += f"\n\n⏱ <i>Время выполнения: {time_str}</i>"
    
    try:
        await status_msg.edit_text(
            text=final_html,
            parse_mode=ParseMode.HTML,
            disable_web_page_preview=True
        )
    except Exception as e:
        logger.error(f"Ошибка финального обновления сообщения: {e}")
        
    return clean_model_answer(answer) if answer else clean_model_answer(thoughts)

async def stream_to_inline(inline_message_id: str, messages: list, model_name: str = None, is_fast: bool = False, start_time: float = None) -> str:
    accumulated_raw_text = ""
    last_update_ts = 0.0
    
    async for chunk in ask_ollama_stream(messages, model_name):
        if not chunk:
            continue
        accumulated_raw_text += chunk
        
        now = time.monotonic()
        if now - last_update_ts >= 1.2:
            thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
            draft_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
            
            if not draft_html.strip():
                draft_html = "⏳ <i>Анализирую собранные данные и формулирую ответ...</i>"
                
            try:
                await bot.edit_message_text(
                    text=draft_html,
                    inline_message_id=inline_message_id,
                    parse_mode=ParseMode.HTML,
                    disable_web_page_preview=True
                )
            except Exception as e:
                logger.error(f"Ошибка стриминга в инлайн: {e}")
            last_update_ts = now
            
    thoughts, answer, is_still_thinking = parse_thoughts_and_answer(accumulated_raw_text)
    final_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=False)
    if not final_html.strip():
        if is_fast:
            final_html = safe_html_cleaner(clean_markdown_from_thoughts(thoughts))
        else:
            final_html = format_draft_html(thoughts, answer, is_still_thinking, show_thoughts=True)
    
    try:
        await bot.edit_message_text(
            text=final_html,
            inline_message_id=inline_message_id,
            parse_mode=ParseMode.HTML,
            disable_web_page_preview=True
        )
    except Exception as e:
        logger.error(f"Ошибка финального инлайн обновления: {e}")
        
    return clean_model_answer(answer) if answer else clean_model_answer(thoughts)

def parse_queries_from_response(response: str, default_query: str) -> list[str]:
    queries = []
    for line in response.strip().split("\n"):
        line = line.strip().strip('"').strip("'").strip("-").strip("*").strip()
        if line and (line[0].isdigit() or line.startswith("•")):
            while line and (line[0].isdigit() or line[0] in [".", ")", "-", " ", "*", "•"]):
                line = line[1:].strip()
        if line:
            queries.append(line)
    if not queries:
        queries = [default_query]
    return queries

def validate_queries(queries: list[str]) -> tuple[bool, str]:
    if not queries:
        return False, "Список поисковых запросов пуст."
        
    for q in queries:
        q_strip = q.strip()
        if not q_strip:
            continue
            
        lower_q = q_strip.lower()
        if "think" in lower_q or "mm:" in lower_q:
            return False, f"Запрос '{q_strip}' содержит служебные слова рассуждений (think, mm:)."
            
        words = q_strip.split()
        if len(words) > 10 or len(q_strip) > 80:
            return False, f"Запрос '{q_strip[:30]}...' слишком длинный (более 10 слов или 80 символов) и похож на обычный текст рассуждений."
            
    return True, ""

async def generate_initial_queries(query: str, chat_history: list, is_fast: bool = False, model_name: str = None, base64_images: list[str] = None, context_summary: str = None) -> tuple[list[str], str]:
    dt = get_current_datetime_str()
    
    if is_fast:
        system_prompt = (
            f"Текущие дата и время: {dt}.\n"
            "Ты — эксперт по быстрому поиску информации. Твоя задача — проанализировать историю переписки и последний вопрос пользователя, "
            "а затем сгенерировать ровно 3 различных поисковых запроса для поисковой системы SearxNG, чтобы собрать ОБЩУЮ информацию по теме вопроса.\n"
            "КРИТИЧЕСКИ ВАЖНО: Твой ответ должен состоять ТОЛЬКО из поисковых запросов. Запрещено использовать тег <think> и писать ход мыслей. Запрещено писать вводные фразы или любые другие пояснения. Просто выведи 3 поисковых запроса, каждый на новой строке, без кавычек и нумерации."
        )
    else:
        system_prompt = (
            f"Текущие дата и время: {dt}.\n"
            "Ты — эксперт по поиску информации. Твоя задача — проанализировать историю переписки и последний вопрос пользователя, "
            "а затем сгенерировать ровно 3 различных поисковых запроса для поисковой системы SearxNG, чтобы собрать ОБЩУЮ информацию по теме вопроса.\n"
            "КРИТИЧЕСКИ ВАЖНО: Твой ответ должен СТРОГО начинаться со своего хода мыслей, завернутого в тег <think>, например:\n"
            "<think>\nЗдесь твои подробные рассуждения о том, почему нужны эти запросы. Пиши только чистый простой текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков).\n</think>\n"
            "После закрывающего тега </think> сразу напиши поисковые запросы, каждый на новой строке, без кавычек и нумерации."
        )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    if is_fast:
        prompt_text = (
            f"Вопрос пользователя: {query}\n\n"
            "Сгенерируй поисковые запросы для этого вопроса.\n"
            "КРИТИЧЕСКИ ВАЖНО:\n"
            "1. Не используй тег <think> и не пиши свои рассуждения.\n"
            "2. Выведи только поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек, без нумерации и без пояснений."
        )
    else:
        prompt_text = (
            f"Вопрос пользователя: {query}\n\n"
            "Сгенерируй поисковые запросы для этого вопроса.\n"
            "КРИТИЧЕСКИ ВАЖНО:\n"
            "1. Начни свой ответ строго со своего хода мыслей на русском языке, завернутого в тег <think>...</think>.\n"
            "2. Внутри тегов <think> пиши исключительно чистый текст, БЕЗ использования Markdown-разметки (запрещено использовать **, *, _, #, списки и т.д.).\n"
            "3. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и без нумерации."
        )
        
    if base64_images and is_model_multimodal(model_name):
        content_structure = [{"type": "text", "text": prompt_text}]
        for b64 in base64_images:
            content_structure.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})
        messages.append({"role": "user", "content": content_structure})
    else:
        if base64_images and context_summary:
            prompt_text = f"Контекст присланных пользователем фото (Описание от визионера):\n{context_summary}\n\n" + prompt_text
        messages.append({"role": "user", "content": prompt_text})
    
    saved_thoughts = ""
    all_api_models = await get_cached_models()
    
    for attempt in range(1, 4):
        active_model = get_model_for_attempt(model_name, attempt, all_api_models)
        try:
            if active_model != (model_name or OLLAMA_MODEL) or attempt > 1:
                await update_status(warning_text=f"Переключаемся на ИИ {active_model} для подбора запросов (Попытка {attempt}/3)...")
                
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (initial_queries, модель {active_model}, попытка {attempt}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not is_fast and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                error_msg = "В твоем ответе СТРОГО отсутствует обязательный ход мыслей внутри тегов <think>...</think>. Ты обязан сначала написать свои рассуждения в теге <think>, а затем запросы."
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
                
            logger.warning(f"Попытка {attempt} ({active_model}) - Ошибка валидации запросов: {error_msg}")
            messages.append({"role": "assistant", "content": response})
            if is_fast:
                user_err_msg = (
                    f"Ошибка валидации запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ровно 3 поисковых запроса (каждый на новой строке, без кавычек, без нумерации и без какого-либо хода мыслей)."
                )
            else:
                user_err_msg = (
                    f"Ошибка форматирования запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ответ строго по формату:\n"
                    "1. Ход мыслей СТРОГО внутри тегов <think>...</think>.\n"
                    "2. После закрывающего тега </think> СРАЗУ выведи только поисковые запросы (каждый на новой строке, без кавычек, без нумерации и без лишних рассуждений)."
                )
            messages.append({
                "role": "user",
                "content": user_err_msg
            })
        except Exception as e:
            logger.error(f"Ошибка на попытке {attempt} ({active_model}) при генерации первичных запросов: {repr(e)}")
            
    # Если 3 основные попытки провалились, перебираем все остальные доступные модели из API
    tried_models = {
        get_model_for_attempt(model_name, 1, all_api_models),
        get_model_for_attempt(model_name, 2, all_api_models),
        get_model_for_attempt(model_name, 3, all_api_models)
    }
    
    full_remaining_chain = [m for m in all_api_models if m not in tried_models]
    
    for active_model in full_remaining_chain:
        await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} для подбора запросов...")
        try:
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (каскад {active_model}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not is_fast and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
        except Exception as e:
            logger.warning(f"Каскадный опрос модели {active_model} завершился ошибкой: {repr(e)}")
            
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла сгенерировать валидные поисковые запросы.")

async def generate_refined_queries(query: str, chat_history: list, research_state: str, model_name: str = None, context_summary: str = None) -> tuple[list[str], str]:
    dt = get_current_datetime_str()
    clean_query = get_clean_fallback_query(query)
    system_prompt = (
        f"Текущие дата и время: {dt}.\n"
        "Ты — эксперт по глубокому поиску. Тебе предоставлен текущий лог исследования (Research State) с первыми результатами.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Твой ответ должен СТРОГО начинаться со своего хода мыслей, завернутого в тег <think>...</think>.\n"
        "2. Внутри тегов <think> пиши исключительно чистый текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков).\n"
        "3. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и нумерации.\n"
        "4. Каждый поисковый запрос должен быть предельно кратким (не более 6-8 слов) и состоять только из ключевых фраз по теме поиска, а не из длинных предложений.\n"
        "5. Любые вводные или завершающие фразы, вежливые обращения, приветствия и мета-размышления после тега </think> СТРОГО ЗАПРЕЩЕНЫ."
    )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
    
    prompt = (
        f"Вопрос пользователя: {query}\n\n"
        f"Текущий лог исследования:\n{research_state}\n\n"
        "Сгенерируй 3 глубоких уточняющих запроса.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Начни свой ответ строго со своего хода мыслей на русском языке, завернутого в тег <think>...</think>.\n"
        "2. Внутри тегов <think> пиши исключительно чистый текст, БЕЗ использования Markdown-разметки (запрещено использовать **, *, _, #, списки и т.д.).\n"
        "3. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и без нумерации."
    )
    if context_summary:
        prompt = f"Контекст присланных пользователем файлов/фото (Краткая выжимка):\n{context_summary}\n\n" + prompt
        
    messages.append({"role": "user", "content": prompt})
    
    saved_thoughts = ""
    all_api_models = await get_cached_models()
    
    for attempt in range(1, 4):
        active_model = get_model_for_attempt(model_name, attempt, all_api_models)
        try:
            if active_model != (model_name or OLLAMA_MODEL) or attempt > 1:
                await update_status(warning_text=f"Переключаемся на ИИ {active_model} для подбора запросов (Попытка {attempt}/3)...")
                
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (refined_queries, модель {active_model}, попытка {attempt}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                error_msg = "В твоем ответе СТРОГО отсутствует обязательный ход мыслей внутри тегов <think>...</think>. Ты обязан сначала написать свои рассуждения в теге <think>, а затем запросы."
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
                
            logger.warning(f"Попытка {attempt} ({active_model}) - Ошибка валидации запросов: {error_msg}")
            await update_status(warning_text=f"Неверный формат запросов (Попытка {attempt}/3). Корректируем и генерируем заново...")
            messages.append({"role": "assistant", "content": response})
            messages.append({
                "role": "user",
                "content": (
                    f"Ошибка форматирования запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ответ строго по формату:\n"
                    "1. Ход мыслей СТРОГО внутри тегов <think>...</think>.\n"
                    "2. После закрывающего тега </think> СРАЗУ выведи только поисковые запросы (каждый на новой строке, без кавычек, без нумерации и без лишних рассуждений)."
                )
            })
        except Exception as e:
            logger.error(f"Ошибка на попытке {attempt} ({active_model}) при генерации уточняющих запросов: {repr(e)}")
            
    # Если 3 основные попытки провалились, перебираем все остальные доступные модели из API
    tried_models = {
        get_model_for_attempt(model_name, 1, all_api_models),
        get_model_for_attempt(model_name, 2, all_api_models),
        get_model_for_attempt(model_name, 3, all_api_models)
    }
    
    full_remaining_chain = [m for m in all_api_models if m not in tried_models]
    
    for active_model in full_remaining_chain:
        await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} для подбора запросов...")
        try:
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (каскад {active_model}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
        except Exception as e:
            logger.warning(f"Каскадный опрос модели {active_model} завершился ошибкой: {repr(e)}")
            
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла сгенерировать поисковые запросы.")

async def generate_opinion_queries(query: str, chat_history: list, research_state: str, model_name: str = None, context_summary: str = None) -> tuple[list[str], str]:
    dt = get_current_datetime_str()
    clean_query = get_clean_fallback_query(query)
    system_prompt = (
        f"Текущие дата и время: {dt}.\n"
        "Ты — эксперт по анализу общественного мнения и отзывов. На основе истории переписки, вопроса пользователя и уже собранных данных исследований:\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Твой ответ должен СТРОГО начинаться со своего хода мыслей, завернутого в тег <think>...</think>.\n"
        "2. Внутри тегов <think> пиши исключительно чистый текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков).\n"
        "3. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и нумерации.\n"
        "4. Каждый поисковый запрос должен быть предельно кратким (не более 6-8 слов) и состоять только из ключевых фраз по теме поиска, а не из длинных предложений.\n"
        "5. Любые вводные или завершающие фразы, вежливые обращения, приветствия и мета-размышления после тега </think> СТРОГО ЗАПРЕЩЕНЫ."
    )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    prompt = (
        f"Вопрос пользователя: {query}\n\n"
        f"Текущий лог исследования:\n{research_state}\n\n"
        "Сгенерируй 3 запроса для сбора мнений и отзывов.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Начни свой ответ строго со своего хода мыслей на русском языке, завернутого в тег <think>...</think>.\n"
        "2. Внутри тегов <think> пиши исключительно чистый текст, БЕЗ использования Markdown-разметки (запрещено использовать **, *, _, #, списки и т.д.).\n"
        "3. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и без нумерации."
    )
    if context_summary:
        prompt = f"Контекст присланных пользователем файлов/фото (Краткая выжимка):\n{context_summary}\n\n" + prompt
        
    messages.append({"role": "user", "content": prompt})
    
    saved_thoughts = ""
    all_api_models = await get_cached_models()
    
    for attempt in range(1, 4):
        active_model = get_model_for_attempt(model_name, attempt, all_api_models)
        try:
            if active_model != (model_name or OLLAMA_MODEL) or attempt > 1:
                await update_status(warning_text=f"Переключаемся на ИИ {active_model} для подбора запросов (Попытка {attempt}/3)...")
                
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (opinion_queries, модель {active_model}, попытка {attempt}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                error_msg = "В твоем ответе СТРОГО отсутствует обязательный ход мыслей внутри тегов <think>...</think>. Ты обязан сначала написать свои рассуждения в теге <think>, а затем запросы."
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
                
            logger.warning(f"Попытка {attempt} ({active_model}) - Ошибка валидации запросов: {error_msg}")
            await update_status(warning_text=f"Неверный формат запросов (Попытка {attempt}/3). Корректируем и генерируем заново...")
            messages.append({"role": "assistant", "content": response})
            messages.append({
                "role": "user",
                "content": (
                    f"Ошибка форматирования запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ответ строго по формату:\n"
                    "1. Ход мыслей СТРОГО внутри тегов <think>...</think>.\n"
                    "2. После закрывающего тега </think> СРАЗУ выведи только поисковые запросы (каждый на новой строке, без кавычек, без нумерации и без лишних рассуждений)."
                )
            })
        except Exception as e:
            logger.error(f"Ошибка на попытке {attempt} ({active_model}) при генерации запросов мнений: {repr(e)}")
            
    # Если 3 основные попытки провалились, перебираем все остальные доступные модели из API
    tried_models = {
        get_model_for_attempt(model_name, 1, all_api_models),
        get_model_for_attempt(model_name, 2, all_api_models),
        get_model_for_attempt(model_name, 3, all_api_models)
    }
    
    full_remaining_chain = [m for m in all_api_models if m not in tried_models]
    
    for active_model in full_remaining_chain:
        await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} для подбора запросов...")
        try:
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (каскад {active_model}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
        except Exception as e:
            logger.warning(f"Каскадный опрос модели {active_model} завершился ошибкой: {repr(e)}")
            
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла сгенерировать поисковые запросы.")

async def generate_cross_queries(query: str, chat_history: list, research_state: str, model_name: str = None, context_summary: str = None) -> tuple[list[str], str]:
    dt = get_current_datetime_str()
    clean_query = get_clean_fallback_query(query)
    system_prompt = (
        f"Текущие дата и время: {dt}.\n"
        "Ты — эксперт по анализу данных. Твоя задача — сопоставить уже найденную общую информацию и собранные мнения/отзывы по вопросу пользователя, "
        "выявить противоречия, сомнительные утверждения или пробелы, которые требуют дополнительной перепроверки в поисковике.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Твой ответ должен СТРОГО начинаться со своего хода мыслей, завернутого в тег <think>...</think>.\n"
        "2. Внутри тегов <think> пиши исключительно чистый текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков).\n"
        "3. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и нумерации.\n"
        "4. Каждый поисковый запрос должен быть предельно кратким (не более 6-8 слов) и состоять только из ключевых фраз по теме поиска, а не из длинных предложений.\n"
        "5. Любые вводные или завершающие фразы, вежливые обращения, приветствия и мета-размышления после тега </think> СТРОГО ЗАПРЕЩЕНЫ."
    )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    prompt = (
        f"Вопрос пользователя: {query}\n\n"
        f"Текущие результаты исследования:\n{research_state}\n\n"
        "Сгенерируй 3 поисковых запроса для проверки противоречий и сомнительных фактов.\n"
        "КРИТИЧЕСКИ ВАЖНО:\n"
        "1. Начни свой ответ строго со своего хода мыслей на русском языке, завернутого в тег <think>...</think>.\n"
        "2. Внутри тегов <think> пиши исключительно чистый текст, БЕЗ использования Markdown-разметки.\n"
        "3. После закрывающего тега </think> сразу напиши поисковые запросы (ровно 3 штуки), каждый на новой строке, без кавычек и без нумерации."
    )
    if context_summary:
        prompt = f"Контекст присланных пользователем файлов/фото (Краткая выжимка):\n{context_summary}\n\n" + prompt
        
    messages.append({"role": "user", "content": prompt})
    
    saved_thoughts = ""
    all_api_models = await get_cached_models()
    
    for attempt in range(1, 4):
        active_model = get_model_for_attempt(model_name, attempt, all_api_models)
        try:
            if active_model != (model_name or OLLAMA_MODEL) or attempt > 1:
                await update_status(warning_text=f"Переключаемся на ИИ {active_model} для подбора запросов (Попытка {attempt}/3)...")
                
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (cross_queries, модель {active_model}, попытка {attempt}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                error_msg = "В твоем ответе СТРОГО отсутствует обязательный ход мыслей внутри тегов <think>...</think>. Ты обязан сначала написать свои рассуждения в теге <think>, а затем запросы."
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
                
            logger.warning(f"Попытка {attempt} ({active_model}) - Ошибка валидации запросов: {error_msg}")
            await update_status(warning_text=f"Неверный формат запросов (Попытка {attempt}/3). Корректируем и генерируем заново...")
            messages.append({"role": "assistant", "content": response})
            messages.append({
                "role": "user",
                "content": (
                    f"Ошибка форматирования запросов: {error_msg}\n"
                    "Пожалуйста, исправь это и напиши ответ строго по формату:\n"
                    "1. Ход мыслей СТРОГО внутри тегов <think>...</think>.\n"
                    "2. После закрывающего тега </think> СРАЗУ выведи только поисковые запросы (каждый на новой строке, без кавычек, без нумерации и без лишних рассуждений)."
                )
            })
        except Exception as e:
            logger.error(f"Ошибка на попытке {attempt} ({active_model}) при генерации перекрестных запросов: {repr(e)}")
            
    # Если 3 основные попытки провалились, перебираем все остальные доступные модели из API
    tried_models = {
        get_model_for_attempt(model_name, 1, all_api_models),
        get_model_for_attempt(model_name, 2, all_api_models),
        get_model_for_attempt(model_name, 3, all_api_models)
    }
    
    full_remaining_chain = [m for m in all_api_models if m not in tried_models]
    
    for active_model in full_remaining_chain:
        await update_status(warning_text=f"Переключаемся на резервный ИИ {active_model} для подбора запросов...")
        try:
            response = await ask_ollama(messages, model_name=active_model)
            logger.info(f"Сырой ответ Ollama (каскад {active_model}):\n{response}")
            thoughts, queries_text, _ = parse_thoughts_and_answer(response)
            queries = parse_queries_from_response(queries_text, clean_query)
            
            if thoughts and thoughts.strip():
                saved_thoughts = thoughts
                
            is_valid, error_msg = validate_queries(queries)
            if is_valid and not thoughts.strip() and not saved_thoughts.strip():
                is_valid = False
                
            if is_valid:
                while len(queries) < 3:
                    queries.append(clean_query)
                return queries[:3], (thoughts if thoughts and thoughts.strip() else saved_thoughts)
        except Exception as e:
            logger.warning(f"Каскадный опрос модели {active_model} завершился ошибкой: {repr(e)}")
            
    raise ValueError("Ни одна из доступных ИИ-моделей не смогла сгенерировать поисковые запросы.")


async def run_cross_verification(query: str, general_context: str, opinions_context: str, model_name: str = None) -> str:
    dt = get_current_datetime_str()
    system_prompt = (
        f"Текущие дата и время: {dt}.\n"
        "Ты — ведущий аналитик данных. Проведи перекрестный анализ собранной информации для ответа на вопрос пользователя.\n"
        "У тебя есть два блока данных:\n"
        "1. Общая справочная информация (результаты этапов 1.1 и 1.2).\n"
        "2. Мнения людей, экспертов, отзывы, клинические исследования или личный опыт использования.\n\n"
        "Проанализируй эти данные и составь структурированный отчет:\n"
        "- Что из найденного совпадает и подтверждается обоими блоками?\n"
        "- Какие есть противоречия (разница между теорией/официальными данными и реальными отзывами)?\n"
        "- Что является доказанным фактом, а что — субъективным мнением или мифом?\n"
        "Пиши строго без использования Markdown-разметки (без звездочек, решеток, жирного шрифта, списков и т.д.), только чистый структурированный текст. На русском языке."
    )
    
    prompt = (
        f"Вопрос пользователя: {query}\n\n"
        f"БЛОК 1 (Общая информация):\n{general_context}\n\n"
        f"БЛОК 2 (Мнения и отзывы):\n{opinions_context}\n\n"
        "Составь аналитический отчет перекрестной проверки."
    )
    
    try:
        return await ask_ollama(prompt, system_prompt, model_name=model_name)
    except Exception as e:
        logger.error(f"Ошибка при перекрестном анализе: {e}")
        return "Не удалось провести перекрестный анализ данных."


async def process_single_query_parallel(
    query_str: str,
    fetched_urls_shared: set,
    limit: int,
    query_state_item: dict,
    mode: str = "deep",
    model_name: str = None,
    user_query: str = "",
    status_updater = None,
    search_queries_state: list = None,
    iteration: int = 1,
    current_thoughts: str = None,
    progress_percent: int = 30
) -> list[dict]:
    logger.info(f"[PARALLEL] Старт обработки запроса: '{query_str}'. Режим: {mode}")
    raw_results = await search_searxng_raw(query_str)
    
    # Фильтруем кандидатов, исключая уже скачанные
    candidates = []
    for res in raw_results:
        url = res.get("url")
        if not url or url in fetched_urls_shared:
            continue
        title = res.get("title", "Без названия").strip()
        snippet = res.get("content", "").strip()
        candidates.append({"url": url, "title": title, "snippet": snippet})
        
    logger.info(f"[PARALLEL] Для запроса '{query_str}' найдено {len(candidates)} потенциальных новых сайтов.")
    if not candidates:
        return []
        
    ordered_candidates = []
    
    if mode == "deep" and len(candidates) > limit:
        # Модель выбирает сайты
        candidates_to_rank = candidates[:8]
        search_list_text = ""
        for idx, cand in enumerate(candidates_to_rank, 1):
            search_list_text += f"{idx}. Заголовок: {cand['title']}\n   Описание: {cand['snippet']}\n   URL: {cand['url']}\n\n"
            
        system_prompt = (
            "Ты — эксперт по анализу релевантности веб-ресурсов. Твоя задача — проанализировать результаты поиска и выбрать ровно 3 наиболее полезных, информативных и авторитетных URL для ответа на вопрос пользователя.\n"
            "КРИТИЧЕСКИ ВАЖНО: Твой ответ должен содержать ТОЛЬКО выбранные URL, каждый на новой строке, без кавычек, без нумерации, без объяснений и без хода мыслей. Не пиши ничего, кроме самих URL."
        )
        
        prompt = (
            f"Оригинальный вопрос пользователя: {user_query}\n"
            f"Текущий поисковый запрос: {query_str}\n\n"
            f"Результаты поиска:\n{search_list_text}\n"
            "Выбери ровно 3 наиболее подходящих URL из списка выше."
        )
        
        try:
            logger.info(f"[PARALLEL] [OLLAMA] Запрос к ИИ на ранжирование сайтов для '{query_str}'...")
            async with ranking_semaphore:
                response = await ask_ollama(prompt, system_prompt, model_name=model_name)
            logger.info(f"[PARALLEL] [OLLAMA] Сырой ответ ИИ по выбору сайтов для '{query_str}':\n{response}")
            
            # Парсим URL из ответа модели
            selected_urls = []
            for line in response.strip().split("\n"):
                line = line.strip().strip('"').strip("'")
                found_url = None
                for cand in candidates_to_rank:
                    if cand["url"] in line or line in cand["url"]:
                        found_url = cand["url"]
                        break
                if found_url and found_url not in selected_urls:
                    selected_urls.append(found_url)
                    
            logger.info(f"[PARALLEL] ИИ выбрал {len(selected_urls)} релевантных URL: {selected_urls}")
            
            # Сначала добавляем в ordered_candidates те, которые выбрал ИИ
            for url in selected_urls:
                for cand in candidates:
                    if cand["url"] == url and cand not in ordered_candidates:
                        ordered_candidates.append(cand)
                        break
        except Exception as e:
            logger.error(f"[PARALLEL] [ОШИБКА] Не удалось ранжировать сайты с помощью ИИ: {repr(e)}")
            
    # Добавляем все оставшиеся кандидаты в ordered_candidates (для резерва или если режим быстрый)
    for cand in candidates:
        if cand not in ordered_candidates:
            ordered_candidates.append(cand)
            
    if not ordered_candidates:
        logger.warning(f"[PARALLEL] Нет сайтов для скачивания по запросу '{query_str}'")
        return []
        
    async def download_helper(url, clean_title):
        logger.info(f"[PARALLEL] Запуск фоновой закачки: {url}")
        try:
            content = await fetch_page_content(url)
            if content and content.strip():
                display_title = clean_title
                if len(display_title) > 40:
                    display_title = display_title[:37] + "..."
                
                # Добавляем во временный список для промежуточного статуса
                if display_title not in query_state_item["sites"]:
                    query_state_item["sites"].append(display_title)
                
                # Сразу обновляем статус в Telegram
                if status_updater and search_queries_state:
                    await status_updater(
                        make_status_html(
                            query=user_query,
                            status_text="Параллельный сбор информации..." if mode == "fast" else "Скачивание сайтов...",
                            progress_percent=progress_percent,
                            search_queries=search_queries_state,
                            attempt=iteration,
                            thoughts=current_thoughts
                        )
                    )
                return {"url": url, "title": display_title, "content": content}
        except Exception as e:
            logger.warning(f"[PARALLEL] [СБОЙ] Ошибка скачивания {url}: {repr(e)}")
        return None

    success_sites = []
    
    if mode == "fast":
        # В быстром поиске качаем первые limit параллельно
        wave1_size = min(limit, len(ordered_candidates))
        wave1_candidates = ordered_candidates[:wave1_size]
        
        for cand in wave1_candidates:
            fetched_urls_shared.add(cand["url"])
            
        tasks = [download_helper(cand["url"], cand["title"]) for cand in wave1_candidates]
        logger.info(f"[PARALLEL] Волна 1 (fast): Запуск параллельного скачивания {len(wave1_candidates)} сайтов...")
        downloaded = await asyncio.gather(*tasks)
        for res in downloaded:
            if res:
                success_sites.append(res)
    else:
        # В глубоком поиске используем две волны
        # Волна 1: Параллельно скачиваем первые 5 кандидатов
        wave1_size = min(5, len(ordered_candidates))
        wave1_candidates = ordered_candidates[:wave1_size]
        
        for cand in wave1_candidates:
            fetched_urls_shared.add(cand["url"])
            
        tasks = [download_helper(cand["url"], cand["title"]) for cand in wave1_candidates]
        logger.info(f"[PARALLEL] Волна 1: Запуск параллельного скачивания {len(wave1_candidates)} сайтов...")
        downloaded_wave1 = await asyncio.gather(*tasks)
        for res in downloaded_wave1:
            if res:
                success_sites.append(res)
                
        logger.info(f"[PARALLEL] Волна 1 завершена. Успешно скачано {len(success_sites)} сайтов.")
        
        # Волна 2: Если успешных сайтов меньше limit, качаем оставшихся по одному
        if len(success_sites) < limit and len(ordered_candidates) > wave1_size:
            logger.info(f"[PARALLEL] Успешных сайтов ({len(success_sites)}) < {limit}. Запуск резервной Волны 2.")
            remaining_candidates = ordered_candidates[wave1_size:]
            
            for cand in remaining_candidates:
                fetched_urls_shared.add(cand["url"])
                res = await download_helper(cand["url"], cand["title"])
                if res:
                    success_sites.append(res)
                    logger.info(f"[PARALLEL] Волна 2: Успешно скачан резервный сайт: {cand['url']}")
                    if len(success_sites) >= limit:
                        break
                        
    # Ограничиваем результат до limit (3) лучших сайтов
    final_success = success_sites[:limit]
    
    # Фиксируем в стейте только заголовки реально оставленных сайтов
    query_state_item["sites"] = [res["title"] for res in final_success]
    
    logger.info(f"[PARALLEL] Итого успешно скачано {len(final_success)} сайтов по запросу '{query_str}'.")
    return final_success


async def run_multistep_search(query: str, user_id: int, status_updater, is_inline: bool, chat_id: int = None, status_msg: types.Message = None, inline_message_id: str = None, mode: str = "deep", base64_images: list[str] = None, context_summary: str = None) -> str:
    start_time = time.monotonic()
    chat_history = user_histories.get(user_id, [])
    fetched_urls = set()
    research_state = ""
    dt = get_current_datetime_str()
    current_thoughts = None
    model_name = user_models.get(user_id, OLLAMA_MODEL)
    
    logger.info(f"[SEARCH-FLOW] Старт поиска. Запрос: '{query}', Режим: {mode}, Модель: {model_name}, Юзер ID: {user_id}, История сообщений: {len(chat_history)}, Картинок: {len(base64_images) if base64_images else 0}")
    current_model_var.set(model_name)
    current_status_updater_var.set(status_updater)
    current_query_var.set(query)
    current_status_text_var.set("Инициализация поиска...")
    current_progress_percent_var.set(0)
    current_search_queries_var.set(None)
    current_attempt_var.set(1)
    current_thoughts_var.set(None)
    
    for iteration in range(1, 4):
        iter_time_range = None
        if iteration == 2:
            iter_time_range = "year"
        elif iteration == 3:
            iter_time_range = "month"
        current_thoughts = None
        await status_updater(
            make_status_html(
                query=query,
                status_text="Первичный сбор информации...",
                progress_percent=20,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        try:
            initial_queries, initial_thoughts = await generate_initial_queries(
                query=query,
                chat_history=chat_history,
                is_fast=(mode == "fast"),
                model_name=current_model_var.get(),
                base64_images=base64_images,
                context_summary=context_summary
            )
        except Exception as e:
            err_text = "❌ К сожалению, все доступные ИИ-модели сейчас перегружены или недоступны. Пожалуйста, повторите запрос позже."
            logger.error(f"[SEARCH-FLOW] Критическая ошибка на этапе 1.1 (первичная генерация): {repr(e)}")
            await status_updater(err_text)
            raise ValueError(err_text)
        logger.info(f"Итерация {iteration} - Первичные запросы: {initial_queries}")
        if initial_thoughts and mode != "fast":
            current_thoughts = initial_thoughts
            
        search_queries_state = [{"query": q, "sites": []} for q in initial_queries]
        
        await status_updater(
            make_status_html(
                query=query,
                status_text="Первичный сбор информации...",
                progress_percent=30 if mode == "deep" else 50,
                search_queries=search_queries_state,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        async def delayed_task(item, delay):
            await asyncio.sleep(delay)
            return await process_single_query_parallel(
                query_str=item["query"],
                fetched_urls_shared=fetched_urls_shared,
                limit=3,
                query_state_item=item,
                mode=mode,
                model_name=current_model_var.get(),
                user_query=query,
                status_updater=status_updater,
                search_queries_state=search_queries_state,
                iteration=iteration,
                current_thoughts=current_thoughts,
                progress_percent=30 if mode == "deep" else 50,
                time_range=iter_time_range
            )
            
        fetched_urls_shared = set(fetched_urls)
        tasks = [
            delayed_task(item, i * 0.8)
            for i, item in enumerate(search_queries_state)
        ]
        
        await status_updater(
            make_status_html(
                query=query,
                status_text="Параллельный сбор информации по 3 запросам...",
                progress_percent=30 if mode == "deep" else 50,
                search_queries=search_queries_state,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        results_lists = await asyncio.gather(*tasks)
        
        initial_sites = []
        for sites_list in results_lists:
            initial_sites.extend(sites_list)
            for s in sites_list:
                fetched_urls.add(s["url"])
                
        await status_updater(
            make_status_html(
                query=query,
                status_text="Первичный сбор информации успешно завершен.",
                progress_percent=35 if mode == "deep" else 60,
                search_queries=search_queries_state,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
            
        initial_context_parts = []
        for i, site in enumerate(initial_sites, 1):
            initial_context_parts.append(
                f"Сайт {i}: {site['url']} ({site['title']})\n"
                f"Текст:\n{site['content']}\n"
            )
        initial_context = "\n".join(initial_context_parts)
        
        research_state += f"\n--- Попытка {iteration} ---\n"
        research_state += f"Первичные запросы: {', '.join(initial_queries)}\n"
        research_state += f"Найденные первичные материалы:\n{initial_context}\n"
        
        if mode == "fast":
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Синтез финального ответа...",
                    progress_percent=80,
                    search_queries=search_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            break
        
        # --- ФАЗА 1.2: Уточняющий сбор ---
        current_thoughts = None
        await status_updater(
            make_status_html(
                query=query,
                status_text="Анализ и уточняющий сбор...",
                progress_percent=40,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        refined_sites = []
        refined_queries = []
        refined_queries_state = []
        try:
            refined_queries, refined_thoughts = await generate_refined_queries(
                query=query,
                chat_history=chat_history,
                research_state=research_state,
                model_name=current_model_var.get(),
                context_summary=context_summary
            )
            logger.info(f"Итерация {iteration} - Уточняющие запросы: {refined_queries}")
            if refined_thoughts:
                current_thoughts = refined_thoughts
                
            refined_queries_state = [{"query": q, "sites": []} for q in refined_queries]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Анализ и уточняющий сбор...",
                    progress_percent=50,
                    search_queries=refined_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            async def delayed_task_refined(item, delay):
                await asyncio.sleep(delay)
                return await process_single_query_parallel(
                    query_str=item["query"],
                    fetched_urls_shared=fetched_urls_shared,
                    limit=3,
                    query_state_item=item,
                    mode=mode,
                    model_name=current_model_var.get(),
                    user_query=query,
                    status_updater=status_updater,
                    search_queries_state=refined_queries_state,
                    iteration=iteration,
                    current_thoughts=current_thoughts,
                    progress_percent=50,
                    time_range=iter_time_range
                )
                
            fetched_urls_shared = set(fetched_urls)
            tasks = [
                delayed_task_refined(item, i * 0.8)
                for i, item in enumerate(refined_queries_state)
            ]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Параллельный сбор уточняющей информации...",
                    progress_percent=50,
                    search_queries=refined_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            results_lists = await asyncio.gather(*tasks)
            
            for sites_list in results_lists:
                refined_sites.extend(sites_list)
                for s in sites_list:
                    fetched_urls.add(s["url"])
                    
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Уточняющий сбор информации успешно завершен.",
                    progress_percent=55,
                    search_queries=refined_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
        except Exception as e:
            logger.error(f"[SEARCH-FLOW] Пропуск уточняющего сбора из-за сбоя ИИ: {repr(e)}")
            break
            
        refined_context_parts = []
        for i, site in enumerate(refined_sites, 1):
            refined_context_parts.append(
                f"Сайт {i}: {site['url']} ({site['title']})\n"
                f"Текст:\n{site['content']}\n"
            )
        refined_context = "\n".join(refined_context_parts)
        
        if refined_queries:
            research_state += f"Уточняющие запросы: {', '.join(refined_queries)}\n"
            research_state += f"Найденные уточняющие материалы:\n{refined_context}\n"
        
        # --- ЭТАП 2.1: Сбор мнений ---
        current_thoughts = None
        await status_updater(
            make_status_html(
                query=query,
                status_text="Сбор мнений, отзывов и опыта...",
                progress_percent=60,
                attempt=iteration,
                thoughts=current_thoughts
            )
        )
        
        opinion_sites = []
        opinion_queries = []
        opinion_queries_state = []
        try:
            opinion_queries, opinion_thoughts = await generate_opinion_queries(
                query=query,
                chat_history=chat_history,
                research_state=research_state,
                model_name=current_model_var.get(),
                context_summary=context_summary
            )
            logger.info(f"Итерация {iteration} - Запросы мнений: {opinion_queries}")
            if opinion_thoughts:
                current_thoughts = opinion_thoughts
                
            opinion_queries_state = [{"query": q, "sites": []} for q in opinion_queries]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Сбор мнений, отзывов и опыта...",
                    progress_percent=70,
                    search_queries=opinion_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            async def delayed_task_opinion(item, delay):
                await asyncio.sleep(delay)
                return await process_single_query_parallel(
                    query_str=item["query"],
                    fetched_urls_shared=fetched_urls_shared,
                    limit=3,
                    query_state_item=item,
                    mode=mode,
                    model_name=current_model_var.get(),
                    user_query=query,
                    status_updater=status_updater,
                    search_queries_state=opinion_queries_state,
                    iteration=iteration,
                    current_thoughts=current_thoughts,
                    progress_percent=70,
                    time_range=iter_time_range
                )
                
            fetched_urls_shared = set(fetched_urls)
            tasks = [
                delayed_task_opinion(item, i * 0.8)
                for i, item in enumerate(opinion_queries_state)
            ]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Параллельный сбор мнений, отзывов и опыта...",
                    progress_percent=70,
                    search_queries=opinion_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            results_lists = await asyncio.gather(*tasks)
            
            for sites_list in results_lists:
                opinion_sites.extend(sites_list)
                for s in sites_list:
                    fetched_urls.add(s["url"])
                    
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Сбор мнений успешно завершен.",
                    progress_percent=75,
                    search_queries=opinion_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
        except Exception as e:
            logger.error(f"[SEARCH-FLOW] Пропуск сбора мнений из-за сбоя ИИ: {repr(e)}")
            break
            
        opinion_context_parts = []
        for i, site in enumerate(opinion_sites, 1):
            opinion_context_parts.append(
                f"Сайт мнений {i}: {site['url']} ({site['title']})\n"
                f"Текст:\n{site['content']}\n"
            )
        opinions_context = "\n".join(opinion_context_parts)
        
        if opinion_queries:
            research_state += f"Запросы мнений: {', '.join(opinion_queries)}\n"
            research_state += f"Найденные мнения/отзывы:\n{opinions_context}\n"
        
        # --- ЭТАП 2.2: Перекрестный анализ и поиск ---
        cross_sites = []
        cross_queries = []
        cross_queries_state = []
        try:
            current_thoughts = None
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Анализ противоречий и подготовка проверочных запросов...",
                    progress_percent=76,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            cross_queries, cross_thoughts = await generate_cross_queries(
                query=query,
                chat_history=chat_history,
                research_state=research_state,
                model_name=current_model_var.get(),
                context_summary=context_summary
            )
            logger.info(f"Итерация {iteration} - Проверочные запросы: {cross_queries}")
            if cross_thoughts:
                current_thoughts = cross_thoughts
                
            cross_queries_state = [{"query": q, "sites": []} for q in cross_queries]
            
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Параллельный сбор данных для перекрестной проверки...",
                    progress_percent=80,
                    search_queries=cross_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            async def delayed_task_cross(item, delay):
                await asyncio.sleep(delay)
                return await process_single_query_parallel(
                    query_str=item["query"],
                    fetched_urls_shared=fetched_urls_shared,
                    limit=3,
                    query_state_item=item,
                    mode=mode,
                    model_name=current_model_var.get(),
                    user_query=query,
                    status_updater=status_updater,
                    search_queries_state=cross_queries_state,
                    iteration=iteration,
                    current_thoughts=current_thoughts,
                    progress_percent=80
                )
                
            fetched_urls_shared = set(fetched_urls)
            tasks = [
                delayed_task_cross(item, i * 0.8)
                for i, item in enumerate(cross_queries_state)
            ]
            
            results_lists = await asyncio.gather(*tasks)
            
            for sites_list in results_lists:
                cross_sites.extend(sites_list)
                for s in sites_list:
                    fetched_urls.add(s["url"])
                    
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Сбор проверочных материалов успешно завершен.",
                    progress_percent=85,
                    search_queries=cross_queries_state,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            
            cross_context_parts = []
            for i, site in enumerate(cross_sites, 1):
                cross_context_parts.append(
                    f"Сайт проверки {i}: {site['url']} ({site['title']})\n"
                    f"Текст:\n{site['content']}\n"
                )
            cross_context = "\n".join(cross_context_parts)
            
            research_state += f"Проверочные запросы перекрестного анализа: {', '.join(cross_queries)}\n"
            research_state += f"Материалы перекрестной проверки:\n{cross_context}\n"
            
            general_combined_context = f"{initial_context}\n\n{refined_context}\n\n{cross_context}"
            logger.info(f"[SEARCH-FLOW] Запуск перекрестного анализа. Символов общего контекста: {len(general_combined_context)}, мнений: {len(opinions_context)}")
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Выполняем перекрестный анализ найденных материалов...",
                    progress_percent=88,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            cross_report = await run_cross_verification(query, general_combined_context, opinions_context, model_name=current_model_var.get())
            logger.info(f"[SEARCH-FLOW] Успешно завершен перекрестный анализ. Размер отчета: {len(cross_report)} символов.")
            research_state += f"Результат перекрестного анализа:\n{cross_report}\n"
            
            logger.info(f"[SEARCH-FLOW] Запуск проверки актуальности. Общая длина накопленного лога: {len(research_state)} символов.")
            await status_updater(
                make_status_html(
                    query=query,
                    status_text="Проверяем актуальность и свежесть информации...",
                    progress_percent=93,
                    attempt=iteration,
                    thoughts=current_thoughts
                )
            )
            is_current = await check_relevance(query, research_state, model_name=current_model_var.get())
            logger.info(f"[SEARCH-FLOW] Результат проверки актуальности: {is_current}")
            
            if is_current:
                current_thoughts = None
                await status_updater(
                    make_status_html(
                        query=query,
                        status_text="Синтез финального ответа...",
                        progress_percent=100,
                        attempt=iteration,
                        thoughts=current_thoughts
                    )
                )
                break
            else:
                research_state += "Результат проверки актуальности: НЕАКТУАЛЬНО. Необходим повторный поиск с уточнением временных параметров.\n"
                if iteration < 3:
                    await status_updater(
                        make_status_html(
                            query=query,
                            status_text="Данные устарели. Запуск повторного поиска с уточнением...",
                            progress_percent=95,
                            attempt=iteration,
                            warning_text="Обнаружены устаревшие данные. Уточняем временные рамки...",
                            thoughts=current_thoughts
                        )
                    )
                    await asyncio.sleep(3.0)
                else:
                    current_thoughts = None
                    await status_updater(
                        make_status_html(
                            query=query,
                            status_text="Синтез финального ответа...",
                            progress_percent=100,
                            attempt=iteration,
                            warning_text="Достигнут лимит попыток. Модель ответит на основе имеющихся данных.",
                            thoughts=current_thoughts
                        )
                    )
        except Exception as e:
            logger.error(f"[SEARCH-FLOW] Пропуск перекрестного анализа из-за сбоя ИИ: {repr(e)}")
            break
                
    # --- СИНТЕЗ ФИНАЛЬНОГО ОТВЕТА (СТРИМИНГ С КРАСИВЫМ HTML-ФОРМАТИРОВАНИЕМ) ---
    if mode == "fast":
        system_prompt = (
            f"Текущие дата и время: {dt}.\n"
            "Ты — полезный ИИ-помощник, который отвечает на вопрос пользователя на основе проведенного исследования.\n"
            "АНАЛИЗ КОНКРЕТНОСТИ ВОПРОСА:\n"
            "- Если вопрос конкретный, точечный и узкий (например, требует точное имя, дату, число, команду, ответ «да/нет» или один точный факт) и в нем нет просьбы расписать подробно, отвечай максимально КРАТКО И ПРЯМО (буквально в одно предложение или даже одно слово/число), без какого-либо расширения темы, предыстории или лишнего контекста.\n"
            "- Если вопрос широкий, открытый или размытый (например, «как работает X», «почему Y», «объясни Z»), твой ответ по умолчанию должен быть коротким, лаконичным и емким (обычно в пределах 1-2 небольших абзацев).\n"
            "КРИТИЧЕСКИ ВАЖНО: Сразу пиши финальный ответ. Запрещено использовать тег <think> и писать ход мыслей/рассуждения.\n"
            "Ты ДОЛЖЕН красиво и стильно оформлять свой ответ с использованием HTML-тегов, которые поддерживает Telegram. ИСПОЛЬЗУЙ МИНИМУМ ЭМОДЗИ (смайликов): максимум 1-2 штуки на все сообщение, только для важнейших акцентов (например, для предупреждения или вывода). Не ставь смайлики в качестве маркеров списков или в начале каждого предложения.\n"
            "- <b>Жирный текст</b> — используй для логических заголовков, выделения ключевых терминов и важных утверждений (не выделяй жирным целые абзацы, только важные фразы).\n"
            "- <i>Курсив</i> — используй для вводных слов, примечаний, названий книг, фильмов, цитат или второстепенного контекста.\n"
            "- <code>моноширинный текст</code> — используй для команд, чисел, дат, версий программ, формул, значений переменных или цен (всего, что пользователю может понадобиться скопировать в один клик).\n"
            "- <pre>блок текста</pre> — используй для многострочных блоков кода, логов или форматированных текстовых таблиц.\n"
            "- <blockquote>Цитата</blockquote> — используй для выделения важных цитат, ключевых выводов или определений.\n\n"
            "Правила оформления списков:\n"
            "- В Telegram Bot API нет тегов <ul> или <li>. Для списков используй символ переноса строки и маркеры: простые символы (например, •, ✦, ➔) или цифры '1.', '2.'. Никаких смайликов в качестве маркеров списков!\n\n"
            "Правила безопасности HTML:\n"
            "- Всегда следи за тем, чтобы все открытые теги (<b>, <i>, <code>, <blockquote>, <pre>) были корректно закрыты.\n"
            "- Не используй никакие другие теги HTML (такие как <p>, <div>, <a>, <ul>, <li>), так как Telegram выдаст ошибку парсинга."
        )
    else:
        system_prompt = (
            f"Текущие дата и время: {dt}.\n"
            "Ты — полезный ИИ-помощник, который отвечает на вопрос пользователя на основе проведенного исследования.\n"
            "КРИТИЧЕСКИ ВАЖНО: Начни свой ответ со своего хода мыслей, завернутого в тег <think>, например:\n"
            "<think>\nЗдесь твои подробные рассуждения и логика на русском языке. Пиши только чистый простой текст на русском языке, без использования Markdown-разметки (никаких **, *, _, # и списков).\n</think>\n"
            "После закрывающего тега </think> сразу напиши свой финальный ответ для пользователя. КРИТИЧЕСКИ ВАЖНО: Никаких вводных слов, пояснений, комментариев вроде 'Вот финальный ответ:', 'Хорошо, я понял задачу...', 'Я проанализировал...' и других мета-сообщений быть НЕ должно. Сразу выводи структурированную полезную суть ответа.\n"
            "АНАЛИЗ КОНКРЕТНОСТИ ВОПРОСА:\n"
            "- Если вопрос конкретный, точечный и узкий (например, требует точное имя, дату, число, команду, ответ «да/нет» или один точный факт) и в нем нет просьбы расписать подробно, отвечай максимально КРАТКО И ПРЯМО (буквально в одно предложение или даже одним словом/числом), без какого-либо расширения темы, предыстории или лишнего контекста.\n"
            "- Если вопрос широкий, открытый или размытый (например, «почему Y», «объясни Z»), то твой ответ по умолчанию должен быть структурированным и емким (обычно не более 2-3 небольших абзацев).\n"
            "Ты ДОЛЖЕН красиво и стильно оформлять свой финальный ответ с использованием HTML-тегов, которые поддерживает Telegram. ИСПОЛЬЗУЙ МИНИМУМ ЭМОДЗИ (смайликов): максимум 2-3 штуки на все сообщение, только для важнейших акцентов. Запрещено украшать ими каждый абзац или пункт списка.\n"
            "- <b>Жирный текст</b> — используй для логических заголовков, выделения ключевых терминов и важных утверждений (не выделяй жирным целые абзацы, только важные фразы).\n"
            "- <i>Курсив</i> — используй для вводных слов, примечаний, названий книг, фильмов, цитат или второстепенного контекста.\n"
            "- <code>моноширинный текст</code> — используй для команд, чисел, дат, версий программ, формул, значений переменных или цен (всего, что пользователю может понадобиться скопировать в один клик).\n"
            "- <pre>блок текста</pre> — используй для многострочных блоков кода, логов или форматированных текстовых таблиц.\n"
            "- <blockquote>Цитата</blockquote> — используй для выделения важных цитат, ключевых выводов или определений.\n\n"
            "Правила оформления списков:\n"
            "- В Telegram Bot API нет тегов <ul> или <li>. Для списков используй символ переноса строки и маркеры: простые символы (например, •, ✦, ➔) или цифры '1.', '2.'. Никаких смайликов в качестве маркеров списков!\n\n"
            "Правила безопасности HTML:\n"
            "- Всегда следи за тем, чтобы все открытые теги (<b>, <i>, <code>, <blockquote>, <pre>) были корректно закрыты.\n"
            "- Не используй никакие другие теги HTML (такие как <p>, <div>, <a>, <ul>, <li>), так как Telegram выдаст ошибку парсинга.\n"
            "- Твой финальный ответ должен включать результаты перекрестного анализа собранных данных (совпадения, противоречия между теорией и мнением людей, доказанные факты и субъективные мнения/мифы)."
        )
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
        
    if mode == "fast":
        prompt = (
            f"Вопрос пользователя: {query}\n\n"
            f"Полный отчет о проведенном исследовании:\n{research_state}\n\n"
            "Сформулируй финальный ответ для пользователя с использованием HTML-разметки. "
            "КРИТИЧЕСКИ ВАЖНО:\n"
            "1. Если вопрос конкретный и точечный (требует факт/число/команду/ответ «да» или «нет»), ответь на него максимально прямо и кратко (буквально 1 предложением или 1 словом), БЕЗ расширения темы и предысторий.\n"
            "2. Если вопрос широкий/размытый, отвечай коротко по существу (1-2 абзаца).\n"
            "3. Используй минимум эмодзи (максимум 1-2 штуки).\n"
            "4. Начни писать ответ сразу, БЕЗ использования тега <think> и без каких-либо рассуждений."
        )
    else:
        prompt = (
            f"Вопрос пользователя: {query}\n\n"
            f"Полный отчет о проведенном исследовании и результаты перекрестного анализа:\n{research_state}\n\n"
            "Сформулируй финальный ответ для пользователя с использованием HTML-разметки. "
            "ОБЯЗАТЕЛЬНО включи в ответ результаты перекрестного анализа собранных данных.\n"
            "КРИТИЧЕСКИ ВАЖНО:\n"
            "1. Твой ответ должен начинаться со своего хода мыслей, завернутого в тег <think>...</think>.\n"
            "2. Внутри тегов <think> пиши исключительно чистый текст на русском языке, БЕЗ использования Markdown-разметки (запрещено использовать **, *, _, #, списки и т.д.).\n"
            "3. После закрывающего тега </think> сразу напиши финальный ответ, без каких-либо вводных комментариев или фраз типа 'Вот финальный ответ:'.\n"
            "4. Если вопрос конкретный и точечный (требует конкретное число/имя/команду/да-нет), ответь на него максимально коротко и прямо (в 1 предложение/слово), не расширяя тему.\n"
            "5. Если вопрос широкий/размытый, отвечай структурированно (до 2-3 абзацев), если пользователь не просил иного.\n"
            "6. Используй минимум эмодзи (максимум 2-3 штуки на весь ответ)."
        )
        
    if base64_images and is_model_multimodal(model_name):
        content_structure = [{"type": "text", "text": prompt}]
        for b64 in base64_images:
            content_structure.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})
        messages.append({"role": "user", "content": content_structure})
    else:
        if base64_images and context_summary:
            prompt = f"Контекст присланных пользователем фото (Описание от визионера):\n{context_summary}\n\n" + prompt
        messages.append({"role": "user", "content": prompt})
        
    logger.info(f"[SEARCH-FLOW] Запуск финального синтеза ответа. Общее время поиска до этого момента: {time.monotonic() - start_time:.2f} сек. Размер лога исследования: {len(research_state)} симв.")
    if is_inline:
        answer_text = await stream_to_inline(inline_message_id, messages, model_name=current_model_var.get(), is_fast=(mode == "fast"), start_time=start_time)
    else:
        answer_text = await stream_to_message(status_msg, messages, model_name=current_model_var.get(), is_fast=(mode == "fast"), start_time=start_time)
    logger.info(f"[SEARCH-FLOW] Финальный ответ успешно отправлен и сохранен. Длина ответа: {len(answer_text)} симв. Полное время: {time.monotonic() - start_time:.2f} сек.")
        
    # Сохраняем историю
    if user_id not in user_histories:
        user_histories[user_id] = []
        
    if base64_images:
        history_query = f"[Фото: {context_summary or 'Изображения'}] {query}"
    elif context_summary and "Содержимое прикрепленного файла" in query:
        history_query = f"[Файл] {query.split('Вопрос пользователя по файлу:')[-1].strip() if 'Вопрос пользователя по файлу:' in query else query[:100]}"
    else:
        history_query = query
        
    user_histories[user_id].append({"role": "user", "content": history_query})
    user_histories[user_id].append({"role": "assistant", "content": answer_text})
    
    return answer_text

@dp.message(Command("start"))
async def start_cmd(message: types.Message):
    await message.answer(
        "👋 Привет! Я инлайн-агент для поиска информации в интернете.\n\n"
        "Для работы со мной начните писать в любом чате: <code>@имя_этого_бота ваш запрос</code> "
        "и выберите предложенный вариант. Сообщение будет отправлено и динамически обновлено на основе результатов поиска и анализа ИИ."
    )

async def clear_cmd(message: types.Message):
    user_id = message.from_user.id
    if user_id in user_histories:
        user_histories[user_id] = []
        logger.info(f"История диалога для пользователя {user_id} очищена.")
        
    current_id = message.message_id
    deleted_count = 0
    status = await message.answer("⏳ Очищаю историю сообщений (до 100 недавних)...")
    
    for msg_id in range(current_id, current_id - 100, -1):
        if msg_id == status.message_id:
            continue
        try:
            await bot.delete_message(chat_id=message.chat.id, message_id=msg_id)
            deleted_count += 1
        except Exception:
            continue
            
    try:
        bot_user = await bot.get_me()
        bot_username = bot_user.username or "SearchXNG_Bot"
        await status.edit_text(
            "🧹 <b>История сообщений успешно очищена! Бот готов к работе.</b>\n\n"
            "Вы можете отправить мне новый поисковый запрос прямо в этом чате или использовать меня в инлайн-режиме в любом другом чате:\n"
            f"🔍 <code>@{bot_username} ваш запрос</code>",
            parse_mode=ParseMode.HTML
        )
        last_clear_message_ids[user_id] = status.message_id
    except Exception as e:
        logger.error(f"Ошибка при обновлении статуса очистки: {e}")

api_models_cache = {"models": [], "expiry": 0}

async def get_available_models_from_api() -> list[str]:
    url = f"{OLLAMA_BASE_URL.rstrip('/')}/v1/models"
    headers = {}
    if OLLAMA_API_KEY:
        headers["Authorization"] = f"Bearer {OLLAMA_API_KEY}"
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            data = response.json()
            models = [m["id"] for m in data.get("data", [])]
            models.sort()
            return models
    except Exception as e:
        logger.error(f"Не удалось получить список моделей из API: {e}")
        return [OLLAMA_MODEL]

async def get_cached_models() -> list[str]:
    now = time.time()
    if not api_models_cache["models"] or now > api_models_cache["expiry"]:
        models = await get_available_models_from_api()
        if models:
            api_models_cache["models"] = models
            api_models_cache["expiry"] = now + 120
    return api_models_cache["models"]

def get_model_pagination_keyboard(models: list[str], current_model: str, page: int = 0) -> types.InlineKeyboardMarkup:
    PAGE_SIZE = 8
    total_pages = (len(models) + PAGE_SIZE - 1) // PAGE_SIZE
    
    if page < 0:
        page = 0
    if page >= total_pages:
        page = total_pages - 1
        
    start_idx = page * PAGE_SIZE
    end_idx = start_idx + PAGE_SIZE
    page_models = models[start_idx:end_idx]
    
    buttons = []
    row = []
    for m_id in page_models:
        # Добавляем галочку, если модель совпадает или является дефолтной
        display_name = f"✅ {m_id}" if current_model == m_id else m_id
        row.append(types.InlineKeyboardButton(text=display_name, callback_data=f"setmodel:{m_id}:{page}"))
        if len(row) == 2:
            buttons.append(row)
            row = []
    if row:
        buttons.append(row)
        
    nav_row = []
    if page > 0:
        nav_row.append(types.InlineKeyboardButton(text="⬅️ Назад", callback_data=f"modelpage:{page-1}"))
    else:
        nav_row.append(types.InlineKeyboardButton(text="❌", callback_data="modelpage_noop"))
        
    nav_row.append(types.InlineKeyboardButton(text=f"📄 {page+1} / {total_pages}", callback_data="modelpage_noop"))
    
    if page < total_pages - 1:
        nav_row.append(types.InlineKeyboardButton(text="Вперед ➡️", callback_data=f"modelpage:{page+1}"))
    else:
        nav_row.append(types.InlineKeyboardButton(text="❌", callback_data="modelpage_noop"))
        
    buttons.append(nav_row)
    close_row = [types.InlineKeyboardButton(text="❌ Закрыть меню", callback_data="close_model_menu")]
    buttons.append(close_row)
    return types.InlineKeyboardMarkup(inline_keyboard=buttons)

@dp.message(Command("model"))
async def model_cmd(message: types.Message):
    user_id = message.from_user.id
    current_model = user_models.get(user_id, OLLAMA_MODEL)
    
    models = await get_cached_models()
    keyboard = get_model_pagination_keyboard(models, current_model, 0)
    
    await message.answer(
        text=(
            "🤖 <b>Выбор языковой модели ИИ</b>\n\n"
            f"Текущая активная модель: <code>{current_model}</code>\n\n"
            "Выберите модель из списка ниже для использования в ваших поисках:"
        ),
        parse_mode=ParseMode.HTML,
        reply_markup=keyboard
    )

@dp.inline_query()
async def inline_query_handler(inline_query: types.InlineQuery):
    query = inline_query.query.strip()
    if not query:
        return
    
    # Генерируем уникальный короткий ключ для кэширования запроса (Telegram ID limit = 64 chars)
    cache_key = uuid.uuid4().hex[:8]
    inline_queries_cache[cache_key] = {"query": query, "timestamp": time.time()}
    
    keyboard = types.InlineKeyboardMarkup(
        inline_keyboard=[[
            types.InlineKeyboardButton(text="⏳ Идет поиск...", callback_data="loading_dummy")
        ]]
    )
    
    results = [
        types.InlineQueryResultArticle(
            id=f"fast:{cache_key}",
            title=f"⚡ Быстрый поиск: {query[:35]}...",
            description="Полноценный сбор за 1 шаг (в пределах 2 минут)",
            input_message_content=types.InputTextMessageContent(
                message_text=(
                    f"🔍 <b>Запрос:</b> <code>{query}</code>\n"
                    f"📊 <code>[■■□□□□□□□□] 20%</code>\n\n"
                    f"🧠 <b>Статус:</b> Инициализация быстрого поиска..."
                ),
                parse_mode=ParseMode.HTML
            ),
            reply_markup=keyboard
        ),
        types.InlineQueryResultArticle(
            id=f"deep:{cache_key}",
            title=f"🔍 Глубокий поиск: {query[:35]}...",
            description="Глубокое многошаговое исследование (от 5 до 10 минут)",
            input_message_content=types.InputTextMessageContent(
                message_text=(
                    f"🔍 <b>Запрос:</b> <code>{query}</code>\n"
                    f"📊 <code>[■■□□□□□□□□] 20%</code>\n\n"
                    f"🧠 <b>Статус:</b> Инициализация глубокого поиска..."
                ),
                parse_mode=ParseMode.HTML
            ),
            reply_markup=keyboard
        )
    ]
    
    try:
        await inline_query.answer(results, cache_time=0, is_personal=True)
    except Exception as e:
        logger.error(f"Ошибка отправки инлайн ответа: {e}")

@dp.chosen_inline_result()
async def chosen_inline_result_handler(chosen_inline_result: types.ChosenInlineResult):
    result_id = chosen_inline_result.result_id
    inline_message_id = chosen_inline_result.inline_message_id
    user_id = chosen_inline_result.from_user.id
    if not inline_message_id:
        logger.warning("inline_message_id отсутствует в ChosenInlineResult")
        return
        
    mode = "deep"
    query = chosen_inline_result.query.strip()
    
    if result_id and ":" in result_id:
        try:
            mode_part, cache_key = result_id.split(":", 1)
            if mode_part in ["fast", "deep"]:
                mode = mode_part
                cached_data = inline_queries_cache.pop(cache_key, None)
                if cached_data:
                    query = cached_data["query"]
        except Exception as e:
            logger.error(f"Ошибка извлечения данных из инлайн-кэша: {e}")
    
    logger.info(f"Пользователь {user_id} выбрал режим {mode} для запроса: '{query}'. ID сообщения: {inline_message_id}")
    asyncio.create_task(process_inline_search(inline_message_id, query, user_id, mode))

async def process_inline_search(inline_message_id: str, query: str, user_id: int, mode: str = "deep"):
    try:
        async def update_status_func(text: str):
            try:
                await bot.edit_message_text(
                    text=text,
                    inline_message_id=inline_message_id,
                    parse_mode=ParseMode.HTML
                )
            except Exception as e:
                logger.error(f"Не удалось обновить инлайн статус: {e}")
                
        await run_multistep_search(
            query=query,
            user_id=user_id,
            status_updater=update_status_func,
            is_inline=True,
            inline_message_id=inline_message_id,
            mode=mode
        )
        
    except Exception as e:
        logger.error(f"Ошибка при обработке инлайн-поиска: {e}")
        try:
            await bot.edit_message_text(
                text=f"❌ Произошла ошибка при обработке запроса: <i>{query}</i>\n\nДетали ошибки: {e}",
                inline_message_id=inline_message_id,
                parse_mode=ParseMode.HTML,
                reply_markup=None
            )
        except Exception as edit_err:
            logger.error(f"Не удалось обновить статус ошибки в инлайн сообщении: {edit_err}")

async def handle_photos_processing(photos: list, caption: str, chat_id: int, user_id: int):
    status_msg = await bot.send_message(
        chat_id=chat_id,
        text="📥 <b>Получено фото. Скачивание и анализ изображений...</b>",
        parse_mode=ParseMode.HTML
    )
    
    try:
        base64_images = []
        descriptions = []
        
        async def process_photo(photo_size):
            try:
                file = await bot.get_file(photo_size.file_id)
                file_io = BytesIO()
                await bot.download_file(file.file_path, file_io)
                b64_data = base64.b64encode(file_io.getvalue()).decode("utf-8")
                try:
                    model_name = user_models.get(user_id, OLLAMA_MODEL)
                    desc = await describe_image(b64_data, model_name=model_name)
                except Exception as ve:
                    logger.error(f"Не удалось распознать изображение через визионер: {ve}")
                    desc = "Ошибка распознавания изображения визионером."
                return b64_data, desc
            except Exception as fe:
                logger.error(f"Не удалось скачать файл изображения: {fe}")
                return None, None
                
        tasks = [process_photo(photo) for photo in photos]
        results = await asyncio.gather(*tasks)
        results = [r for r in results if r is not None and r[0] is not None]
        
        if not results:
            raise ValueError("Не удалось загрузить ни одно изображение.")
            
        for b64, desc in results:
            base64_images.append(b64)
            descriptions.append(desc)
            
        combined_description = ""
        if len(descriptions) == 1:
            combined_description = descriptions[0]
        else:
            for idx, desc in enumerate(descriptions, 1):
                combined_description += f"Изображение {idx}:\n{desc}\n\n"
        
        combined_description = combined_description.strip()
        user_query = caption if caption else "Проанализируй присланные изображения."
        
        cache_key = uuid.uuid4().hex[:8]
        private_queries_cache[cache_key] = {
            "query": f"[Фото-запрос] Вопрос пользователя по фото: {user_query}",
            "base64_images": base64_images,
            "context_summary": combined_description,
            "timestamp": time.time()
        }
        
        preview = combined_description[:300] + "..." if len(combined_description) > 300 else combined_description
        
        keyboard = types.InlineKeyboardMarkup(
            inline_keyboard=[[
                types.InlineKeyboardButton(text="⚡ Быстрый поиск", callback_data=f"fast:{cache_key}"),
                types.InlineKeyboardButton(text="🔍 Глубокий поиск", callback_data=f"deep:{cache_key}")
            ]]
        )
        
        await status_msg.edit_text(
            text=(
                f"📷 <b>Анализ изображений завершен (всего: {len(photos)}).</b>\n"
                f"📝 <b>Описание:</b> <i>{preview}</i>\n"
                f"❓ <b>Запрос:</b> <code>{user_query[:50]}...</code>\n\n"
                f"Выберите режим исследования:"
            ),
            parse_mode=ParseMode.HTML,
            reply_markup=keyboard
        )
    except Exception as e:
        logger.error(f"Ошибка при обработке фото: {e}")
        await status_msg.edit_text(
            text=f"❌ Произошла ошибка при обработке фото: {e}",
            parse_mode=ParseMode.HTML
        )

async def process_media_group_after_delay(media_group_id: str, chat_id: int, user_id: int):
    await asyncio.sleep(1.5)
    data = media_groups_data.pop(media_group_id, None)
    if not data:
        return
    messages = data["messages"]
    messages.sort(key=lambda m: m.message_id)
    
    photos = []
    caption = ""
    for msg in messages:
        if msg.photo:
            photos.append(msg.photo[-1])
        if msg.caption and not caption:
            caption = msg.caption.strip()
            
    await handle_photos_processing(photos, caption, chat_id, user_id)

@dp.message(F.photo)
async def photo_message_handler(message: types.Message):
    logger.info(f"Получено фото: ID={message.message_id}, MediaGroup={message.media_group_id}")
    media_group_id = message.media_group_id
    if media_group_id:
        if media_group_id not in media_groups_data:
            media_groups_data[media_group_id] = {
                "messages": [],
                "timer": asyncio.create_task(process_media_group_after_delay(media_group_id, message.chat.id, message.from_user.id))
            }
        media_groups_data[media_group_id]["messages"].append(message)
    else:
        await handle_photos_processing([message.photo[-1]], message.caption.strip() if message.caption else "", message.chat.id, message.from_user.id)

@dp.message(F.document)
async def document_message_handler(message: types.Message):
    file_id = message.document.file_id
    file_name = message.document.file_name or "document"
    logger.info(f"Получен документ: ID={message.message_id}, Имя={file_name}")
    
    status_msg = await bot.send_message(
        chat_id=message.chat.id,
        text=f"📥 <b>Скачивание и анализ документа '{file_name}'...</b>",
        parse_mode=ParseMode.HTML
    )
    
    try:
        file_info = await bot.get_file(file_id)
        
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, file_name)
            await bot.download_file(file_info.file_path, temp_file_path)
            file_text = await extract_text_from_document(temp_file_path, file_name)
            
        if not file_text.strip():
            raise ValueError("Документ пуст или не содержит распознаваемого текста.")
            
        file_text = file_text[:20000]
        
        model_name = user_models.get(message.from_user.id, OLLAMA_MODEL)
        summary = await generate_document_summary(file_name, file_text, model_name=model_name)
        
        user_query = message.caption.strip() if message.caption else "Проанализируй содержимое прикрепленного файла."
        
        cache_key = uuid.uuid4().hex[:8]
        private_queries_cache[cache_key] = {
            "query": f"[Файл-запрос] Содержимое прикрепленного файла '{file_name}':\n\n{file_text}\n\nВопрос пользователя по файлу: {user_query}",
            "base64_images": None,
            "context_summary": summary,
            "timestamp": time.time()
        }
        
        preview = summary[:300] + "..." if len(summary) > 300 else summary
        
        keyboard = types.InlineKeyboardMarkup(
            inline_keyboard=[[
                types.InlineKeyboardButton(text="⚡ Быстрый поиск", callback_data=f"fast:{cache_key}"),
                types.InlineKeyboardButton(text="🔍 Глубокий поиск", callback_data=f"deep:{cache_key}")
            ]]
        )
        
        await status_msg.edit_text(
            text=(
                f"📄 <b>Файл '{file_name}' успешно обработан.</b>\n"
                f"📝 <b>Выжимка:</b> <i>{preview}</i>\n"
                f"❓ <b>Запрос:</b> <code>{user_query[:50]}...</code>\n\n"
                f"Выберите режим исследования:"
            ),
            parse_mode=ParseMode.HTML,
            reply_markup=keyboard
        )
        
    except Exception as e:
        logger.error(f"Ошибка при обработке документа: {e}")
        await status_msg.edit_text(
            text=f"❌ Произошла ошибка при обработке документа: {e}",
            parse_mode=ParseMode.HTML
        )

@dp.message()
async def private_message_handler(message: types.Message):
    print(f"DEBUG LOG: Получено сообщение! ID: {message.message_id}, Текст: {message.text}, Тип чата: {message.chat.type}", flush=True)
    
    if message.text and message.text.startswith("/"):
        if message.text.startswith("/start"):
            await start_cmd(message)
        elif message.text.startswith("/clear"):
            await clear_cmd(message)
        elif message.text.startswith("/model"):
            await model_cmd(message)
        return
        
    query = message.text.strip() if message.text else ""
    if not query:
        print("DEBUG LOG: Пустой запрос, игнорируем.", flush=True)
        return
        
    user_id = message.from_user.id
    
    if user_id in last_clear_message_ids:
        clear_msg_id = last_clear_message_ids.pop(user_id)
        try:
            await bot.delete_message(chat_id=message.chat.id, message_id=clear_msg_id)
            logger.info(f"Сообщение об очистке {clear_msg_id} удалено перед новым запросом.")
        except Exception as e:
            logger.warning(f"Не удалось удалить сообщение об очистке {clear_msg_id}: {e}")
            
    # Генерируем уникальный ключ и сохраняем запрос в кэш
    cache_key = uuid.uuid4().hex[:8]
    private_queries_cache[cache_key] = {"query": query, "timestamp": time.time()}
    
    keyboard = types.InlineKeyboardMarkup(
        inline_keyboard=[[
            types.InlineKeyboardButton(text="⚡ Быстрый поиск", callback_data=f"fast:{cache_key}"),
            types.InlineKeyboardButton(text="🔍 Глубокий поиск", callback_data=f"deep:{cache_key}")
        ]]
    )
    
    await message.answer(
        text=(
            f"❓ <b>Запрос:</b> <code>{query[:50]}...</code>\n\n"
            "Выберите режим исследования:"
        ),
        parse_mode=ParseMode.HTML,
        reply_markup=keyboard
    )

@dp.callback_query()
async def callback_query_handler(callback_query: types.CallbackQuery):
    if callback_query.data == "loading_dummy":
        await callback_query.answer()
        return
        
    if callback_query.data == "close_model_menu":
        try:
            await callback_query.message.delete()
        except Exception:
            try:
                await callback_query.message.edit_reply_markup(reply_markup=None)
            except Exception:
                pass
        await callback_query.answer()
        return
        
    data_str = callback_query.data or ""
    if ":" not in data_str:
        await callback_query.answer()
        return
        
    prefix, payload_part = data_str.split(":", 1)
    
    if prefix == "setmodel":
        user_id = callback_query.from_user.id
        if ":" in payload_part:
            model_id, page_str = payload_part.rsplit(":", 1)
            try:
                page = int(page_str)
            except ValueError:
                page = 0
        else:
            model_id = payload_part
            page = 0
            
        models = await get_cached_models()
        if model_id in models:
            user_models[user_id] = model_id
            save_user_models()
            logger.info(f"Пользователь {user_id} переключил модель на '{model_id}' и настройки были сохранены.")
            
            try:
                await callback_query.message.delete()
            except Exception:
                try:
                    await callback_query.message.edit_text(
                        text=f"🤖 <b>Выбор языковой модели ИИ</b>\n\n✅ Модель изменена на: <code>{model_id}</code>",
                        parse_mode=ParseMode.HTML,
                        reply_markup=None
                    )
                except Exception:
                    pass
                
            await callback_query.answer(f"Модель изменена на {model_id}!", show_alert=False)
        else:
            await callback_query.answer("⚠️ Неизвестная модель ИИ.", show_alert=True)
        return
        
    elif prefix == "modelpage":
        user_id = callback_query.from_user.id
        current_model = user_models.get(user_id, OLLAMA_MODEL)
        try:
            page = int(payload_part)
        except ValueError:
            page = 0
            
        models = await get_cached_models()
        keyboard = get_model_pagination_keyboard(models, current_model, page)
        try:
            await callback_query.message.edit_reply_markup(reply_markup=keyboard)
        except Exception as e:
            logger.error(f"Не удалось переключить страницу моделей: {e}")
        await callback_query.answer()
        return
        
    if prefix not in ["fast", "deep"]:
        await callback_query.answer()
        return
        
    mode = prefix
    cache_key = payload_part
    
    query_data = private_queries_cache.pop(cache_key, None)
    if not query_data:
        await callback_query.answer("⚠️ Запрос устарел. Пожалуйста, отправьте его заново.", show_alert=True)
        try:
            await callback_query.message.delete()
        except Exception:
            pass
        return
        
    query = query_data["query"]
    base64_images = query_data.get("base64_images")
    context_summary = query_data.get("context_summary")
    status_text = "Инициализация быстрого поиска..." if mode == "fast" else "Инициализация глубокого поиска..."
    status_msg = callback_query.message
    
    try:
        await status_msg.edit_text(
            text=(
                f"🔍 <b>Запрос:</b> <code>{query[:50]}...</code>\n"
                f"📊 <code>[■■□□□□□□□□] 20%</code>\n\n"
                f"🧠 <b>Статус:</b> {status_text}"
            ),
            parse_mode=ParseMode.HTML,
            reply_markup=None
        )
    except Exception as e:
        logger.error(f"Не удалось переключить сообщение в статус поиска: {e}")
        await callback_query.answer("Произошла ошибка инициализации.", show_alert=True)
        return
        
    async def run_search_task():
        try:
            async def update_status_func(text: str):
                try:
                    await status_msg.edit_text(text=text, parse_mode=ParseMode.HTML)
                except Exception as e:
                    logger.error(f"Не удалось обновить статус: {e}")
                    
            async with ChatActionSender.typing(bot=bot, chat_id=status_msg.chat.id):
                await run_multistep_search(
                    query=query,
                    user_id=callback_query.from_user.id,
                    status_updater=update_status_func,
                    is_inline=False,
                    chat_id=status_msg.chat.id,
                    status_msg=status_msg,
                    mode=mode,
                    base64_images=base64_images,
                    context_summary=context_summary
                )
        except Exception as e:
            logger.error(f"Ошибка при обработке ЛС в асинхронном таске: {e}")
            try:
                await status_msg.answer(
                    text=f"❌ Произошла ошибка при обработке запроса: <i>{query}</i>\n\nДетали ошибки: {e}",
                    parse_mode=ParseMode.HTML
                )
            except Exception as edit_err:
                logger.error(f"Не удалось отправить сообщение об ошибке: {edit_err}")

    asyncio.create_task(run_search_task())
    await callback_query.answer()

async def clean_expired_caches():
    logger.info("[CACHE] Запущена фоновая задача очистки устаревших кэшей.")
    while True:
        await asyncio.sleep(600)  # Каждые 10 минут
        try:
            now = time.time()
            expiry_time = 3600  # 1 час
            
            # Очищаем inline_queries_cache
            expired_inline = [
                k for k, v in inline_queries_cache.items()
                if isinstance(v, dict) and now - v.get("timestamp", 0) > expiry_time
            ]
            for k in expired_inline:
                inline_queries_cache.pop(k, None)
                
            # Очищаем private_queries_cache
            expired_private = [
                k for k, v in private_queries_cache.items()
                if isinstance(v, dict) and now - v.get("timestamp", 0) > expiry_time
            ]
            for k in expired_private:
                private_queries_cache.pop(k, None)
                
            if expired_inline or expired_private:
                logger.info(f"[CACHE] Очищено устаревших записей: inline={len(expired_inline)}, private={len(expired_private)}")
        except Exception as e:
            logger.error(f"[CACHE] Ошибка при очистке кэшей: {e}")

async def main():
    logger.info("Регистрация команд в Telegram...")
    await bot.set_my_commands([
        types.BotCommand(command="start", description="Запустить бота"),
        types.BotCommand(command="model", description="Выбрать модель ИИ"),
        types.BotCommand(command="clear", description="Очистить историю сообщений")
    ])
    logger.info("Запуск фоновых задач...")
    asyncio.create_task(clean_expired_caches())
    logger.info("Запуск Telegram-бота...")
    await dp.start_polling(bot, allowed_updates=["message", "inline_query", "chosen_inline_result", "callback_query"])

if __name__ == "__main__":
    asyncio.run(main())
